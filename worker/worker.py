import base64, hashlib, html, json, os, re, time, traceback, urllib.error, urllib.request, uuid, unicodedata
from difflib import SequenceMatcher
from datetime import datetime, timezone
from pathlib import Path
import fitz, psycopg, redis
from json_repair import repair_json
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

DB = os.environ["DATABASE_URL"]
R = redis.from_url(os.environ["REDIS_URL"], decode_responses=True)
UP = Path(os.getenv("UPLOAD_DIR", "/app/data/uploads"))
EXPORT_DIR = UP.parent / "exports"
ARTIFACT_DIR = UP.parent / "artifacts"
VERSION = "table-aware-parser-2.0"
MODEL = os.getenv("AI_MODEL", "deepseek-v4-pro")
DEEPSEEK_TEXT_MODEL = MODEL
GPT_TEXT_MODEL = os.getenv("GPT_TEXT_MODEL", "gpt-5.5")
AI_BASE_URL = os.getenv("AI_BASE_URL") or os.getenv("SUB2API_BASE_URL", "http://sub2api:8080")
AI_URL = AI_BASE_URL.rstrip("/") + "/v1/chat/completions"
AI_KEY = os.getenv("AI_API_KEY") or os.environ.get("SUB2API_API_KEY", "")
IMAGE_BASE_URL = os.getenv("SUB2API_BASE_URL", "http://sub2api:8080")
IMAGE_URL = IMAGE_BASE_URL.rstrip("/") + "/v1/images/generations"
IMAGE_KEY = os.getenv("SUB2API_API_KEY", "")
IMAGE_MODEL = os.getenv("IMAGE_MODEL", "gpt-image-2")
AI_THINKING = os.getenv("AI_THINKING", "disabled").strip().lower()
BILLING_DB = os.environ.get("SUB2API_USAGE_DATABASE_URL", "")
INPUT_PRICE = 0.435 / 1_000_000
OUTPUT_PRICE = 0.87 / 1_000_000
TYPE_SET = {"qualification", "disqualification", "scoring", "deadline", "deposit", "deliverable", "technical", "commercial", "other"}
DELIVERY_ARCHETYPES = {
    "goods": "货物采购",
    "equipment_integration": "设备采购与系统集成",
    "software": "软件与信息化建设",
    "professional_service": "专业服务",
    "operation_service": "运营服务",
    "construction": "工程施工",
    "mixed": "混合型项目",
}
SCORING_ROUTE_LABELS = {
    "technical_solution": "技术方案正文",
    "technical_parameter": "技术参数响应表",
    "commercial_response": "商务响应",
    "qualification_evidence": "资格与证明材料",
    "pricing_policy": "价格政策与声明",
    "compliance_response": "实质性响应",
    "evaluation_rule": "评审规则提醒",
}
CAPABILITY_MODULES = {
    "requirements_analysis": {
        "name": "需求理解与范围界定",
        "keywords": ("需求", "背景", "现状", "范围", "目标"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "从采购目标、对象、边界、约束和验收结果五方面建立需求基线。",
        "artifacts": ["需求响应矩阵", "项目范围边界表"],
    },
    "overall_solution": {
        "name": "总体方案与技术路线",
        "keywords": ("总体", "架构", "技术路线", "实施路径", "方案"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "用总体目标、分层设计、实施主线和关键控制点串联各专业工作。",
        "artifacts": ["总体架构图", "实施路线图"],
    },
    "organization_resource": {
        "name": "组织与资源配置",
        "keywords": ("组织", "团队", "人员", "岗位", "资源"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "按职责、接口、授权、替补和升级路径配置项目组织。",
        "artifacts": ["组织架构图", "岗位职责表"],
    },
    "schedule_control": {
        "name": "进度与里程碑控制",
        "keywords": ("进度", "工期", "计划", "里程碑", "节点"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "将总期限分解为阶段、依赖和里程碑，并设置监测与纠偏机制。",
        "artifacts": ["进度计划表", "里程碑检查表"],
    },
    "quality_control": {
        "name": "全过程质量控制",
        "keywords": ("质量", "检查", "检验", "复核", "审核"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "设置事前预防、过程检查、成果复核和不符合项整改。",
        "artifacts": ["质量检查表", "不符合项整改台账"],
    },
    "procurement_supply": {
        "name": "采购与供货管理",
        "keywords": ("采购", "供货", "备货", "厂商", "供应链"),
        "archetypes": ("goods", "equipment_integration", "mixed"),
        "method": "控制选型、订单、生产跟踪、出厂检查和异常替代。",
        "artifacts": ["采购控制表", "供货进度跟踪表"],
    },
    "packaging_logistics": {
        "name": "包装运输与到货",
        "keywords": ("包装", "运输", "物流", "装卸", "到货"),
        "archetypes": ("goods", "equipment_integration", "mixed"),
        "method": "按货物特性规划包装、装卸、在途跟踪、交接和货损处置。",
        "artifacts": ["运输控制表", "到货清点表"],
    },
    "site_implementation": {
        "name": "现场实施与施工组织",
        "keywords": ("现场", "施工", "改造", "安装", "进场"),
        "archetypes": ("equipment_integration", "construction", "mixed"),
        "method": "从现场条件、资源进场、工序衔接、作业控制和成品保护组织实施。",
        "artifacts": ["现场实施流程图", "工序检查表"],
    },
    "system_integration": {
        "name": "系统集成与联调",
        "keywords": ("系统", "集成", "接口", "调试", "联调", "试运行"),
        "archetypes": ("equipment_integration", "software", "mixed"),
        "method": "按接口确认、单元测试、分系统联调、端到端验证和试运行推进。",
        "artifacts": ["接口矩阵", "联调测试用例表"],
    },
    "software_delivery": {
        "name": "软件与数据交付",
        "keywords": ("软件", "平台", "数据", "开发", "部署", "接口"),
        "archetypes": ("software", "mixed"),
        "method": "覆盖需求、设计、开发配置、数据接口、测试、上线和版本管理。",
        "artifacts": ["功能清单", "测试用例与问题单"],
    },
    "service_operation": {
        "name": "运营与专业服务执行",
        "keywords": ("运营", "服务", "作业", "日常", "巡检", "考核"),
        "archetypes": ("professional_service", "operation_service", "mixed"),
        "method": "将服务对象转化为岗位动作、频次、检查标准、记录和考核结果。",
        "artifacts": ["服务流程图", "服务记录表"],
    },
    "training_transfer": {
        "name": "培训与知识转移",
        "keywords": ("培训", "课程", "实操", "知识转移"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "按对象、场景、课程、实操、考核和补训形成知识转移。",
        "artifacts": ["培训计划表", "培训考核记录"],
    },
    "acceptance_handover": {
        "name": "验收与成果移交",
        "keywords": ("验收", "交付", "移交", "成果"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "明确验收对象、依据、准备、实施、整改复验和移交资料。",
        "artifacts": ["验收检查表", "交付资料清单"],
    },
    "aftersales_maintenance": {
        "name": "售后与持续运维",
        "keywords": ("售后", "运维", "质保", "维护", "响应"),
        "archetypes": ("goods", "equipment_integration", "software", "operation_service", "mixed"),
        "method": "按受理、分级、派单、处理、验证、关闭和复盘管理服务。",
        "artifacts": ["服务响应流程", "运维工单台账"],
    },
    "safety_confidentiality": {
        "name": "安全、保密与数据保护",
        "keywords": ("安全", "保密", "数据保护", "隐私", "权限"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "从人员、场所、设备、网络、数据和制度控制安全与保密风险。",
        "artifacts": ["安全检查表", "保密责任矩阵"],
    },
    "risk_emergency": {
        "name": "风险与应急管理",
        "keywords": ("风险", "应急", "异常", "故障", "预案"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "识别风险场景、触发条件、预防监测、响应升级和恢复验证。",
        "artifacts": ["风险矩阵", "应急处置流程"],
    },
    "document_delivery": {
        "name": "文档与过程证据管理",
        "keywords": ("文档", "资料", "记录", "台账", "报告"),
        "archetypes": tuple(DELIVERY_ARCHETYPES),
        "method": "统一文档编号、版本、审批、归档和交付，形成可追溯证据链。",
        "artifacts": ["文档交付清单", "过程记录台账"],
    },
}
SECTION_RETRY_DELAYS = (30, 120, 300, 900)
DELAYED_QUEUE = "ai_bid:delayed_jobs"
KEYWORDS = {
    "废标": 12, "无效": 12, "资格": 10, "评分": 10, "评审": 8, "投标保证金": 10,
    "响应文件": 7, "截止时间": 9, "截止": 6, "★": 8, "必须": 6, "应当": 4,
    "联合体": 7, "中小企业": 7, "履约": 6, "交付": 5, "技术要求": 6, "采购需求": 5,
}


def make_chat_body(messages, max_tokens, temperature=0, model_override=None):
    selected_model = model_override or MODEL
    payload = {
        "model": selected_model,
        "stream": False,
        "max_tokens": max_tokens,
        "messages": messages,
    }
    if selected_model.startswith("deepseek-v4"):
        payload["thinking"] = {
            "type": "enabled" if AI_THINKING == "enabled" else "disabled"
        }
    if AI_THINKING != "enabled":
        payload["temperature"] = temperature
    return json.dumps(payload, ensure_ascii=False).encode("utf-8")


CATEGORY_KEYWORDS = {
    "qualification": {"资格": 10, "资质": 10, "联合体": 8, "中小企业": 8, "供应商应具备": 12, "投标人应具备": 12},
    "disqualification": {"废标": 12, "无效响应": 12, "无效投标": 12, "否决": 10, "实质性": 8, "不予通过": 10},
    "scoring": {"评分": 12, "评审标准": 12, "分值": 10, "评分办法": 12, "加分": 8},
    "milestones": {"截止时间": 12, "开标时间": 10, "提交时间": 10, "保证金": 10, "履约保证": 10},
    "technical": {"技术要求": 12, "采购需求": 10, "参数": 8, "性能": 8, "功能要求": 10},
    "commercial": {"付款": 10, "合同": 8, "交付": 10, "服务期限": 10, "工期": 10, "验收": 8},
}

CATEGORY_LABELS = {"qualification": "资格条件", "disqualification": "废标/无效响应", "scoring": "评分标准", "milestones": "时间节点/保证金", "technical": "技术要求", "commercial": "商务/合同/付款/交付"}


def blocks_docx(path):
    doc = Document(path); out = []; paragraph_no = table_no = 0
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph_no += 1; p = Paragraph(child, doc); text = p.text.strip()
            if text: out.append({"id": f"paragraph-{paragraph_no}", "kind": "paragraph", "text": text, "position": len(out) + 1})
        elif isinstance(child, CT_Tbl):
            table_no += 1; table = Table(child, doc)
            for row_no, row in enumerate(table.rows, 1):
                cells = [unicodedata.normalize("NFKC", cell.text).strip() for cell in row.cells]
                if any(cells): out.append({"id": f"table-{table_no}-row-{row_no}", "kind": "table-row", "tableId": f"table-{table_no}", "row": row_no, "cells": cells, "text": " | ".join(cells), "position": len(out) + 1})
    return out, None, "docx"


def blocks_pdf(path):
    doc = fitz.open(path); out = []
    for page_no, page in enumerate(doc, 1):
        for block_no, block in enumerate(page.get_text("blocks"), 1):
            text = unicodedata.normalize("NFKC", " ".join(str(block[4]).split())).strip()
            if text: out.append({"id": f"p{page_no}-b{block_no}", "kind": "paragraph", "text": text, "page": page_no, "bbox": [round(x, 2) for x in block[:4]]})
        if hasattr(page, "find_tables"):
            try:
                for table_no, table in enumerate(page.find_tables().tables, 1):
                    rows = table.extract()
                    header = " | ".join(
                        unicodedata.normalize("NFKC", str(cell or "")).strip()
                        for row in rows[:2] for cell in row
                        if str(cell or "").strip()
                    )[:800]
                    for row_no, row in enumerate(rows, 1):
                        cells = [
                            unicodedata.normalize("NFKC", str(cell or "")).strip()
                            for cell in row
                        ]
                        row_text = " | ".join(cell for cell in cells if cell)
                        if not row_text:
                            continue
                        out.append({
                            "id": f"p{page_no}-t{table_no}-r{row_no}",
                            "kind": "table-row",
                            "tableId": f"p{page_no}-t{table_no}",
                            "row": row_no,
                            "cells": cells,
                            "tableHeader": header,
                            "text": row_text,
                            "page": page_no,
                            "bbox": [round(x, 2) for x in table.bbox],
                        })
            except Exception as exc:
                print(json.dumps({"event": "pdf_table_parse_skipped", "page": page_no, "error": str(exc)[:160]}), flush=True)
    return out, len(doc), "pdf"


def clean_parameter_text(value):
    value = unicodedata.normalize("NFKC", str(value or ""))
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\s*\n\s*", "\n", value)
    value = re.sub(r"-第\s*\d+\s*页-", "", value)
    return value.strip()


def pdf_parameter_row_intervals(page):
    values = []
    for drawing in page.get_drawings():
        for item in drawing["items"]:
            if item[0] != "re":
                continue
            rect = item[1]
            if (
                rect.width <= 1.2
                and rect.height >= 6
                and abs(rect.x0 - page.rect.width * 0.4655) <= 3
            ):
                top = max(0.0, float(rect.y0))
                bottom = min(float(page.rect.height), float(rect.y1))
                if bottom - top >= 6:
                    values.append((top, bottom, float(rect.y0), float(rect.y1)))
    grouped = []
    for value in sorted(values):
        if (
            not grouped
            or abs(value[0] - grouped[-1][0]) > 1.2
            or abs(value[1] - grouped[-1][1]) > 1.2
        ):
            grouped.append(value)
    return grouped


def split_parameter_clauses(text):
    note_match = re.search(r"[\(（]\s*以上带[“\"]?[★▲]", text)
    note = clean_parameter_text(text[note_match.start():]) if note_match else ""
    body = text[:note_match.start()] if note_match else text
    body = re.sub(
        r"([;；。])\s*([★▲]?\s*\d+\s*[、.．])",
        r"\1\n\2",
        body,
    )
    lines = [line.strip() for line in body.splitlines() if line.strip()]
    clauses = []
    current = []
    marker = ""
    number = ""
    start_pattern = re.compile(r"^\s*([★▲]?)\s*(\d+)\s*[、.．]\s*(.*)$")
    for line in lines:
        match = start_pattern.match(line)
        if match:
            if current:
                clauses.append({
                    "marker": marker,
                    "number": number,
                    "text": clean_parameter_text(" ".join(current)),
                })
            marker, number = match.group(1), match.group(2)
            current = [match.group(3)]
        elif current:
            current.append(line)
        else:
            current = [line]
    if current:
        clauses.append({
            "marker": marker,
            "number": number,
            "text": clean_parameter_text(" ".join(current)),
        })
    return [item for item in clauses if item["text"]], note


def extract_pdf_parameter_matrix(path):
    doc = fitz.open(path)
    start = end = None
    for index, page in enumerate(doc):
        text = clean_parameter_text(page.get_text())
        if start is None and "序号 标的名称 技术参数" in text.replace("\n", " "):
            start = index
        if start is not None and (
            re.search(r"3\.3\.\s*商务要求", text)
            or ("商务要求名称" in text and "商务要求内容" in text)
        ):
            end = index
            break
    if start is None:
        return []
    end = end or len(doc)
    products = []
    current_product = None
    pending_fragments = []
    for page_index in range(start + 1, end):
        page = doc[page_index]
        intervals = pdf_parameter_row_intervals(page)
        if not intervals:
            continue
        width = page.rect.width
        x0, x1, x2, x3 = (
            width * 0.331,
            width * 0.382,
            width * 0.466,
            width * 0.905,
        )
        for top, bottom, raw_top, _raw_bottom in intervals:
            sequence = clean_parameter_text(
                page.get_textbox(fitz.Rect(x0, top, x1, bottom))
            ).replace("\n", "")
            name = clean_parameter_text(
                page.get_textbox(fitz.Rect(x1, top, x2, bottom))
            ).replace("\n", "")
            parameters = clean_parameter_text(
                page.get_textbox(fitz.Rect(x2, top, x3, bottom))
            )
            match = re.search(r"\d+", sequence)
            if match and name:
                product = {
                    "page": page_index + 1,
                    "productNo": int(match.group()),
                    "productName": name,
                    "fragments": [*pending_fragments, parameters],
                }
                pending_fragments = []
                products.append(product)
                current_product = product
                continue
            if len(parameters) < 4:
                continue
            if raw_top < 0 and current_product is not None and not pending_fragments:
                current_product["fragments"].append(parameters)
            elif raw_top > 0:
                pending_fragments = [parameters]
            elif pending_fragments:
                pending_fragments.append(parameters)
    items = []
    for product in products:
        parameter_text = clean_parameter_text("\n".join(product["fragments"]))
        clauses, proof_note = split_parameter_clauses(parameter_text)
        for clause in clauses:
            items.append({
                "itemIndex": len(items) + 1,
                "productNo": product["productNo"],
                "productName": product["productName"],
                "parameterNo": clause["number"],
                "marker": clause["marker"],
                "requirement": clause["text"],
                "sourcePage": product["page"],
                "proofRequirement": proof_note if clause["marker"] else "",
            })
    return items


def replace_parameter_matrix(conn, project_id, items):
    conn.execute(
        "DELETE FROM technical_parameter_items WHERE project_id=%s",
        (project_id,),
    )
    for item in items:
        conn.execute(
            """INSERT INTO technical_parameter_items(
                id,project_id,item_index,product_no,product_name,parameter_no,
                marker,requirement_text,source_page,proof_requirement
            ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
            (
                str(uuid.uuid4()),
                project_id,
                item["itemIndex"],
                item["productNo"],
                item["productName"],
                item["parameterNo"],
                item["marker"],
                item["requirement"],
                item["sourcePage"],
                item["proofRequirement"],
            ),
        )


def norm(value):
    return re.sub(r"\s+", "", unicodedata.normalize("NFKC", str(value or ""))).strip()


def evidence_matches(quote, source):
    q, text = norm(quote), norm(source)
    if not q or not text: return False
    if q in text: return True
    if len(q) < 8: return False
    matcher = SequenceMatcher(None, q, text, autojunk=False)
    return matcher.ratio() >= 0.78 or matcher.find_longest_match().size / len(q) >= 0.85


def make_chunks(blocks, limit=4500):
    chunks = []; current = []; size = 0
    for block in blocks:
        line = f"[{block.get('id')}] {block.get('text', '')}".strip()
        if not line: continue
        if current and size + len(line) + 1 > limit:
            chunks.append(current); current = []; size = 0
        current.append(block); size += len(line) + 1
    if current: chunks.append(current)
    return chunks


def chunk_score(chunk):
    text = "\n".join(str(x.get("text", "")) for x in chunk)
    score = sum(text.count(word) * weight for word, weight in KEYWORDS.items())
    return score + min(len(text), 4500) / 4500


def category_score(chunk, weights):
    text = "\n".join(str(x.get("text", "")) for x in chunk)
    return sum(text.count(word) * weight for word, weight in weights.items())


def select_chunks(blocks, count=16):
    chunks = make_chunks(blocks)
    selected = []
    seen = set()
    category_quotas = {
        "scoring": 6,
        "technical": 4,
        "qualification": 2,
        "disqualification": 2,
        "commercial": 2,
        "milestones": 1,
    }
    for category, quota in category_quotas.items():
        weights = CATEGORY_KEYWORDS[category]
        ranked = sorted(enumerate(chunks), key=lambda item: category_score(item[1], weights), reverse=True)
        added = 0
        for item in ranked:
            if category_score(item[1], weights) <= 0:
                break
            if item[0] in seen:
                continue
            selected.append(item)
            seen.add(item[0])
            added += 1
            if added >= quota or len(selected) >= count:
                break
        if len(selected) >= count:
            break
    for item in sorted(enumerate(chunks), key=lambda item: chunk_score(item[1]), reverse=True):
        if len(selected) >= count: break
        if item[0] not in seen: selected.append(item); seen.add(item[0])
    output = []
    for index, chunk in selected[:count]:
        focus_scores = sorted(((category_score(chunk, weights), name) for name, weights in CATEGORY_KEYWORDS.items()), reverse=True)
        focus = "、".join(CATEGORY_LABELS[name] for score, name in focus_scores[:2] if score > 0) or "关键响应要求"
        output.append((index + 1, chunk, focus))
    return output


COVERAGE_CATEGORY_SIGNALS = {
    "scoring": (
        "\u8bc4\u5206\u529e\u6cd5", "\u8bc4\u5ba1\u5185\u5bb9",
        "\u8bc4\u5ba1\u6807\u51c6", "\u5f97\u5206",
    ),
    "technical": (
        "\u6280\u672f\u53c2\u6570", "\u6027\u80fd\u6307\u6807",
        "\u6280\u672f\u8981\u6c42\u540d\u79f0",
    ),
    "qualification": (
        "\u8d44\u683c\u5ba1\u67e5", "\u8d44\u683c\u8981\u6c42",
        "\u8d44\u683c\u8bc1\u660e",
    ),
    "disqualification": (
        "\u7b26\u5408\u6027\u5ba1\u67e5", "\u5b9e\u8d28\u6027\u8981\u6c42",
        "\u65e0\u6548\u6295\u6807", "\u5e9f\u6807",
    ),
    "commercial": (
        "\u6700\u9ad8\u9650\u4ef7", "\u4ed8\u6b3e\u65b9\u5f0f",
        "\u4ea4\u4ed8\u671f\u9650", "\u6295\u6807\u4fdd\u8bc1\u91d1",
        "\u5c65\u7ea6\u4fdd\u8bc1\u91d1", "\u5546\u52a1\u8981\u6c42",
    ),
}

REQUIREMENT_COVERAGE_CATEGORY = {
    "scoring": "scoring",
    "technical": "technical",
    "qualification": "qualification",
    "disqualification": "disqualification",
    "commercial": "commercial",
    "deadline": "commercial",
    "deposit": "commercial",
    "deliverable": "commercial",
}


def coverage_categories(text, table_cells=None):
    value = unicodedata.normalize("NFKC", str(text or ""))
    categories = {
        category for category, signals in COVERAGE_CATEGORY_SIGNALS.items()
        if any(signal in value for signal in signals)
    }
    cells = table_cells or []
    joined_cells = " ".join(str(cell or "") for cell in cells)
    if (
        re.search(r"(?:\u4e3b\u89c2|\u5ba2\u89c2)", joined_cells)
        and re.search(r"\d+(?:\.\d+)?\s*\u5206?", joined_cells)
    ) or re.search(r"(?:\u5f97|\u6263)\s*\d+(?:\.\d+)?\s*\u5206", value):
        categories.add("scoring")
    return categories


def evidence_page(block_id):
    match = re.match(r"p(\d+)-", str(block_id or ""))
    return int(match.group(1)) if match else None


def build_coverage_audit(blocks, requirements):
    table_groups = {}
    page_groups = {}
    for block in blocks:
        block_id = str(block.get("id") or "")
        page = block.get("page")
        if block.get("kind") == "table-row" and block.get("tableId"):
            group = table_groups.setdefault(
                str(block["tableId"]),
                {"kind": "table", "page": page, "blocks": [], "texts": [], "cells": [], "rowCells": []},
            )
            group["blocks"].append(block_id)
            group["texts"].append(str(block.get("text") or ""))
            group["cells"].extend(block.get("cells") or [])
            group["rowCells"].append(block.get("cells") or [])
            continue
        categories = coverage_categories(block.get("text"))
        for category in categories:
            key = (page, category)
            group = page_groups.setdefault(
                key,
                {"kind": "page", "page": page, "category": category, "blocks": [], "texts": []},
            )
            group["blocks"].append(block_id)
            group["texts"].append(str(block.get("text") or ""))

    zones = []
    for table_id, group in table_groups.items():
        combined = "\n".join(group["texts"])
        categories = coverage_categories(combined, group["cells"])
        if not categories:
            continue
        rows = len(group["blocks"])
        for category in categories:
            matching_indexes = [
                index for index, text in enumerate(group["texts"])
                if category in coverage_categories(text, group["rowCells"][index])
            ]
            if not matching_indexes:
                matching_indexes = list(range(rows))
            relevant_blocks = [group["blocks"][index] for index in matching_indexes]
            relevant_texts = [group["texts"][index] for index in matching_indexes]
            header_rows = sum(
                1 for text in relevant_texts[:2]
                if len(norm(text)) < 120 and any(
                    token in text for token in (
                        "\u5e8f\u53f7", "\u8bc4\u5ba1\u5185\u5bb9",
                        "\u8bc4\u5ba1\u6807\u51c6", "\u6280\u672f\u53c2\u6570",
                    )
                )
            )
            item_count = max(1, len(relevant_blocks) - header_rows)
            zones.append({
                "id": f"{table_id}:{category}",
                "kind": "table",
                "category": category,
                "page": group["page"],
                "blockIds": relevant_blocks,
                "itemCount": item_count,
                "sample": re.sub(r"\s+", " ", "\n".join(relevant_texts))[:220],
                "coveredItemCount": 0,
            })
    table_zone_keys = {(zone["page"], zone["category"]) for zone in zones}
    for group in page_groups.values():
        if (group["page"], group["category"]) in table_zone_keys:
            continue
        zones.append({
            "id": f"page-{group['page']}:{group['category']}",
            "kind": "page",
            "category": group["category"],
            "page": group["page"],
            "blockIds": group["blocks"],
            "itemCount": 1,
            "sample": re.sub(r"\s+", " ", " ".join(group["texts"]))[:220],
            "coveredItemCount": 0,
        })

    requirement_refs = []
    for index, requirement in enumerate(requirements):
        category = REQUIREMENT_COVERAGE_CATEGORY.get(str(requirement.get("type") or ""))
        if not category:
            continue
        evidence = requirement.get("evidence") or []
        if isinstance(evidence, dict):
            evidence = [evidence]
        block_ids = {
            str(item.get("blockId") or "") for item in evidence
            if isinstance(item, dict)
        }
        pages = {page for page in (evidence_page(item) for item in block_ids) if page}
        requirement_refs.append({
            "id": str(requirement.get("id") or index),
            "category": category,
            "blockIds": block_ids,
            "pages": pages,
        })

    assigned = set()
    for zone in zones:
        exact = [
            item for item in requirement_refs
            if item["id"] not in assigned
            and item["category"] == zone["category"]
            and item["blockIds"].intersection(zone["blockIds"])
        ]
        for item in exact[:zone["itemCount"]]:
            assigned.add(item["id"])
            zone["coveredItemCount"] += 1
    for zone in zones:
        same_page = [
            item for item in requirement_refs
            if item["id"] not in assigned
            and item["category"] == zone["category"]
            and zone["page"] in item["pages"]
        ]
        available = max(0, zone["itemCount"] - zone["coveredItemCount"])
        for item in same_page[:available]:
            assigned.add(item["id"])
            zone["coveredItemCount"] += 1

    categories = {}
    for category in COVERAGE_CATEGORY_SIGNALS:
        category_zones = [zone for zone in zones if zone["category"] == category]
        candidates = sum(zone["itemCount"] for zone in category_zones)
        covered = sum(zone["coveredItemCount"] for zone in category_zones)
        categories[category] = {
            "candidateItems": candidates,
            "coveredItems": covered,
            "possibleMissing": max(0, candidates - covered),
            "coverageRate": round(covered / candidates * 100) if candidates else 100,
        }
    total_candidates = sum(item["candidateItems"] for item in categories.values())
    total_covered = sum(item["coveredItems"] for item in categories.values())
    missing = [
        {
            "category": zone["category"],
            "page": zone["page"],
            "kind": zone["kind"],
            "possibleMissing": zone["itemCount"] - zone["coveredItemCount"],
            "sample": zone["sample"],
        }
        for zone in zones if zone["coveredItemCount"] < zone["itemCount"]
    ]
    return {
        "parserVersion": VERSION,
        "candidateItems": total_candidates,
        "coveredItems": total_covered,
        "possibleMissing": max(0, total_candidates - total_covered),
        "coverageRate": round(total_covered / total_candidates * 100) if total_candidates else 100,
        "categories": categories,
        "missingSamples": sorted(
            missing,
            key=lambda item: (item["page"] is None, item["page"] or 0, item["category"]),
        )[:20],
        "generatedAt": int(time.time()),
    }


def parse_json(content):
    text = str(content or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text)
    start, end = text.find("{"), text.rfind("}")
    if start < 0 or end <= start: raise ValueError("模型未返回JSON对象")
    candidate = text[start:end + 1]
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        repaired = repair_json(candidate, return_objects=True)
        if not isinstance(repaired, dict): raise ValueError("JSON自动修复后仍不是对象")
        return repaired


def fetch_billing(client_request_id):
    if not BILLING_DB or not client_request_id: return None
    request_id = "client:" + client_request_id
    for delay in (0, 0.25, 0.75, 1.5):
        if delay: time.sleep(delay)
        try:
            with psycopg.connect(BILLING_DB, connect_timeout=3) as conn:
                row = conn.execute("SELECT input_tokens,output_tokens,total_cost FROM ai_bid_usage WHERE request_id=%s", (request_id,)).fetchone()
                if row: return int(row[0]), int(row[1]), float(row[2])
        except Exception as exc:
            print(json.dumps({"event": "billing_lookup_failed", "error": str(exc)[:160]}), flush=True)
            break
    return None


def save_run(run_id, project_id, job_id, chunk_no, status, input_tokens, output_tokens, duration, attempt, error=None, actual_cost=None, run_type="extraction", model_override=None):
    cost = actual_cost if actual_cost is not None else input_tokens * INPUT_PRICE + output_tokens * OUTPUT_PRICE
    with psycopg.connect(DB) as conn:
        conn.execute("""INSERT INTO ai_runs(id,project_id,job_id,chunk_number,run_type,model,status,input_tokens,output_tokens,cost_usd,duration_ms,retries,error_message,finished_at) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,now())""", (run_id, project_id, job_id, chunk_no, run_type, model_override or MODEL, status, input_tokens, output_tokens, cost, duration, attempt, str(error)[:500] if error else None))
        conn.commit()


def call_model(chunk_no, chunk, focus, project_id, job_id):
    block_lines = "\n".join(f"[{b.get('id')}] {b.get('text', '')}" for b in chunk)
    system = """你是招标文件条款提取助手。只根据给定原文提取关键要求，不推测。目录、资格审查表、评分表、采购需求和合同模板中明确写出的条件同样必须提取；“不收取”“应提供”“不得”“限期”等都是有效要求，不能因上下文不完整而返回空。输出严格JSON对象，不要Markdown。每项必须给出可在原文逐字找到的短证据。类型仅限qualification,disqualification,scoring,deadline,deposit,deliverable,technical,commercial,other。"""
    user = f"""本片段重点检查：{focus}。逐条检查明确条件后再输出，不要只做摘要。格式：{{\"requirements\":[{{\"type\":\"qualification\",\"title\":\"简短标题\",\"normalizedValue\":\"规范化要求\",\"mandatory\":true,\"evidence\":{{\"blockId\":\"原块ID\",\"quote\":\"原文短句\"}}}}]}}。最多8项；没有就返回空数组。\n\n{block_lines}"""
    user += "\n\nTreat each scoring-table row as an independent requirement. Extract its factor, score, deduction rule, evidence and linked response file separately; do not stop after the price formula."
    body = make_chat_body([{"role": "system", "content": system}, {"role": "user", "content": user}], 2600, 0)
    last_error = None
    for attempt in range(3):
        started = time.monotonic(); run_id = str(uuid.uuid4()); input_tokens = output_tokens = 0; actual_cost = None
        try:
            request = urllib.request.Request(AI_URL, data=body, headers={"Authorization": f"Bearer {AI_KEY}", "Content-Type": "application/json"}, method="POST")
            with urllib.request.urlopen(request, timeout=300) as response:
                client_request_id = response.headers.get("X-Client-Request-Id")
                payload = json.loads(response.read().decode("utf-8"))
            usage = payload.get("usage") or {}; input_tokens = int(usage.get("prompt_tokens") or 0); output_tokens = int(usage.get("completion_tokens") or 0); actual_cost = None
            billing = fetch_billing(client_request_id)
            if billing: input_tokens, output_tokens, actual_cost = billing
            try:
                parsed = parse_json(payload["choices"][0]["message"]["content"])
            except Exception as exc:
                last_error = exc; save_run(run_id, project_id, job_id, chunk_no, "failed", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), attempt, exc, actual_cost)
                if attempt < 2: time.sleep(2 * (attempt + 1)); continue
                break
            save_run(run_id, project_id, job_id, chunk_no, "succeeded", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), attempt, actual_cost=actual_cost)
            return parsed
        except Exception as exc:
            last_error = exc; save_run(run_id, project_id, job_id, chunk_no, "failed", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), attempt, exc, actual_cost)
            if attempt < 2: time.sleep(2 * (attempt + 1))
    raise RuntimeError(f"AI片段{chunk_no}提取失败: {last_error}")


def validate_requirements(payload, chunk_no, chunk):
    block_map = {str(b.get("id")): str(b.get("text", "")) for b in chunk}; output = []
    context_map = {str(block.get("id")): "".join(str(x.get("text", "")) for x in chunk[index:index + 3]) for index, block in enumerate(chunk)}
    values = payload.get("requirements", []) if isinstance(payload, dict) else []
    if not isinstance(values, list): return output
    for item in values[:24]:
        if not isinstance(item, dict): continue
        kind = str(item.get("type", "other")); title = str(item.get("title", "")).strip()[:200]; value = str(item.get("normalizedValue", "")).strip()[:2000]
        evidence = item.get("evidence") or {}; block_id = str(evidence.get("blockId", "")); quote = str(evidence.get("quote", "")).strip()[:1000]
        if kind not in TYPE_SET or not title or not value or block_id not in block_map or not quote: continue
        if not evidence_matches(quote, context_map[block_id]): continue
        output.append({"type": kind, "title": title, "value": value, "mandatory": bool(item.get("mandatory")), "evidence": [{"blockId": block_id, "quote": quote, "chunk": chunk_no}]})
    return output


def call_audit_model(batch_no, items, project_id, job_id):
    system = """你是投标要求二次审核员。仅依据提供的招标原文上下文，复核首次AI提取结果。
逐项判断：
1. auto_pass：要求及限定条件被原文明确支持；
2. needs_review：基本相关但存在歧义、遗漏限定、概括过度，必须人工确认；
3. rejected：证据不支持该结论、把目录/说明误当要求，或明显误提取。
不要把签字、盖章、纸质装订等只能线下核验的事项判定为已满足。输出严格JSON，不要Markdown。每个输入ID必须且只能返回一次。"""
    user = json.dumps({
        "outputSchema": {"reviews": [{"id": "输入ID", "status": "auto_pass|needs_review|rejected", "reason": "简短、具体的判断依据", "suggestion": "人工应该核对或修正什么；通过时可为空", "confidence": 0.0}]},
        "requirements": items,
    }, ensure_ascii=False)
    body = make_chat_body([{"role": "system", "content": system}, {"role": "user", "content": user}], 1800, 0)
    last_error = None
    for attempt in range(3):
        started = time.monotonic(); run_id = str(uuid.uuid4()); input_tokens = output_tokens = 0; actual_cost = None
        try:
            request = urllib.request.Request(AI_URL, data=body, headers={"Authorization": f"Bearer {AI_KEY}", "Content-Type": "application/json"}, method="POST")
            with urllib.request.urlopen(request, timeout=300) as response:
                client_request_id = response.headers.get("X-Client-Request-Id")
                payload = json.loads(response.read().decode("utf-8"))
            usage = payload.get("usage") or {}; input_tokens = int(usage.get("prompt_tokens") or 0); output_tokens = int(usage.get("completion_tokens") or 0)
            billing = fetch_billing(client_request_id)
            if billing: input_tokens, output_tokens, actual_cost = billing
            parsed = parse_json(payload["choices"][0]["message"]["content"])
            reviews = parsed.get("reviews") if isinstance(parsed, dict) else None
            if not isinstance(reviews, list): raise ValueError("模型未返回reviews数组")
            save_run(run_id, project_id, job_id, batch_no, "succeeded", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), attempt, actual_cost=actual_cost, run_type="audit")
            return reviews
        except Exception as exc:
            last_error = exc
            save_run(run_id, project_id, job_id, batch_no, "failed", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), attempt, exc, actual_cost, "audit")
            if attempt < 2: time.sleep(2 * (attempt + 1))
    raise RuntimeError(f"AI审核批次{batch_no}失败: {last_error}")


def process_audit(job):
    job_id, project_id = job["jobId"], job["projectId"]
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s", (job_id,))
        conn.execute("UPDATE projects SET status='auditing',progress=5,error_message=NULL,updated_at=now() WHERE id=%s", (project_id,))
        rows = conn.execute("""SELECT id,type,title,normalized_value,mandatory,evidence FROM requirements WHERE project_id=%s AND review_status<>'rejected' ORDER BY created_at""", (project_id,)).fetchall()
        doc_row = conn.execute("SELECT blocks FROM documents WHERE project_id=%s", (project_id,)).fetchone()
        conn.commit()
    try:
        if not rows or not doc_row: raise ValueError("没有可复核的条款或解析原文")
        blocks = doc_row[0] if isinstance(doc_row[0], list) else json.loads(doc_row[0])
        block_positions = {str(block.get("id")): index for index, block in enumerate(blocks)}
        candidates = []
        for row in rows:
            evidence = row[5] if isinstance(row[5], list) else json.loads(row[5] or "[]")
            block_id = str(evidence[0].get("blockId", "")) if evidence else ""
            index = block_positions.get(block_id)
            context = ""
            if index is not None:
                context = "\n".join(f"[{block.get('id')}] {block.get('text', '')}" for block in blocks[max(0, index - 1):index + 3])
            candidates.append({"id": str(row[0]), "type": row[1], "title": row[2], "normalizedValue": row[3], "mandatory": bool(row[4]), "evidence": evidence, "sourceContext": context[:5000]})
        batches = [candidates[index:index + 8] for index in range(0, len(candidates), 8)]
        completed = 0; errors = []
        allowed = {"auto_pass", "needs_review", "rejected"}
        for batch_no, batch in enumerate(batches, 1):
            try:
                reviews = call_audit_model(batch_no, batch, project_id, job_id)
                review_map = {str(item.get("id")): item for item in reviews if isinstance(item, dict) and str(item.get("status")) in allowed}
                with psycopg.connect(DB) as conn:
                    for candidate in batch:
                        result = review_map.get(candidate["id"])
                        if result:
                            status = str(result["status"])
                            reason = str(result.get("reason") or "")[:1000]
                            suggestion = str(result.get("suggestion") or "")[:1000] or None
                            try: confidence = max(0.0, min(1.0, float(result.get("confidence", 0.5))))
                            except (TypeError, ValueError): confidence = 0.5
                        else:
                            status, reason, suggestion, confidence = "needs_review", "模型未返回本条审核结果", "请人工核对原文证据", 0.0
                        conn.execute("""UPDATE requirements SET ai_review_status=%s,ai_review_reason=%s,ai_review_suggestion=%s,ai_review_confidence=%s,ai_reviewed_at=now(),updated_at=now() WHERE id=%s""", (status, reason, suggestion, confidence, candidate["id"]))
                    conn.commit()
                completed += len(batch)
            except Exception as exc:
                errors.append(str(exc))
            with psycopg.connect(DB) as conn:
                conn.execute("UPDATE projects SET progress=%s,updated_at=now() WHERE id=%s", (5 + int(batch_no / len(batches) * 90), project_id)); conn.commit()
        if not completed: raise RuntimeError("；".join(errors) or "二次审核没有完成任何批次")
        with psycopg.connect(DB) as conn:
            unresolved = conn.execute("SELECT count(*) FROM requirements WHERE project_id=%s AND review_status<>'rejected' AND ai_review_status='unreviewed'", (project_id,)).fetchone()[0]
            error_message = ("部分审核批次失败：" + "；".join(errors))[:500] if errors else None
            conn.execute("UPDATE projects SET status='reviewing',progress=100,error_message=%s,updated_at=now() WHERE id=%s", (error_message, project_id))
            conn.execute("UPDATE jobs SET status='succeeded',finished_at=now(),error_message=%s WHERE id=%s", (error_message, job_id))
            conn.commit()
        print(json.dumps({"event": "audited", "projectId": project_id, "requirements": len(candidates), "completed": completed, "unreviewed": unresolved, "batches": len(batches), "partialErrors": len(errors)}), flush=True)
    except Exception as exc:
        fail_job(job_id, project_id, exc)


def enqueue_extract(project_id):
    job_id = str(uuid.uuid4())
    with psycopg.connect(DB) as conn:
        conn.execute("INSERT INTO jobs(id,project_id,type,status) VALUES(%s,%s,'extract','queued')", (job_id, project_id)); conn.commit()
    R.lpush("ai_bid:jobs", json.dumps({"jobId": job_id, "projectId": project_id, "type": "extract"}))


def enqueue_outline(project_id):
    job_id = str(uuid.uuid4())
    with psycopg.connect(DB) as conn:
        conn.execute("INSERT INTO jobs(id,project_id,type,status) VALUES(%s,%s,'outline','queued')", (job_id, project_id))
        conn.execute("""INSERT INTO outlines(project_id,status,model) VALUES(%s,'generating',%s) ON CONFLICT(project_id) DO UPDATE SET status='generating',error_message=NULL,updated_at=now()""", (project_id, MODEL))
        conn.execute("UPDATE projects SET status='outlining',progress=0,updated_at=now() WHERE id=%s", (project_id,))
        conn.commit()
    R.lpush("ai_bid:jobs", json.dumps({"jobId": job_id, "projectId": project_id, "type": "outline"}))


def process_parse(job):
    job_id, project_id = job["jobId"], job["projectId"]; path = UP / job["storedName"]
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s", (job_id,)); conn.execute("UPDATE projects SET status='parsing',progress=25,updated_at=now(),error_message=NULL WHERE id=%s", (project_id,)); conn.commit()
    try:
        blocks, pages, fmt = blocks_pdf(path) if path.suffix.lower() == ".pdf" else blocks_docx(path); chars = sum(len(x.get("text", "")) for x in blocks)
        parameter_items = extract_pdf_parameter_matrix(path) if fmt == "pdf" else []
        with psycopg.connect(DB) as conn:
            conn.execute("""INSERT INTO documents(project_id,format,page_count,character_count,blocks,parser_version) VALUES(%s,%s,%s,%s,%s::jsonb,%s) ON CONFLICT(project_id) DO UPDATE SET format=excluded.format,page_count=excluded.page_count,character_count=excluded.character_count,blocks=excluded.blocks,parser_version=excluded.parser_version,parsed_at=now()""", (project_id, fmt, pages, chars, json.dumps(blocks, ensure_ascii=False), VERSION)); replace_parameter_matrix(conn, project_id, parameter_items); conn.execute("UPDATE projects SET status='extracting',progress=55,updated_at=now() WHERE id=%s", (project_id,)); conn.execute("UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s", (job_id,)); conn.commit()
        enqueue_extract(project_id)
        print(json.dumps({"event": "parsed", "projectId": project_id, "format": fmt, "blocks": len(blocks), "chars": chars, "technicalParameterItems": len(parameter_items)}), flush=True)
    except Exception as exc:
        fail_job(job_id, project_id, exc)


def process_extract(job):
    job_id, project_id = job["jobId"], job["projectId"]
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s", (job_id,)); conn.execute("UPDATE projects SET status='extracting',progress=60,updated_at=now(),error_message=NULL WHERE id=%s", (project_id,)); row = conn.execute("SELECT blocks FROM documents WHERE project_id=%s", (project_id,)).fetchone(); conn.commit()
    try:
        if not row: raise ValueError("项目尚无解析结果")
        blocks = row[0] if isinstance(row[0], list) else json.loads(row[0]); selected = select_chunks(blocks, 16)
        all_items = []; success = 0; errors = []
        for offset, (chunk_no, chunk, focus) in enumerate(selected):
            try:
                payload = call_model(chunk_no, chunk, focus, project_id, job_id); all_items.extend(validate_requirements(payload, chunk_no, chunk)); success += 1
            except Exception as exc: errors.append(str(exc))
            with psycopg.connect(DB) as conn:
                conn.execute("UPDATE projects SET progress=%s,updated_at=now() WHERE id=%s", (70 + int((offset + 1) / max(len(selected), 1) * 25), project_id)); conn.commit()
        if not success: raise RuntimeError("；".join(errors) or "没有可处理片段")
        unique = {}
        for item in all_items: unique[(item["type"], norm(item["title"]), norm(item["value"]))] = item
        with psycopg.connect(DB) as conn:
            reviewed = conn.execute("SELECT type,title,normalized_value FROM requirements WHERE project_id=%s AND review_status<>'pending'", (project_id,)).fetchall()
            protected = {(row[0], norm(row[1]), norm(row[2])) for row in reviewed}
            conn.execute("DELETE FROM requirements WHERE project_id=%s AND review_status='pending'", (project_id,))
            inserted = 0
            for key, item in unique.items():
                if key in protected: continue
                conn.execute("""INSERT INTO requirements(id,project_id,type,title,normalized_value,mandatory,evidence) VALUES(%s,%s,%s,%s,%s,%s,%s::jsonb)""", (str(uuid.uuid4()), project_id, item["type"], item["title"], item["value"], item["mandatory"], json.dumps(item["evidence"], ensure_ascii=False))); inserted += 1
            total = conn.execute("SELECT count(*) FROM requirements WHERE project_id=%s", (project_id,)).fetchone()[0]
            requirement_rows = conn.execute(
                "SELECT id,type,evidence FROM requirements WHERE project_id=%s AND review_status<>'rejected'",
                (project_id,),
            ).fetchall()
            coverage_audit = build_coverage_audit(
                blocks,
                [
                    {"id": str(value[0]), "type": value[1], "evidence": value[2]}
                    for value in requirement_rows
                ],
            )
            conn.execute(
                "UPDATE documents SET coverage_audit=%s::jsonb WHERE project_id=%s",
                (json.dumps(coverage_audit, ensure_ascii=False), project_id),
            )
            conn.execute("UPDATE projects SET status='reviewing',progress=100,findings_count=%s,updated_at=now(),error_message=%s WHERE id=%s", (total, ("部分片段失败：" + "；".join(errors))[:500] if errors else None, project_id)); conn.execute("UPDATE jobs SET status='succeeded',finished_at=now(),error_message=%s WHERE id=%s", (("；".join(errors))[:500] if errors else None, job_id)); conn.commit()
        print(json.dumps({"event": "extracted", "projectId": project_id, "requirements": total, "inserted": inserted, "protected": len(protected), "chunks": len(selected), "partialErrors": len(errors), "coverage": coverage_audit}), flush=True)
        enqueue_outline(project_id)
    except Exception as exc:
        fail_job(job_id, project_id, exc)


def sanitize_outline_node(
    node, valid_ids, valid_block_ids=None, depth=1, inherited_blocks=None
):
    if not isinstance(node, dict) or depth > 5: return None
    title = str(node.get("title") or "").strip()[:200]
    if not title: return None
    description = str(node.get("description") or "").strip()[:1000]
    requirement_ids = [str(value) for value in (node.get("requirementIds") or []) if str(value) in valid_ids][:30]
    requested_blocks = [
        str(value) for value in (node.get("sourceBlockIds") or [])
        if not valid_block_ids or str(value) in valid_block_ids
    ][:20]
    source_block_ids = requested_blocks or list(inherited_blocks or [])[:20]
    children = []
    for child in (node.get("children") or [])[:20]:
        clean = sanitize_outline_node(
            child,
            valid_ids,
            valid_block_ids,
            depth + 1,
            source_block_ids,
        )
        if clean: children.append(clean)
    content_form = str(node.get("contentForm") or "").strip()[:80]
    semantic_role = str(node.get("semanticRole") or "").strip()[:80]
    volume_type = str(node.get("volumeType") or "").strip()[:40]
    return {
        "title": title,
        "volumeType": volume_type,
        "description": description,
        "requirementIds": requirement_ids,
        "sourceBlockIds": source_block_ids,
        "contentForm": content_form,
        "semanticRole": semantic_role,
        "projectSpecific": bool(node.get("projectSpecific", bool(source_block_ids))),
        "children": children,
    }


def title_intent(title):
    text = str(title or "")
    groups = (
        ("difficulty", ("重点难点", "难点分析", "重难点")),
        ("supply", ("供货", "采购", "供应链", "备货")),
        ("transport", ("运输", "包装", "装卸")),
        ("installation", ("安装", "部署", "调试", "联调")),
        ("schedule", ("进度", "工期", "里程碑")),
        ("organization", ("组织", "岗位", "人员", "团队")),
        ("quality", ("质量", "检查", "检验")),
        ("acceptance", ("验收", "移交", "交付")),
        ("training", ("培训", "知识转移")),
        ("service", ("售后", "运维", "质保", "响应")),
        ("risk", ("风险", "应急", "安全", "保密")),
    )
    for name, keywords in groups:
        if any(keyword in text for keyword in keywords):
            return name
    return ""


def preserve_generated_content(new_nodes, old_nodes):
    old_items = []
    used = set()

    def collect(nodes):
        for item in nodes or []:
            if not isinstance(item, dict):
                continue
            if item.get("content"):
                old_items.append(item)
            collect(item.get("children") or [])

    def restore(nodes):
        for item in nodes:
            previous = None
            best_score = 0
            new_title = norm(item.get("title"))
            new_requirements = set(item.get("requirementIds") or [])
            for candidate in old_items:
                candidate_key = id(candidate)
                if candidate_key in used:
                    continue
                old_title = norm(candidate.get("title"))
                title_score = SequenceMatcher(
                    None, new_title, old_title, autojunk=False
                ).ratio()
                old_requirements = set(
                    candidate.get("requirementIds") or []
                )
                union = new_requirements | old_requirements
                requirement_score = (
                    len(new_requirements & old_requirements) / len(union)
                    if union else 0
                )
                score = max(title_score, title_score * 0.7 + requirement_score * 0.3)
                if (
                    title_intent(new_title)
                    and title_intent(new_title) == title_intent(old_title)
                ):
                    score = max(score, 0.68)
                if score > best_score:
                    previous, best_score = candidate, score
            if previous and best_score >= 0.58 and not item.get("children"):
                used.add(id(previous))
                item["content"] = previous.get("content")
                item["contentStatus"] = previous.get("contentStatus") or "ready"
                item["contentMode"] = previous.get("contentMode")
                item["contentUpdatedAt"] = previous.get("contentUpdatedAt")
            restore(item.get("children") or [])

    collect(old_nodes)
    restore(new_nodes)


def build_safe_outline(requirements, project_name=""):
    def node(title, description, types=()):
        ids = [item["id"] for item in requirements if item["type"] in set(types)]
        return {"title": title, "description": description, "requirementIds": ids[:30], "children": []}

    service_word = "服务" if any(word in project_name for word in ("服务", "管理", "培训", "运营")) else "项目"
    return [
        {"title": "项目理解与需求分析", "description": "从采购目标、范围和交付成果出发形成针对性理解", "requirementIds": [], "children": [
            node("项目背景与建设目标", "结合招标文件说明项目背景、现状、目标和预期成果", ("technical", "deliverable")),
            node("服务范围与需求理解", f"梳理本{service_word}的工作范围、对象、边界和关键任务", ("technical", "commercial")),
            node("重点难点分析及应对思路", "识别实施、协同、进度、质量和合规方面的重点难点", ("scoring", "technical")),
        ]},
        {"title": "总体实施方案", "description": "说明总体思路、实施路径和工作机制", "requirementIds": [], "children": [
            node("总体思路与原则", "提出与采购目标匹配的实施理念、原则和总体策略", ("technical",)),
            node("实施路径与工作流程", "按启动、实施、检查、验收和移交阶段说明工作流程", ("technical", "deliverable")),
            node("工作方法与协同机制", "说明任务分解、沟通协调、问题闭环和过程留痕机制", ("scoring", "technical")),
        ]},
        {"title": "项目组织与资源配置", "description": "建立职责明确、响应高效的项目组织保障体系", "requirementIds": [], "children": [
            node("项目组织架构", "说明项目组织层级、管理关系和接口机制，具体人员信息待补充", ("scoring",)),
            node("岗位职责与人员安排", "说明岗位设置、职责分工和人员投入原则，具体人员信息待补充", ("qualification", "scoring")),
            node("资源配置与保障", "说明场地、设备、工具、资料和后勤等资源保障安排", ("technical", "commercial")),
        ]},
        {"title": "实施进度与交付计划", "description": "形成阶段清晰、节点可控的实施和交付安排", "requirementIds": [], "children": [
            node("项目阶段划分", "划分准备、实施、试运行或检查、验收和移交阶段", ("deadline", "deliverable")),
            node("进度计划与里程碑", "根据招标工期和交付要求设置里程碑，具体日期以合同为准", ("deadline", "commercial")),
            node("进度控制与纠偏措施", "说明计划跟踪、偏差分析、资源调整和赶工保障机制", ("scoring", "deliverable")),
        ]},
        {"title": "质量控制与验收保障", "description": "建立覆盖全过程的质量管理和成果验收机制", "requirementIds": [], "children": [
            node("质量目标与管理体系", "说明质量目标、责任体系和过程质量控制原则", ("technical", "scoring")),
            node("过程检查与成果审核", "说明自检、复核、审批、问题整改和版本管理方法", ("technical", "deliverable")),
            node("验收配合与资料移交", "响应验收要求并说明成果、记录和资料移交安排", ("deliverable", "commercial")),
        ]},
        {"title": "安全、保密与风险管理", "description": "控制实施过程中的安全、保密、合规和履约风险", "requirementIds": [], "children": [
            node("安全管理措施", "说明人员、场地、设备和作业过程的安全管理措施", ("technical", "commercial")),
            node("保密与数据管理", "说明资料访问、传输、存储、使用和销毁的控制措施", ("technical", "commercial")),
            node("风险识别与应急预案", "建立风险清单、预警、响应、升级和恢复机制", ("disqualification", "scoring")),
        ]},
        {"title": "服务保障与持续改进", "description": "说明培训、响应、服务保障和持续优化安排", "requirementIds": [], "children": [
            node("培训与知识转移", "说明培训对象、内容、方式、考核和资料交付安排", ("technical", "deliverable")),
            node("响应与售后服务", "说明服务渠道、问题分级、响应闭环和持续支持机制", ("commercial", "deliverable")),
            node("满意度与持续改进", "通过回访、评价、复盘和改进计划提升服务效果", ("scoring",)),
        ]},
        {"title": "商务及实质性条款响应", "description": "集中核对资格、商务、报价和否决性要求", "requirementIds": [], "children": [
            node("资格与证明材料响应", "逐项准备资格条件和证明材料，所有企业信息均须据实补充", ("qualification",)),
            node("商务条款响应", "响应服务期限、付款、履约、交付等商务条款", ("commercial", "deadline", "deposit")),
            node("实质性要求与偏离核对", "集中核对强制性、无效响应和偏离事项", ("disqualification", "other")),
        ]},
    ]


def _text_list(value, limit=12, item_limit=300):
    if not isinstance(value, list):
        return []
    return [str(item).strip()[:item_limit] for item in value if str(item).strip()][:limit]


ARCHETYPE_KEYWORDS = {
    "goods": ("货物", "产品", "采购清单", "供货", "交货", "制造商"),
    "equipment_integration": ("设备", "安装", "调试", "联调", "机电", "集成", "试运行", "弱电"),
    "software": ("软件", "软件平台", "信息化", "数字化", "开发", "接口", "数据库", "系统建设"),
    "professional_service": ("咨询", "审计", "检测", "调查", "研究", "规划", "培训服务", "设计服务"),
    "operation_service": ("运营", "运维", "购买服务", "物业", "保安", "保洁", "值守", "服务人员"),
    "construction": ("施工", "工程量清单", "竣工", "改造", "土建", "装修", "建筑工程", "场地建设", "外墙", "涂装"),
}


ARCHETYPE_PHASES = {
    "goods": ["需求与清单确认", "选型与采购", "生产备货", "出厂检查", "包装运输", "到货验收", "交付移交", "售后服务"],
    "equipment_integration": ["需求确认", "深化设计", "采购与备货", "出厂检查", "包装运输", "到货验收", "安装集成", "联调试运行", "培训验收", "售后维护"],
    "software": ["需求调研", "方案与架构设计", "开发或配置", "数据与接口准备", "部署集成", "测试整改", "试运行", "验收移交", "运行维护"],
    "professional_service": ["项目启动", "资料收集", "调研分析", "专业实施", "成果编制", "内部复核", "汇报评审", "修改定稿", "成果移交"],
    "operation_service": ["项目启动", "需求确认", "制度与流程建立", "资源进场", "日常服务实施", "质量巡检", "考核与整改", "持续改进", "成果移交"],
    "construction": ["施工准备", "深化设计与技术交底", "资源进场", "分部分项施工", "过程检查与隐蔽验收", "系统调试", "竣工验收", "资料移交", "缺陷责任服务"],
    "mixed": ["项目启动", "需求确认", "方案深化", "采购或资源准备", "分专业实施", "集成与协同", "检查或试运行", "验收移交", "持续服务"],
}


def classify_delivery_archetype(requirements, project_name=""):
    corpus = str(project_name or "") + "\n" + "\n".join(
        f"{item.get('title', '')} {item.get('requirement', '')}"
        for item in requirements
    )
    scores = {
        key: sum(corpus.count(word) for word in words)
        for key, words in ARCHETYPE_KEYWORDS.items()
    }
    ranked = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    if not ranked or ranked[0][1] == 0:
        return "mixed"
    first, second = ranked[0], ranked[1]
    if first[0] == "construction":
        return "construction"
    if first[0] in {"goods", "equipment_integration"} and (
        scores["equipment_integration"] >= 2
        and scores["goods"] >= 1
    ):
        return "equipment_integration"
    if first[1] >= 3 and second[1] >= 3 and second[1] >= first[1] * 0.75:
        return "mixed"
    return first[0]


def classify_archetype_components(
    requirements,
    project_name="",
    extra_text="",
    delivery_archetype=None,
):
    corpus = str(project_name or "") + "\n" + str(extra_text or "") + "\n" + "\n".join(
        f"{item.get('title', '')} {item.get('requirement', '')}"
        for item in requirements
    )
    scores = {
        key: sum(corpus.count(word) for word in words)
        for key, words in ARCHETYPE_KEYWORDS.items()
    }
    ranked = [
        key for key, score in sorted(
            scores.items(), key=lambda item: item[1], reverse=True
        )
        if score >= 1
    ]
    primary = (
        delivery_archetype
        if delivery_archetype in DELIVERY_ARCHETYPES
        else classify_delivery_archetype(requirements, project_name)
    )
    if primary != "mixed":
        return [primary]
    return ranked[:4] or ["mixed"]


def build_safe_blueprint(requirements, project_name=""):
    archetype = classify_delivery_archetype(requirements, project_name)
    project_type = DELIVERY_ARCHETYPES[archetype]
    phase_names = ARCHETYPE_PHASES[archetype]
    phases = [
        {
            "name": phase,
            "objective": f"完成{phase}阶段的规定工作并形成可核验成果",
            "tasks": ["明确输入条件", "执行阶段任务", "检查阶段成果", "确认后进入下一阶段"],
            "outputs": [f"{phase}工作记录", f"{phase}阶段成果"],
            "qualityGate": "阶段成果经检查确认，问题完成闭环处理",
            "risks": ["输入条件不完整", "任务衔接不及时", "成果记录不充分"],
        }
        for phase in phase_names
    ]
    return {
        "projectProfile": {
            "deliveryArchetype": archetype,
            "projectType": project_type,
            "objectives": [],
            "scope": [],
            "deliverables": [
                item["requirement"]
                for item in requirements
                if item.get("type") == "deliverable"
            ][:8],
            "constraints": [
                item["requirement"]
                for item in requirements
                if item.get("type") in {"deadline", "commercial", "disqualification"}
            ][:10],
            "stakeholders": ["采购人", "投标人项目团队", "相关配合单位", "验收或评审人员"],
        },
        "implementationBlueprint": {
            "templateName": project_type,
            "phases": phases,
        },
    }


def build_safe_project_analysis(requirements, project_name, blueprint):
    profile = blueprint["projectProfile"]
    archetype = profile.get("deliveryArchetype") or classify_delivery_archetype(
        requirements, project_name
    )
    technical = [
        item["requirement"]
        for item in requirements
        if item.get("type") in {"technical", "deliverable"}
    ]
    phases = blueprint["implementationBlueprint"]["phases"]
    return {
        "deliveryArchetype": archetype,
        "deliveryArchetypeLabel": DELIVERY_ARCHETYPES.get(archetype, "混合型项目"),
        "archetypeComponents": classify_archetype_components(
            requirements, project_name
        ),
        "procurementObjects": list(dict.fromkeys(technical))[:8],
        "workstreams": [
            {
                "name": phase["name"],
                "objective": phase.get("objective", ""),
                "outputs": phase.get("outputs", [])[:4],
            }
            for phase in phases[:12]
        ],
        "acceptanceObjects": profile.get("deliverables", [])[:10],
        "keyConstraints": profile.get("constraints", [])[:12],
        "domainSignals": [],
        "enterpriseInputsNeeded": [
            "项目负责人及团队人员资料",
            "与本项目相关的企业业绩和证明材料",
            "拟投入产品、设备、工具或服务资源资料",
            "企业售后、网点、库存等真实保障能力资料",
        ],
        "knowledgeGaps": [],
    }


def _suggested_artifacts(text):
    artifacts = []
    mappings = (
        (("进度", "工期", "节点"), "进度计划或里程碑表"),
        (("人员", "组织", "团队", "岗位"), "组织架构与岗位职责表"),
        (("质量", "检查", "审核"), "质量检查表"),
        (("风险", "应急", "安全"), "风险识别与应对表"),
        (("培训",), "培训计划表"),
        (("验收", "交付", "成果"), "验收或交付清单"),
        (("供货", "运输", "采购"), "供货与运输控制表"),
        (("服务", "售后", "运维"), "服务响应流程表"),
    )
    for keywords, artifact in mappings:
        if any(word in text for word in keywords):
            artifacts.append(artifact)
    return artifacts[:4] or ["方案要点与响应对照表"]


def route_scoring_requirement(item):
    title = str(item.get("title") or "")
    text = f"{title} {item.get('requirement', '')}"
    if any(word in text for word in (
        "价格扣除", "价格优惠", "报价得分", "价格分", "小微企业",
        "中小企业", "监狱企业", "残疾人福利性单位",
    )):
        return "pricing_policy", "form_or_declaration", False
    if any(word in text for word in (
        "\u8282\u80fd\u4ea7\u54c1", "\u73af\u5883\u6807\u5fd7\u4ea7\u54c1",
        "\u4f18\u5148\u91c7\u8d2d\u4ea7\u54c1", "\u8ba4\u8bc1\u8bc1\u4e66",
    )) and any(word in text for word in (
        "\u52a0\u5206", "\u5f97\u5206", "\u8bc4\u5ba1", "\u8bc4\u5206",
    )):
        return "evaluation_rule", "reminder_only", False
    if any(word in text for word in (
        "得分相同", "推荐顺序", "评审基准价", "计算公式",
        "评标办法", "评审方法", "排名规则",
    )):
        return "evaluation_rule", "reminder_only", False
    parameter_terms = (
            "参数", "指标", "精度", "尺寸", "承载", "续航", "功率",
            "分辨率", "频率", "容量", "带宽", "响应",
    )
    if (
        (
            any(word in title for word in parameter_terms[:-1])
            and not any(word in title for word in (
                "方案", "措施", "流程", "组织", "计划", "分析",
            ))
        )
        or (
            any(word in text for word in parameter_terms)
            and any(word in text for word in (
            "符合", "满足", "响应", "≥", "≤", "不低于", "不大于",
            ))
            and not any(word in title for word in (
                "方案", "措施", "流程", "组织", "计划",
            ))
        )
    ):
        return "technical_parameter", "technical_response_matrix", False
    if (
        any(word in text for word in (
            "证书", "资质", "业绩", "合同复印件", "证明材料",
            "社保", "学历", "职称", "获奖",
        ))
        and not any(word in text for word in (
            "组织方案", "团队与职责", "人员管理方案", "培训方案",
        ))
    ):
        return "qualification_evidence", "evidence_checklist", False
    if any(word in text for word in (
        "付款", "履约保证金", "报价", "合同条款", "商务条款",
        "服务期限承诺", "交货期承诺",
    )):
        return "commercial_response", "response_matrix", False
    if any(word in text for word in (
        "实质性", "无偏离", "偏离表", "强制性", "符合性审查",
        "承诺函",
    )):
        return "compliance_response", "compliance_matrix", False
    return "technical_solution", "section_narrative", True


def capability_ids_for_text(text, allowed_ids=None):
    allowed = set(allowed_ids or CAPABILITY_MODULES)
    ranked = []
    for module_id, module in CAPABILITY_MODULES.items():
        if module_id not in allowed:
            continue
        score = sum(str(text or "").count(word) for word in module["keywords"])
        if score:
            ranked.append((score, module_id))
    return [
        module_id for _, module_id in sorted(
            ranked, key=lambda item: (-item[0], item[1])
        )[:5]
    ]


def build_capability_plan(project_analysis, chapters):
    archetype = str(project_analysis.get("deliveryArchetype") or "mixed")
    components = set(
        project_analysis.get("archetypeComponents") or [archetype]
    )
    components.add(archetype)
    context = json.dumps({
        "workstreams": project_analysis.get("workstreams") or [],
        "objects": project_analysis.get("procurementObjects") or [],
        "acceptance": project_analysis.get("acceptanceObjects") or [],
        "chapters": chapters,
    }, ensure_ascii=False)
    core_ids = {
        "requirements_analysis", "overall_solution", "organization_resource",
        "schedule_control", "quality_control", "acceptance_handover",
        "risk_emergency", "document_delivery",
    }
    selected = []
    for module_id, module in CAPABILITY_MODULES.items():
        compatible = bool(components & set(module["archetypes"]))
        keyword_hit = any(word in context for word in module["keywords"])
        if module_id not in core_ids and not (compatible and keyword_hit):
            continue
        reasons = []
        if module_id in core_ids:
            reasons.append("全项目通用控制能力")
        if compatible:
            reasons.append("与项目交付形态匹配")
        if keyword_hit:
            reasons.append("项目工作流或章节明确涉及")
        selected.append({
            "id": module_id,
            "name": module["name"],
            "methodPattern": module["method"],
            "suggestedArtifacts": module["artifacts"],
            "reason": "；".join(reasons),
        })
    return selected


def build_safe_scoring_tasks(requirements):
    tasks = []
    for item in requirements:
        if item.get("type") != "scoring":
            continue
        text = f"{item.get('title', '')} {item.get('requirement', '')}"
        route_type, response_mode, generates_narrative = (
            route_scoring_requirement(item)
        )
        enterprise_inputs = []
        if any(word in text for word in ("人员", "证书", "资质", "社保")):
            enterprise_inputs.append("人员、证书或资质证明")
        if any(word in text for word in ("业绩", "案例", "合同")):
            enterprise_inputs.append("企业业绩、合同或验收证明")
        if any(word in text for word in ("产品", "设备", "参数", "检测报告")):
            enterprise_inputs.append("拟投产品参数和证明材料")
        tasks.append({
            "requirementId": item["id"],
            "title": item["title"],
            "responseObjective": f"完整、针对性响应“{item['title']}”的评价要求",
            "mustCover": [item["requirement"]],
            "suggestedArtifacts": _suggested_artifacts(text),
            "tenderFacts": [item["requirement"]],
            "enterpriseInputsNeeded": enterprise_inputs,
            "riskIfMissing": "可能造成该评分点响应不完整或缺少证明支撑",
            "targetSections": [],
            "routeType": route_type,
            "routeLabel": SCORING_ROUTE_LABELS[route_type],
            "responseMode": response_mode,
            "generatesNarrative": generates_narrative,
            "capabilityModuleIds": (
                capability_ids_for_text(text) or ["overall_solution"]
                if route_type == "technical_solution" else []
            ),
        })
    return tasks[:40]


def sanitize_project_analysis(value, requirements, project_name, blueprint):
    fallback = build_safe_project_analysis(
        requirements, project_name, blueprint
    )
    if not isinstance(value, dict):
        return fallback
    archetype = str(value.get("deliveryArchetype") or "").strip()
    if archetype not in DELIVERY_ARCHETYPES:
        archetype = fallback["deliveryArchetype"]
    workstreams = []
    for stream in (value.get("workstreams") or [])[:15]:
        if not isinstance(stream, dict):
            continue
        name = str(stream.get("name") or "").strip()[:100]
        if name:
            workstreams.append({
                "name": name,
                "objective": str(stream.get("objective") or "").strip()[:300],
                "outputs": _text_list(stream.get("outputs"), 6),
            })
    return {
        "deliveryArchetype": archetype,
        "deliveryArchetypeLabel": DELIVERY_ARCHETYPES[archetype],
        "archetypeComponents": [
            item for item in _text_list(
                value.get("archetypeComponents"), 7, 50
            )
            if item in DELIVERY_ARCHETYPES
        ] or classify_archetype_components(
            requirements,
            project_name,
            json.dumps(value, ensure_ascii=False),
            archetype,
        ),
        "procurementObjects": _text_list(
            value.get("procurementObjects"), 15
        ) or fallback["procurementObjects"],
        "workstreams": workstreams or fallback["workstreams"],
        "acceptanceObjects": _text_list(
            value.get("acceptanceObjects"), 12
        ) or fallback["acceptanceObjects"],
        "keyConstraints": _text_list(
            value.get("keyConstraints"), 15
        ) or fallback["keyConstraints"],
        "domainSignals": _text_list(value.get("domainSignals"), 15, 120),
        "enterpriseInputsNeeded": _text_list(
            value.get("enterpriseInputsNeeded"), 12
        ) or fallback["enterpriseInputsNeeded"],
        "knowledgeGaps": _text_list(value.get("knowledgeGaps"), 12),
    }


def sanitize_scoring_tasks(value, requirements):
    requirement_map = {
        item["id"]: item for item in requirements
        if item.get("type") == "scoring"
    }
    clean = []
    for task in value if isinstance(value, list) else []:
        if not isinstance(task, dict):
            continue
        requirement_id = str(task.get("requirementId") or "")
        item = requirement_map.get(requirement_id)
        if not item:
            continue
        route_type, response_mode, generates_narrative = (
            route_scoring_requirement(item)
        )
        capability_ids = (
            capability_ids_for_text(
                f"{item.get('title', '')} {item.get('requirement', '')}"
            ) or ["overall_solution"]
            if route_type == "technical_solution" else []
        )
        clean.append({
            "requirementId": requirement_id,
            "title": str(task.get("title") or item["title"]).strip()[:200],
            "responseObjective": str(
                task.get("responseObjective") or
                f"完整响应“{item['title']}”"
            ).strip()[:500],
            "mustCover": _text_list(task.get("mustCover"), 12),
            "suggestedArtifacts": _text_list(
                task.get("suggestedArtifacts"), 8, 120
            ),
            "tenderFacts": _text_list(task.get("tenderFacts"), 12),
            "enterpriseInputsNeeded": _text_list(
                task.get("enterpriseInputsNeeded"), 10
            ),
            "riskIfMissing": str(
                task.get("riskIfMissing") or "评分响应不完整"
            ).strip()[:300],
            "targetSections": _text_list(
                task.get("targetSections"), 8, 150
            ),
            "routeType": route_type,
            "routeLabel": SCORING_ROUTE_LABELS[route_type],
            "responseMode": response_mode,
            "generatesNarrative": generates_narrative,
            "capabilityModuleIds": capability_ids,
        })
    existing = {task["requirementId"] for task in clean}
    clean.extend(
        task for task in build_safe_scoring_tasks(requirements)
        if task["requirementId"] not in existing
    )
    return clean[:40]


def sanitize_blueprint(value, requirements, project_name):
    fallback = build_safe_blueprint(requirements, project_name)
    if not isinstance(value, dict):
        return fallback
    profile = value.get("projectProfile")
    implementation = value.get("implementationBlueprint")
    if not isinstance(profile, dict):
        profile = {}
    if not isinstance(implementation, dict):
        implementation = {}
    phases = []
    for phase in (implementation.get("phases") or [])[:12]:
        if not isinstance(phase, dict):
            continue
        name = str(phase.get("name") or "").strip()[:100]
        if not name:
            continue
        phases.append({
            "name": name,
            "objective": str(phase.get("objective") or "").strip()[:300],
            "tasks": _text_list(phase.get("tasks"), 10),
            "outputs": _text_list(phase.get("outputs"), 8),
            "qualityGate": str(phase.get("qualityGate") or "").strip()[:300],
            "risks": _text_list(phase.get("risks"), 8),
        })
    if len(phases) < 4:
        phases = fallback["implementationBlueprint"]["phases"]
    delivery_archetype = str(
        profile.get("deliveryArchetype")
        or fallback["projectProfile"]["deliveryArchetype"]
    ).strip()
    if delivery_archetype not in DELIVERY_ARCHETYPES:
        delivery_archetype = fallback["projectProfile"]["deliveryArchetype"]
    return {
        "projectProfile": {
            "deliveryArchetype": delivery_archetype,
            "projectType": str(profile.get("projectType") or fallback["projectProfile"]["projectType"]).strip()[:100],
            "objectives": _text_list(profile.get("objectives"), 10),
            "scope": _text_list(profile.get("scope"), 12),
            "deliverables": _text_list(profile.get("deliverables"), 12),
            "constraints": _text_list(profile.get("constraints"), 12),
            "stakeholders": _text_list(profile.get("stakeholders"), 10),
        },
        "implementationBlueprint": {
            "templateName": str(implementation.get("templateName") or fallback["implementationBlueprint"]["templateName"]).strip()[:100],
            "phases": phases,
        },
    }


SECTION_PROFILES = (
    {
        "keywords": ("资格", "证明材料", "资质材料"),
        "formFactor": "qualification_evidence",
        "writingPattern": "采用资格条件与证明材料清单逐项核对，不扩写为实施方案；企业名称、证书编号、业绩和人员资料缺失时保留待补充项。",
        "mustCover": ["资格条件与证明材料一一对应", "注明材料名称、有效性和装订位置", "不得虚构企业资质、业绩或人员信息"],
        "tables": ["资格条件与证明材料核对表"],
    },
    {
        "keywords": ("商务条款", "合同条款", "报价响应", "商务响应"),
        "formFactor": "commercial_response",
        "writingPattern": "按招标文件商务条款形成响应矩阵，明确条款原文、响应结论、偏离情况和证明位置；只有需要解释的履约安排才使用正文段落。",
        "mustCover": ["服务或交付期限", "付款及合同条件", "保证金与报价要求", "偏离情况和证明位置"],
        "tables": ["商务条款响应表"],
    },
    {
        "keywords": ("实质性", "偏离", "否决", "废标", "无效投标"),
        "formFactor": "compliance_matrix",
        "writingPattern": "以实质性要求和否决风险核对表为主体，逐项记录响应结论、所需证明和复核状态，不把否决条款扩写成技术方案。",
        "mustCover": ["强制性或否决性条件", "响应结论", "证明材料或正文位置", "人工复核状态"],
        "tables": ["实质性要求与偏离核对表"],
    },
    {
        "keywords": ("技术参数", "参数响应", "技术指标"),
        "formFactor": "technical_response_matrix",
        "writingPattern": "以技术参数响应矩阵为主体，逐项记录招标参数、拟响应值、偏离情况和证明材料；不得把参数列表改写成空泛叙述。",
        "mustCover": ["招标参数原文", "拟响应值", "正负偏离", "检测报告、彩页或其他证明位置"],
        "tables": ["技术参数响应表"],
    },
    {
        "keywords": ("重点", "难点"),
        "formFactor": "diagnostic_narrative",
        "writingPattern": "按项目专属业务环节识别4至7个具体难点；每个难点自然说明成因、影响和针对性措施，不重复固定栏目；最后可用一张简洁对照表收束。",
        "mustCover": ["难点必须出现具体设备、系统、作业场景或协同对象", "措施必须与难点一一对应，避免通用口号", "区分采购、运输、安装集成、调试、培训、售后及合规等不同环节"],
        "tables": ["重点难点与应对措施对照表"],
    },
    {
        "keywords": ("采购", "供货", "备货"),
        "formFactor": "operational_plan",
        "writingPattern": "先说明本项目货物构成和采购约束，再按选型确认、供应商协同、生产跟踪、出厂检查、备货和异常替代自然展开；步骤可编号，但不要重复目标、输入、责任等栏目。",
        "mustCover": ["具体货物或核心设备", "上游厂商与交付周期协调", "技术参数核验", "出厂检查和随货资料", "缺货、延期或型号不匹配时的处置"],
        "tables": ["采购控制要点表"],
    },
    {
        "keywords": ("运输", "包装", "装卸", "到货"),
        "formFactor": "operational_plan",
        "writingPattern": "围绕设备特性说明包装、防震、防潮、装卸、路线、在途跟踪、到货清点和异常处置；以连续作业流程为主，辅以控制点表。",
        "mustCover": ["设备分类包装", "精密或大型设备保护", "装卸与固定", "在途状态和交接", "到货开箱与损坏处理"],
        "tables": ["运输环节控制表"],
    },
    {
        "keywords": ("安装", "集成", "调试", "联调", "试运行"),
        "formFactor": "technical_process",
        "writingPattern": "按现场勘察、条件确认、安装部署、接口配置、单机测试、系统联调、试运行和问题整改组织；突出设备和系统之间的技术关系。",
        "mustCover": ["安装前置条件", "软硬件接口", "单机与系统测试", "问题分级整改", "试运行转验收条件"],
        "tables": ["安装调试检查表"],
    },
    {
        "keywords": ("进度", "工期", "里程碑"),
        "formFactor": "schedule",
        "writingPattern": "用阶段计划或里程碑表表达主线，正文只解释关键依赖、跟踪和纠偏机制，不重复逐阶段套话。",
        "mustCover": ["阶段划分", "关键依赖", "里程碑输出", "偏差监测与纠偏"],
        "tables": ["进度计划表", "里程碑检查表"],
    },
    {
        "keywords": ("组织", "岗位", "人员", "职责"),
        "formFactor": "organization",
        "writingPattern": "先说明组织关系和协同原则，再用岗位职责表呈现分工；企业真实姓名、证书和人数不足时保留待补充项。",
        "mustCover": ["组织层级", "岗位职责", "协同接口", "替补与升级机制"],
        "tables": ["岗位职责表", "沟通接口表"],
    },
    {
        "keywords": ("质量", "检查", "审核"),
        "formFactor": "quality_control",
        "writingPattern": "围绕事前预防、过程检查、成果复核和问题整改展开，使用检查表承载重复字段，正文解释关键控制逻辑。",
        "mustCover": ["质量责任", "检查对象与方法", "不符合项整改", "记录和版本控制"],
        "tables": ["质量检查表"],
    },
    {
        "keywords": ("验收", "移交", "交付"),
        "formFactor": "acceptance",
        "writingPattern": "依照招标文件明确的验收对象和依据，说明验收准备、实施、问题整改、复验和资料移交；不自行发明验收指标。",
        "mustCover": ["验收对象", "验收前准备", "验收步骤", "整改复验", "资料移交"],
        "tables": ["验收资料清单"],
    },
    {
        "keywords": ("培训", "知识转移"),
        "formFactor": "training_plan",
        "writingPattern": "按培训对象和使用场景组织课程、实操、考核与补训；以课程计划表为主，正文突出针对性。",
        "mustCover": ["培训对象", "理论与实操内容", "培训材料", "考核与补训", "培训记录"],
        "tables": ["培训计划表"],
    },
    {
        "keywords": ("售后", "响应", "运维", "质保"),
        "formFactor": "service_process",
        "writingPattern": "按受理、判断、派单、处理、验证、关闭和复盘说明服务闭环；仅引用招标文件明确的时限和承诺。",
        "mustCover": ["服务入口", "问题分级", "处理与升级", "验证关闭", "复盘改进"],
        "tables": ["服务响应流程表"],
    },
    {
        "keywords": ("风险", "应急", "安全", "保密", "数据"),
        "formFactor": "risk_control",
        "writingPattern": "先解释本项目特有风险，再用风险矩阵归纳触发条件、预防和处置；避免把每项风险重复写成八段。",
        "mustCover": ["具体风险场景", "触发条件", "预防和监测", "响应与恢复"],
        "tables": ["风险识别与应对表"],
    },
)


def section_profile(title):
    for profile in SECTION_PROFILES:
        if any(keyword in title for keyword in profile["keywords"]):
            return profile
    return {
        "formFactor": "professional_narrative",
        "writingPattern": "使用自然的专业段落和必要的小标题展开，依据内容选择步骤、项目符号或一张表格，不套用固定栏目。",
        "mustCover": ["紧扣本项目对象和招标要求", "说明可执行的方法和控制要点", "避免空泛承诺和重复总结"],
        "tables": [],
    }


def requirement_link_allowed(node, item):
    text = f"{node.get('title', '')}{node.get('description', '')}"
    requirement_text = f"{item.get('title', '')}{item.get('requirement', '')}"
    item_type = item.get("type")
    if any(word in text for word in ("资格", "证书", "业绩", "人员资质", "证明材料")):
        return item_type in {"qualification", "scoring"}
    if any(word in text for word in ("商务", "报价", "付款", "保证金", "合同", "价格")):
        return item_type in {"commercial", "deadline", "deposit", "scoring", "other"}
    if any(word in text for word in ("废标", "无效", "实质性", "偏离", "否决")):
        return item_type in {"disqualification", "other", "qualification"}
    blocked = ("价格扣除", "价格分", "中小企业", "业绩", "资质", "证书", "社保", "保证金", "报价修正")
    if item_type == "scoring" and any(word in requirement_text for word in blocked):
        return False
    if item_type in {"qualification", "deposit"}:
        return False
    return True


def refine_requirement_links(nodes, requirement_map):
    for node in nodes:
        node["requirementIds"] = [
            value
            for value in (node.get("requirementIds") or [])
            if value in requirement_map
            and requirement_link_allowed(node, requirement_map[value])
        ]
        refine_requirement_links(node.get("children") or [], requirement_map)


ROUTING_GENERIC_TERMS = {
    "项目", "方案", "要求", "响应", "服务", "工作", "实施", "管理",
    "采购", "相关", "进行", "提供", "说明", "内容", "投标", "招标",
}


def outline_leaf_entries(nodes, prefix=()):
    leaves = []
    for index, node in enumerate(nodes or []):
        path = prefix + (index,)
        children = node.get("children") or []
        if children:
            leaves.extend(outline_leaf_entries(children, path))
        else:
            leaves.append({"path": path, "node": node})
    return leaves


def requirement_route_lane(item):
    item_type = str(item.get("type") or "other")
    text = f"{item.get('title', '')} {item.get('requirement', '')}"
    if item_type == "scoring":
        route_type, _, _ = route_scoring_requirement(item)
        return route_type
    if item_type == "qualification":
        return "qualification_evidence"
    if item_type == "disqualification":
        return "compliance_response"
    if item_type in {"commercial", "deadline", "deposit"}:
        return "commercial_response"
    if any(word in text for word in ("资格条件", "资格证明", "资质证明")):
        return "qualification_evidence"
    if any(word in text for word in ("废标", "无效投标", "否决", "实质性要求")):
        return "compliance_response"
    if any(word in text for word in ("付款", "保证金", "报价", "合同条款", "商务条款")):
        return "commercial_response"
    if (
        any(word in text for word in ("参数", "指标", "规格", "型号"))
        and any(word in text for word in ("满足", "响应", "不低于", "不大于", "≥", "≤"))
    ):
        return "technical_parameter"
    return "technical_solution"


def section_route_lane(node):
    form_factor = section_profile(str(node.get("title") or ""))["formFactor"]
    if form_factor == "qualification_evidence":
        return "qualification_evidence"
    if form_factor == "commercial_response":
        return "commercial_response"
    if form_factor == "compliance_matrix":
        return "compliance_response"
    if form_factor == "technical_response_matrix":
        return "technical_parameter"
    return "technical_solution"


def route_match_score(item, entry, scoring_task=None):
    node = entry["node"]
    requirement_text = f"{item.get('title', '')} {item.get('requirement', '')}"
    section_text = f"{node.get('title', '')} {node.get('description', '')}"
    requirement_terms = {
        term for term in re.findall(r"[\u4e00-\u9fff]{2,6}", requirement_text)
        if term not in ROUTING_GENERIC_TERMS
    }
    score = sum(
        1 + min(3, len(term) - 1)
        for term in requirement_terms
        if term in section_text
    )
    requirement_capabilities = set(capability_ids_for_text(requirement_text))
    section_capabilities = set(capability_ids_for_text(section_text))
    score += 7 * len(requirement_capabilities & section_capabilities)
    if scoring_task:
        for target in scoring_task.get("targetSections") or []:
            if target and (
                target in str(node.get("title") or "")
                or str(node.get("title") or "") in target
            ):
                score += 15
    intent = title_intent(requirement_text)
    if intent and intent == title_intent(section_text):
        score += 12
    return score


def apply_requirement_routing(chapters, requirements, scoring_tasks):
    leaves = outline_leaf_entries(chapters)
    if not leaves:
        return {
            "version": "strict-primary-v1",
            "totalRequirements": len(requirements),
            "primaryAssignments": 0,
            "secondaryAssignments": 0,
            "totalPlacements": 0,
            "averagePlacements": 0,
            "laneCounts": {},
            "assignments": [],
        }
    for entry in leaves:
        entry["node"]["requirementIds"] = []
    scoring_map = {
        str(task.get("requirementId") or ""): task
        for task in scoring_tasks or []
    }
    lane_counts = {}
    assignments = []
    leaf_loads = {entry["path"]: 0 for entry in leaves}
    for item in requirements:
        requirement_id = str(item.get("id") or "")
        lane = requirement_route_lane(item)
        lane_counts[lane] = lane_counts.get(lane, 0) + 1
        virtual_targets = {
            "technical_parameter": "产品技术参数响应矩阵",
            "pricing_policy": "价格政策与声明",
            "evaluation_rule": "评审规则提醒",
        }
        if lane in virtual_targets:
            assignments.append({
                "requirementId": requirement_id,
                "requirementType": item.get("type"),
                "lane": lane,
                "primaryKind": "virtual_register",
                "primaryPath": None,
                "primaryTitle": virtual_targets[lane],
                "primaryScore": None,
                "secondaryPath": None,
                "secondaryTitle": None,
                "reason": "该要求由独立响应矩阵或提醒清单承载，不扩写进技术方案正文",
            })
            continue
        candidates = [
            entry for entry in leaves
            if section_route_lane(entry["node"]) == lane
        ]
        if not candidates and lane == "technical_parameter":
            candidates = [
                entry for entry in leaves
                if section_route_lane(entry["node"]) == "technical_solution"
            ]
        if not candidates:
            candidates = leaves
        ranked = sorted(
            [
                (
                route_match_score(
                    item, entry, scoring_map.get(requirement_id)
                ),
                -leaf_loads[entry["path"]],
                -index,
                entry,
                )
                for index, entry in enumerate(candidates)
            ],
            key=lambda value: (value[0], value[1], value[2]),
            reverse=True,
        )
        if ranked[0][0] <= 0 and lane == "technical_solution":
            assignments.append({
                "requirementId": requirement_id,
                "requirementType": item.get("type"),
                "lane": lane,
                "primaryKind": "unassigned_register",
                "primaryPath": None,
                "primaryTitle": None,
                "primaryScore": ranked[0][0],
                "secondaryPath": None,
                "secondaryTitle": None,
                "reason": "没有与该要求语义相关的技术正文，不强行绑定到无关章节",
            })
            continue
        primary = ranked[0][3]
        primary["node"].setdefault("requirementIds", []).append(requirement_id)
        leaf_loads[primary["path"]] += 1
        assignment = {
            "requirementId": requirement_id,
            "requirementType": item.get("type"),
            "lane": lane,
            "primaryKind": "outline_section",
            "primaryPath": list(primary["path"]),
            "primaryTitle": primary["node"].get("title"),
            "primaryScore": ranked[0][0],
            "secondaryPath": None,
            "secondaryTitle": None,
            "reason": "按要求类型硬分流后，依据章节语义、能力模块和评分任务目标选择唯一主承载章节",
        }
        assignments.append(assignment)
    for task in scoring_tasks or []:
        assignment = next(
            (
                value for value in assignments
                if value["requirementId"] == str(task.get("requirementId") or "")
            ),
            None,
        )
        if assignment and assignment.get("primaryTitle"):
            task["targetSections"] = [assignment["primaryTitle"]]
    narrative_placements = sum(
        len(entry["node"].get("requirementIds") or []) for entry in leaves
    )
    return {
        "version": "strict-primary-v1",
        "totalRequirements": len(requirements),
        "primaryAssignments": sum(
            1 for value in assignments
            if value.get("primaryKind") == "outline_section"
        ),
        "secondaryAssignments": 0,
        "virtualAssignments": sum(
            1 for value in assignments
            if value.get("primaryKind") == "virtual_register"
        ),
        "unassignedRequirements": sum(
            1 for value in assignments
            if value.get("primaryKind") == "unassigned_register"
        ),
        "narrativePlacements": narrative_placements,
        "totalPlacements": len(assignments),
        "averagePlacements": round(
            len(assignments) / max(1, len(requirements)), 2
        ),
        "laneCounts": lane_counts,
        "assignments": assignments,
    }


def build_page_budget(profile, requirements, scoring_obligations, suggested_tables):
    form_factor = profile["formFactor"]
    base_pages = {
        "qualification_evidence": 2.0,
        "commercial_response": 2.0,
        "compliance_matrix": 2.0,
        "technical_response_matrix": 3.0,
        "diagnostic_narrative": 4.0,
        "operational_plan": 4.0,
        "technical_process": 5.0,
        "schedule": 3.0,
        "organization": 3.0,
        "quality_control": 4.0,
        "acceptance": 3.0,
        "training_plan": 3.0,
        "service_process": 3.0,
        "risk_control": 4.0,
        "professional_narrative": 3.0,
    }.get(form_factor, 3.0)
    narrative_types = {"technical", "deliverable", "scoring"}
    narrative_count = sum(
        1 for item in requirements if item.get("type") in narrative_types
    )
    evidence_count = max(0, len(requirements) - narrative_count)
    mandatory_count = sum(
        1 for item in requirements if item.get("mandatory")
    )
    target_pages = (
        base_pages
        + min(3.0, narrative_count * 0.38)
        + min(1.5, evidence_count * 0.16)
        + min(1.0, mandatory_count * 0.12)
        + min(1.0, len(scoring_obligations) * 0.35)
        + min(0.8, len(suggested_tables) * 0.25)
    )
    target_pages = max(1.5, min(9.0, round(target_pages * 2) / 2))
    min_pages = max(1, round((target_pages * 0.72) * 2) / 2)
    max_pages = min(11, round((target_pages * 1.32) * 2) / 2)
    priority = (
        "high" if len(scoring_obligations) >= 2 or mandatory_count >= 4
        else "medium" if requirements or suggested_tables
        else "standard"
    )
    return {
        "minPages": min_pages,
        "targetPages": target_pages,
        "maxPages": max_pages,
        "targetCharacters": int(target_pages * 720),
        "priority": priority,
        "basis": {
            "baseFormFactor": form_factor,
            "requirementCount": len(requirements),
            "mandatoryCount": mandatory_count,
            "scoringObligationCount": len(scoring_obligations),
            "plannedTableCount": len(suggested_tables),
        },
    }


def build_section_brief(node, requirements, blueprint):
    title = str(node.get("title") or "")
    profile = section_profile(title)
    must_cover = [str(node.get("description") or "说明本章节的目标、范围和实施方法")]
    must_cover.extend(profile["mustCover"])
    project_analysis = blueprint.get("projectAnalysis") or {}
    if profile["formFactor"] == "diagnostic_narrative":
        must_cover.extend(
            f"{stream.get('name')}：{stream.get('objective')}"
            for stream in (project_analysis.get("workstreams") or [])[:4]
            if isinstance(stream, dict) and stream.get("name")
        )
    phases = blueprint.get("implementationBlueprint", {}).get("phases", [])
    matched = [
        phase for phase in phases
        if str(phase.get("name") or "") in title
        or any(term in title for term in re.findall(r"[\u4e00-\u9fff]{2,4}", str(phase.get("name") or "")))
    ]
    if not matched and any(word in title for word in ("总体", "实施", "流程", "阶段")):
        matched = phases
    requirement_ids = {
        str(item.get("id") or "") for item in requirements
    }
    scoring_obligations = [
        task for task in (blueprint.get("scoringTasks") or [])
        if str(task.get("requirementId") or "") in requirement_ids
        and task.get("generatesNarrative") is not False
    ]
    capability_plan = blueprint.get("capabilityPlan") or []
    planned_ids = {
        str(module.get("id") or "") for module in capability_plan
        if isinstance(module, dict)
    }
    matched_capability_ids = set(
        capability_ids_for_text(
            f"{title} {node.get('description', '')}", planned_ids
        )
    )
    for task in scoring_obligations:
        matched_capability_ids.update(
            task.get("capabilityModuleIds") or []
        )
    capability_modules = [
        module for module in capability_plan
        if module.get("id") in matched_capability_ids
    ][:6]
    page_budget = build_page_budget(
        profile, requirements, scoring_obligations, profile["tables"][:4]
    )
    if node.get("xiqueNode"):
        page_budget = {
            **page_budget,
            "minPages": 1.4,
            "targetPages": 2.1,
            "maxPages": 3.2,
            "targetCharacters": 1500,
            "priority": (
                "high" if scoring_obligations else page_budget.get("priority")
            ),
            "basis": {
                **(page_budget.get("basis") or {}),
                "xiqueCompiledLeaf": True,
            },
        }
    return {
        "purpose": str(node.get("description") or title)[:300],
        "formFactor": profile["formFactor"],
        "writingPattern": profile["writingPattern"],
        "pageBudget": page_budget,
        "mustCover": list(dict.fromkeys(must_cover))[:10],
        "relatedPhases": matched[:10],
        "requirementTitles": [item["title"] for item in requirements][:20],
        "requiredOutputs": list(dict.fromkeys(
            output for phase in matched for output in phase.get("outputs", [])
        ))[:12],
        "scoringObligations": scoring_obligations[:12],
        "capabilityModules": capability_modules,
        "suggestedTables": profile["tables"][:4],
        "avoidPatterns": ["不要逐项重复“目标、输入、实施步骤、责任角色、输出成果、质量关卡、风险处置、过程记录”", "不要用“围绕项目蓝图”“建立闭环机制”等套话作为每节开头"],
        "forbiddenClaims": [
            "招标文件未给出的期限、金额、数量、品牌、参数和量化承诺",
            "投标人未提供的资质、人员姓名、证书、业绩、设备数量和服务网点",
            "未经企业资料证明的成熟供应链、长期合作厂家、现有设备、自有车辆、固定库存、认证人员和既有客服系统",
        ],
    }


def attach_section_briefs(nodes, requirement_map, blueprint):
    for node in nodes:
        children = node.get("children") or []
        if children:
            attach_section_briefs(children, requirement_map, blueprint)
            continue
        requirements = [
            requirement_map[value]
            for value in collect_node_requirement_ids(node)
            if value in requirement_map
        ]
        node["brief"] = build_section_brief(node, requirements, blueprint)


def summarize_document_budget(chapters):
    budgets = [
        entry["node"].get("brief", {}).get("pageBudget", {})
        for entry in outline_leaf_entries(chapters)
    ]
    budgets = [value for value in budgets if value]
    return {
        "version": "semantic-page-budget-v1",
        "sectionCount": len(budgets),
        "minPages": round(sum(float(value.get("minPages") or 0) for value in budgets), 1),
        "targetPages": round(sum(float(value.get("targetPages") or 0) for value in budgets), 1),
        "maxPages": round(sum(float(value.get("maxPages") or 0) for value in budgets), 1),
        "targetCharacters": sum(
            int(value.get("targetCharacters") or 0) for value in budgets
        ),
        "highPrioritySections": sum(
            1 for value in budgets if value.get("priority") == "high"
        ),
    }


def call_planner_json(
    messages,
    max_tokens,
    project_id,
    job_id,
    chunk_no,
    run_type,
    target_model,
    timeout=300,
):
    if target_model.startswith("gpt-"):
        request_url = IMAGE_BASE_URL.rstrip("/") + "/v1/chat/completions"
        request_key = IMAGE_KEY
        if not request_key:
            raise ValueError("sub2api GPT 密钥未配置")
    else:
        request_url = AI_URL
        request_key = AI_KEY
    body = make_chat_body(
        messages, max_tokens, 0.1, model_override=target_model
    )
    last_error = None
    retry_delays = (0, 5, 15, 30)
    for attempt, retry_delay in enumerate(retry_delays):
        if retry_delay:
            time.sleep(retry_delay)
        started = time.monotonic()
        run_id = str(uuid.uuid4())
        input_tokens = output_tokens = 0
        actual_cost = None
        try:
            request = urllib.request.Request(
                request_url,
                data=body,
                headers={
                    "Authorization": f"Bearer {request_key}",
                    "Content-Type": "application/json",
                },
                method="POST",
            )
            with urllib.request.urlopen(request, timeout=timeout) as response:
                client_request_id = response.headers.get(
                    "X-Client-Request-Id"
                )
                payload = json.loads(response.read().decode("utf-8"))
            usage = payload.get("usage") or {}
            input_tokens = int(usage.get("prompt_tokens") or 0)
            output_tokens = int(usage.get("completion_tokens") or 0)
            billing = fetch_billing(client_request_id)
            if billing:
                input_tokens, output_tokens, actual_cost = billing
            parsed = parse_json(
                payload["choices"][0]["message"]["content"]
            )
            save_run(
                run_id,
                project_id,
                job_id,
                chunk_no,
                "succeeded",
                input_tokens,
                output_tokens,
                int((time.monotonic() - started) * 1000),
                attempt,
                actual_cost=actual_cost,
                run_type=run_type,
                model_override=target_model,
            )
            return parsed
        except Exception as exc:
            last_error = exc
            save_run(
                run_id,
                project_id,
                job_id,
                chunk_no,
                "failed",
                input_tokens,
                output_tokens,
                int((time.monotonic() - started) * 1000),
                attempt,
                exc,
                actual_cost,
                run_type,
                target_model,
            )
    raise RuntimeError(f"{run_type}调用失败: {last_error}")


def call_outline_model(
    requirements,
    project_name,
    source_context,
    project_id,
    job_id,
    target_model=None,
    dynamic=False,
):
    target_model = target_model or MODEL
    prompt_requirements = requirements
    if dynamic:
        prompt_requirements = [
            {
                **item,
                "requirement": str(item.get("requirement") or "")[:450],
            }
            for item in requirements
        ]
    system = """你是政府采购投标文件总策划师。先识别项目类型并规划可执行的项目生命周期，再据此生成可扩写为完整投标文件的专业大纲。
项目画像中的事实只能来自招标文件。实施阶段、工作方法、质量关卡和风险处置可以使用通用专业方法，但不得伪装成采购人明确要求。
大纲规则：
1. 生成8至12个一级章节，每章3至7个二级章节，不生成三级章节。
2. 技术方案必须按本项目真实业务生命周期组织，而不是按“资格、商务、评分项、技术参数、废标项”分类堆砌。
3. 先从goods、equipment_integration、software、professional_service、operation_service、construction、mixed中选择交付形态，再根据招标文件规划真实工作流。不得因为示例或既有项目而默认采用某个行业流程。
4. 评分表中的技术方案评分点应落入最适合得分的方案章节；价格扣除、企业业绩、人员证书等不得错误关联到实施难点章节。
5. 技术参数作为设备方案或技术响应的依据，不要把每个参数单独生成一个正文小节。
6. 资格、报价、商务和实质性响应可以集中在末章，但不能挤占技术方案主体。
7. 为每条scoring要求生成独立评分任务，明确必须回答的问题、建议成果形式、招标事实、所需企业资料、缺失风险和目标章节。
   同时判断该任务应进入技术方案正文、商务响应、资格证明、价格政策、实质性响应还是仅作为评审规则提醒；评审规则和价格计算不得扩写成技术方案。
8. 项目作战图必须区分采购或服务对象、工作流、验收对象、关键约束、行业信号、企业资料需求和知识缺口。
9. 不得杜撰企业资质、人员姓名、业绩、品牌、参数和承诺。输出严格JSON，不要Markdown。"""
    if dynamic:
        system += """
10. 本次只规划项目专属骨架：先生成且只生成第14条规定的四个一级分卷，每卷3至8个二级章节。标题必须使用本项目原文中的场所、对象、系统、设备、服务事项、实施条件或验收动作，禁止复制通用章节套装。
11. 不得把所有行业统一套入“任务分解、关键环节、执行标准、持续改进”等固定标题。相邻章节的业务对象、解决问题和成果形式必须不同。
12. 每个二级章节必须填写sourceBlockIds，引用输入tenderExcerpts中真实存在的块ID；同时填写contentForm和semanticRole。正文方法可以通用，但目录的业务名词与边界必须来自本项目。
13. 发现本项目自己的实体类型、实体关系、业务工作流和评价维度，写入projectSemantics。不要套用预设行业字段。
14. 动态模式采用四个功能分卷，但卷内目录必须完全由本项目诱导：
    - overall_solution：服务整体方案，承载管理重点难点、管理服务模式、内部管理架构、整体设想策划，并选取本项目最复杂、最能体现水平的场所或对象形成专项方案；
    - specific_service：服务具体方案，承载本项目的日常行为制度、安全制度、岗位职责、工作标准、奖惩办法和突发事件预案，并按真实对象细化；
    - resource_configuration：人员与设备配备，承载项目负责人、人员投入、专业岗位、排班、资格要求、工具物资和设备投入；
    - additional_service_response：其他服务响应，承载评分表或采购需求中尚未被前三章覆盖的专项服务、临时任务、急修、巡检、活动保障和增值响应。
    必须生成且只生成这四个一级分卷，并为每个一级分卷填写volumeType。四个一级标题可以使用项目化名称，但不得改变分卷边界。
15. 四个技术方案分卷均禁止出现商务条款响应表、偏离情况、证明位置、报价、付款、合同签订、电子投标、联合体、企业资格证明和企业业绩证明。这些内容只写入scoringTasks或后台requirementRouting，不生成技术方案目录。人员岗位资格和设备投入只进入resource_configuration。
16. 目录可以规划投标人自主设计的实施频次、会议机制、巡检节奏、内部质量指标、培训安排、记录表单和应急分级，但必须标记为我方方案或内部控制标准，不得写成招标文件原文事实；不得规划虚构的企业资质、人员姓名、证书编号、既有业绩、自有设备、服务网点或品牌型号。"""
        system += """
17. 联合体、电子投标、投标文件组成、签章、报价、付款、合同条款、企业资格和企业业绩均属于非技术响应材料，不得以“责任统筹”“项目约束”“其他响应”等改名后放入四个技术方案分卷。
18. 目录标题必须像正式投标文件目录，不得使用“确认采购需求、锁定对象范围、证明本章、响应评分点、对应评分要求、对应采购需求、评分要求到本章内容的响应结构、后续章节处理”等后台任务或写作指令式标题。标题直接写业务对象、现场问题、实施动作或管理机制。"""
    user = json.dumps({
        "projectName": project_name,
        "outputSchema": {
            "projectSemantics": {
                "entityTypes": [{
                    "name": "本项目发现的实体类型",
                    "examples": ["招标原文中的实体"],
                    "sourceBlockIds": ["原文块ID"],
                }],
                "relations": [{
                    "from": "实体或对象",
                    "relation": "本项目关系",
                    "to": "实体、动作、指标或成果",
                    "sourceBlockIds": ["原文块ID"],
                }],
                "workflows": [{
                    "name": "本项目特有工作流",
                    "steps": ["按业务先后排列的动作"],
                    "sourceBlockIds": ["原文块ID"],
                }],
                "evaluationDimensions": [{
                    "name": "评分、验收或考核维度",
                    "criteria": ["原文指标或判定方式"],
                    "sourceBlockIds": ["原文块ID"],
                }],
            },
            "projectAnalysis": {
                "deliveryArchetype": "goods|equipment_integration|software|professional_service|operation_service|construction|mixed",
                "archetypeComponents": ["混合型项目包含的交付形态代码；非混合项目只写自身"],
                "procurementObjects": ["采购、建设或服务对象"],
                "workstreams": [{
                    "name": "工作流名称",
                    "objective": "该工作流要解决的问题",
                    "outputs": ["可核验输出"],
                }],
                "acceptanceObjects": ["验收或评价对象"],
                "keyConstraints": ["招标文件明确约束"],
                "domainSignals": ["从招标文件识别的专业领域和术语"],
                "enterpriseInputsNeeded": ["需要投标人补充的真实资料"],
                "knowledgeGaps": ["招标文件未明确、正文不得擅自承诺的事项"],
            },
            "projectProfile": {
                "deliveryArchetype": "与projectAnalysis一致的交付形态代码",
                "projectType": "项目类型",
                "objectives": ["招标文件明确的目标"],
                "scope": ["招标文件明确的工作或供货范围"],
                "deliverables": ["明确交付成果"],
                "constraints": ["期限、参数、合规等明确约束"],
                "stakeholders": ["参与或配合角色"],
            },
            "implementationBlueprint": {
                "templateName": "适用的方法模板名称",
                "phases": [{
                    "name": "阶段名称",
                    "objective": "阶段目标",
                    "tasks": ["具体工作"],
                    "outputs": ["阶段输出物"],
                    "qualityGate": "进入下一阶段前的检查条件",
                    "risks": ["本阶段主要风险"],
                }],
            },
            "chapters": [{
                "title": "一级章节",
                "volumeType": "overall_solution|specific_service|resource_configuration|additional_service_response",
                "description": "本章写作目标",
                "requirementIds": ["输入ID"],
                "children": [{
                    "title": "二级章节",
                    "description": "写作要点",
                    "requirementIds": ["输入ID"],
                    "sourceBlockIds": ["招标原文块ID"],
                    "contentForm": "流程|作业方案|配置表|检查表|风险表|叙述等",
                    "semanticRole": "本节在本项目中的业务作用",
                    "projectSpecific": True,
                    "children": [],
                }],
            }],
            "scoringTasks": [{
                "requirementId": "scoring类型的输入ID",
                "title": "评分任务名称",
                "responseObjective": "本项要证明什么",
                "mustCover": ["必须回答的问题"],
                "suggestedArtifacts": ["建议表格、清单或图示"],
                "tenderFacts": ["可直接使用的招标事实"],
                "enterpriseInputsNeeded": ["需要企业提供的证明或事实"],
                "riskIfMissing": "遗漏后的失分风险",
                "targetSections": ["适合承载该任务的章节标题"],
                "routeType": "technical_solution|technical_parameter|commercial_response|qualification_evidence|pricing_policy|compliance_response|evaluation_rule",
            }],
        },
        "requirements": prompt_requirements,
        "tenderExcerpts": source_context,
    }, ensure_ascii=False)
    parsed = call_planner_json(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        3500 if dynamic else 8000,
        project_id,
        job_id,
        1,
        "outline_dynamic_plan" if dynamic else "outline",
        target_model,
        timeout=600 if dynamic else 360,
    )
    chapters = parsed.get("chapters") if isinstance(parsed, dict) else None
    if not isinstance(chapters, list):
        raise ValueError("模型未返回chapters数组")
    return parsed


def outline_node_values(node, key):
    values = []
    if isinstance(node, dict):
        values.extend(str(value) for value in (node.get(key) or []))
        for child in node.get("children") or []:
            values.extend(outline_node_values(child, key))
    return list(dict.fromkeys(value for value in values if value))


def dynamic_chapter_context(chapter, source_context, limit=16000):
    lines = [line for line in str(source_context or "").splitlines() if line]
    seed = " ".join(
        [str(chapter.get("title") or ""), str(chapter.get("description") or "")]
        + [
            str(child.get("title") or "")
            for child in (chapter.get("children") or [])
            if isinstance(child, dict)
        ]
    )
    generic = {
        "项目", "方案", "工作", "服务", "管理", "实施", "保障",
        "要求", "响应", "总体", "技术", "采购", "投标",
    }
    terms = {
        term
        for term in re.findall(r"[\u4e00-\u9fff]{2,10}", seed)
        if term not in generic
    }
    ranked = []
    for index, line in enumerate(lines):
        score = sum((2 + len(term)) * line.count(term) for term in terms)
        if score:
            ranked.append((score, index))
    selected_indexes = set()
    for _, index in sorted(ranked, reverse=True)[:80]:
        selected_indexes.update(
            range(max(0, index - 1), min(len(lines), index + 2))
        )
    for index, line in enumerate(lines):
        block_id = line[1:line.find("]")] if line.startswith("[") else ""
        if block_id in set(outline_node_values(chapter, "sourceBlockIds")):
            selected_indexes.update(
                range(max(0, index - 1), min(len(lines), index + 2))
            )
    if not selected_indexes:
        return "\n".join(lines)[:limit]
    return "\n".join(lines[index] for index in sorted(selected_indexes))[:limit]


def save_outline_planning_checkpoint(
    project_id, job_id, planned, expanded, progress
):
    checkpoint = {
        "version": "dynamic-outline-checkpoint-v1",
        "jobId": str(job_id),
        "planned": planned,
        "expanded": expanded,
        "completedVolumes": len(expanded),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    with psycopg.connect(DB) as conn:
        conn.execute(
            """
            UPDATE outlines
            SET content=jsonb_set(
                    coalesce(content,'{}'::jsonb),
                    '{_outlinePlanningCheckpoint}',
                    %s::jsonb,
                    true
                ),
                updated_at=now()
            WHERE project_id=%s
            """,
            (json.dumps(checkpoint, ensure_ascii=False), project_id),
        )
        conn.execute(
            "UPDATE projects SET progress=%s,updated_at=now() WHERE id=%s",
            (progress, project_id),
        )
        conn.commit()


def call_dynamic_outline_model(
    requirements, project_name, source_context, project_id, job_id,
    checkpoint=None,
):
    checkpoint = checkpoint if isinstance(checkpoint, dict) else {}
    if (
        checkpoint.get("version") == "dynamic-outline-checkpoint-v1"
        and str(checkpoint.get("jobId") or "") == str(job_id)
        and isinstance(checkpoint.get("planned"), dict)
    ):
        planned = checkpoint["planned"]
        expanded = [
            value for value in (checkpoint.get("expanded") or [])
            if isinstance(value, dict)
        ]
        print(json.dumps({
            "event": "outline_dynamic_checkpoint_resumed",
            "projectId": project_id,
            "jobId": job_id,
            "completedVolumes": len(expanded),
        }, ensure_ascii=False), flush=True)
    else:
        planned = call_outline_model(
            requirements,
            project_name,
            source_context,
            project_id,
            job_id,
            target_model=GPT_TEXT_MODEL,
            dynamic=True,
        )
        expanded = []
        save_outline_planning_checkpoint(
            project_id, job_id, planned, expanded, 20
        )
    base_chapters = [
        value for value in (planned.get("chapters") or [])
        if isinstance(value, dict)
    ][:12]
    if not base_chapters:
        raise ValueError("GPT没有返回项目专属一级目录")
    requirement_map = {
        str(item.get("id") or ""): item for item in requirements
    }
    for index, chapter in enumerate(base_chapters):
        if index < len(expanded):
            continue
        linked_ids = outline_node_values(chapter, "requirementIds")
        linked_requirements = [
            requirement_map[value]
            for value in linked_ids
            if value in requirement_map
        ]
        volume_type = str(chapter.get("volumeType") or "")
        allowed_types = {
            "overall_solution": {"scoring", "technical", "deliverable"},
            "specific_service": {"scoring", "technical", "deliverable"},
            "resource_configuration": {
                "scoring", "technical", "deliverable", "qualification"
            },
            "additional_service_response": {
                "scoring", "technical", "deliverable"
            },
        }.get(volume_type, {"scoring", "technical", "deliverable"})
        linked_requirements = [
            item for item in linked_requirements
            if item.get("type") in allowed_types
        ]
        if not linked_requirements:
            linked_requirements = [
                item for item in requirements
                if item.get("type") in allowed_types
            ][:45]
        chapter_context = dynamic_chapter_context(
            chapter, source_context
        )
        system = """你是投标技术方案目录深化专家。只深化输入中的一个一级章节，输出严格JSON对象。
规则：
1. 保留一级章节主题，根据本项目真实场所、对象、系统、设备、服务事项、评分动作、实施条件和验收要求，扩展为4至5级可直接撰写的目录；复杂对象必须到第五级，简单对象可停在第四级。
2. 目录必须体现本项目独有的业务分解。例如应写“具体场所/具体对象+具体动作或标准”，不得写成任何行业都能使用的抽象标题。
3. 禁止批量复用“任务分解、作业准备与连续实施”“关键环节、协作接口与过程记录”“执行标准、检查要点与合格判定”“异常处置、成果复核与持续改进”等固定标题或同义改写。
4. 同级标题之间必须解决不同对象、不同场景或不同问题；禁止只替换序号。叶子标题应当让评审人员仅看目录就能知道要响应哪项具体工作。
5. 每个叶子节点必须填写至少一个真实sourceBlockIds；只能使用tenderExcerpts中方括号里的块ID。尽量填写对应requirementIds。
6. 为每个叶子指定适合的contentForm，如专业论述、逐步流程、配置表、频次表、检查表、验收矩阵、风险清单、应急流程或图文说明，避免全文只有长段落。
7. 通用质量、安全、进度和应急方法必须结合本章具体对象命名，不能反过来主导目录。
8. 严格服从chapterSkeleton.volumeType的分卷边界。四个分卷均不得生成商务条款、偏离情况、证明位置、报价、付款、合同签订、电子投标、联合体、企业资格或企业业绩；resource_configuration只写人员岗位与资源配置；additional_service_response只写前三章未覆盖的项目专项服务。
9. 可以设计投标人自主实施的作业频次、响应分级、内部检查、培训安排、会议机制、表单记录和可执行质量目标，并在description或semanticRole中说明其属于“我方方案/内部控制标准”；不得把这些设计写成采购人已规定的事实。
10. 不得虚构企业资质、人员姓名、证书编号、既有业绩、客户名称、自有设备、库存、服务网点、品牌型号或招标文件未给出的强制参数。"""
        system += """
11. 联合体、电子投标、文件组成、签章、报价、付款、合同条款、企业资格和企业业绩不得在任何四个技术方案分卷中深化，只保留在后台路由。
12. 叶子标题必须是正式方案标题，禁止使用“确认、锁定、对应评分点、对应采购需求、回应评分要求、评分要求到本章内容的响应结构、后续章节、证明本章”等后台策划口吻。标题直接写业务对象、现场问题、实施动作或管理机制。"""
        system += """
13. 目录密度目标：
    - overall_solution生成5至7个二级章节、25至40个三级章节，并继续细化到70至120个四级或五级叶子；必须同时包含管理重难点、管理模式、内部架构、整体策划和本项目复杂场景专项方案。
    - specific_service生成5至7个二级章节、25至40个三级章节，并细化到70至120个叶子；重点为制度、职责、标准、奖惩和应急预案。
    - resource_configuration生成3至5个二级章节、12至25个三级章节，并细化到35至80个叶子。
    - additional_service_response按尚未覆盖的专项服务生成3至8个二级章节，并细化到30至80个叶子。
14. 不得为了达到数量批量复制同一套标题。每个叶子必须对应不同的场所、对象、动作、标准、异常或记录成果；若原文信息不足，应减少数量而不是用抽象模板凑数。"""
        user = json.dumps(
            {
                "projectName": project_name,
                "projectSemantics": planned.get("projectSemantics") or {},
                "projectAnalysis": planned.get("projectAnalysis") or {},
                "chapterSkeleton": chapter,
                "linkedRequirements": linked_requirements,
                "tenderExcerpts": chapter_context,
                "outputSchema": {
                    "chapter": {
                        "title": "一级章节标题",
                        "volumeType": "与chapterSkeleton相同的分卷类型",
                        "description": "本章响应目标",
                        "requirementIds": ["输入要求ID"],
                        "sourceBlockIds": ["原文块ID"],
                        "contentForm": "本层内容形态",
                        "semanticRole": "本层业务作用",
                        "projectSpecific": True,
                        "children": ["递归使用相同节点结构，最大5级"],
                    }
                },
            },
            ensure_ascii=False,
        )
        try:
            result = call_planner_json(
                [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                9000,
                project_id,
                job_id,
                index + 2,
                "outline_dynamic_expand",
                GPT_TEXT_MODEL,
                timeout=600,
            )
            expanded_chapter = (
                result.get("chapter") if isinstance(result, dict) else None
            )
            expanded.append(
                expanded_chapter
                if isinstance(expanded_chapter, dict)
                else chapter
            )
        except Exception as exc:
            expanded.append(chapter)
            print(
                json.dumps(
                    {
                        "event": "outline_dynamic_expand_fallback",
                        "projectId": project_id,
                        "chapter": index + 1,
                        "error": str(exc)[:200],
                    },
                    ensure_ascii=False,
                ),
                flush=True,
            )
        progress = 20 + int((index + 1) / len(base_chapters) * 70)
        save_outline_planning_checkpoint(
            project_id, job_id, planned, expanded, progress
        )
    planned["chapters"] = expanded
    return planned


def sanitize_project_semantics(value, valid_block_ids):
    value = value if isinstance(value, dict) else {}

    def clean_group(name, fields, limit):
        output = []
        for item in (value.get(name) or [])[:limit]:
            if not isinstance(item, dict):
                continue
            clean = {}
            for field in fields:
                raw = item.get(field)
                clean[field] = (
                    _text_list(raw, 12, 200)
                    if isinstance(raw, list)
                    else str(raw or "").strip()[:300]
                )
            clean["sourceBlockIds"] = [
                str(block_id)
                for block_id in (item.get("sourceBlockIds") or [])
                if str(block_id) in valid_block_ids
            ][:20]
            if any(clean.get(field) for field in fields):
                output.append(clean)
        return output

    return {
        "version": "project-induced-v1",
        "entityTypes": clean_group(
            "entityTypes", ("name", "examples"), 30
        ),
        "relations": clean_group(
            "relations", ("from", "relation", "to"), 40
        ),
        "workflows": clean_group(
            "workflows", ("name", "steps"), 25
        ),
        "evaluationDimensions": clean_group(
            "evaluationDimensions", ("name", "criteria"), 30
        ),
    }


def outline_specificity_audit(chapters):
    titles = []
    leaf_nodes = []
    depth_counts = {}
    generic_patterns = (
        "任务分解", "作业准备", "关键环节", "协作接口", "过程记录",
        "执行标准", "检查要点", "合格判定", "异常处置", "成果复核",
        "持续改进",
    )

    def visit(nodes, depth=1):
        for node in nodes or []:
            title = str(node.get("title") or "").strip()
            if title:
                titles.append(title)
                depth_counts[str(depth)] = depth_counts.get(str(depth), 0) + 1
            children = node.get("children") or []
            if children:
                visit(children, depth + 1)
            else:
                leaf_nodes.append(node)

    visit(chapters)
    unique_titles = len(set(titles))
    generic_titles = sum(
        1 for title in titles
        if any(pattern in title for pattern in generic_patterns)
    )
    bound_leaves = sum(
        1 for node in leaf_nodes if node.get("sourceBlockIds")
    )
    return {
        "version": "dynamic-outline-audit-v1",
        "titleCount": len(titles),
        "uniqueTitleCount": unique_titles,
        "titleUniquenessRate": round(
            unique_titles / max(len(titles), 1) * 100
        ),
        "leafCount": len(leaf_nodes),
        "sourceBoundLeaves": bound_leaves,
        "sourceBindingRate": round(
            bound_leaves / max(len(leaf_nodes), 1) * 100
        ),
        "genericTitleCount": generic_titles,
        "genericTitleRate": round(
            generic_titles / max(len(titles), 1) * 100
        ),
        "depthCounts": depth_counts,
    }


def outline_volume_boundary_audit(chapters):
    forbidden_by_volume = {
        "overall_solution": (
            "商务条款", "偏离情况", "证明位置", "投标报价", "分项报价",
            "合同签订", "资格证明", "业绩证明", "证书复印件",
            "设备投入承诺", "资格审查", "联合体", "电子投标",
            "签字盖章", "投标文件组成", "付款方式",
        ),
        "specific_service": (
            "商务条款", "偏离情况", "证明位置", "投标报价", "分项报价",
            "合同签订", "资格证明", "业绩证明", "证书复印件",
            "设备投入承诺", "资格审查", "联合体", "电子投标",
            "签字盖章", "投标文件组成", "付款方式",
        ),
        "resource_configuration": (
            "投标报价", "分项报价", "合同签订", "商务条款响应",
            "价格扣除", "联合体", "电子投标", "付款方式", "企业业绩",
        ),
        "additional_service_response": (
            "商务条款", "偏离情况", "证明位置", "投标报价", "分项报价",
            "合同签订", "资格审查", "联合体", "电子投标",
            "签字盖章", "投标文件组成", "付款方式", "企业业绩",
        ),
    }
    expected = (
        "overall_solution", "specific_service",
        "resource_configuration", "additional_service_response",
    )
    violations = []

    def visit(node, root_index, volume_type, path):
        title = str(node.get("title") or "").strip()
        description = str(node.get("description") or "").strip()
        semantic_role = str(node.get("semanticRole") or "").strip()
        content_form = str(node.get("contentForm") or "").strip()
        haystack = " ".join((title, description, semantic_role, content_form))
        for phrase in forbidden_by_volume.get(volume_type, ()):
            if phrase in haystack:
                violations.append({
                    "rootIndex": root_index,
                    "volumeType": volume_type,
                    "path": path + [title],
                    "phrase": phrase,
                })
                break
        for child in node.get("children") or []:
            if isinstance(child, dict):
                visit(child, root_index, volume_type, path + [title])

    actual = []
    for index, chapter in enumerate(chapters or []):
        volume_type = str(chapter.get("volumeType") or "").strip()
        actual.append(volume_type)
        visit(chapter, index, volume_type, [])
    missing = [value for value in expected if value not in actual]
    duplicates = [
        value for value in expected if actual.count(value) > 1
    ]
    return {
        "version": "volume-boundary-audit-v1",
        "expectedVolumeTypes": list(expected),
        "actualVolumeTypes": actual,
        "missingVolumeTypes": missing,
        "duplicateVolumeTypes": duplicates,
        "boundaryViolationCount": len(violations),
        "violations": violations[:30],
        "passed": (
            len(chapters or []) == 4
            and not missing
            and not duplicates
            and not violations
        ),
    }


def enforce_outline_volume_boundaries(chapters):
    roots = {
        str(node.get("volumeType") or ""): node
        for node in chapters or []
        if isinstance(node, dict)
    }
    resource_root = roots.get("resource_configuration")
    if not resource_root:
        return {"movedCount": 0, "moves": []}
    response_phrases = (
        "商务条款", "偏离情况", "证明位置", "投标报价", "分项报价",
        "合同签订", "资格证明", "业绩证明", "资格审查",
        "联合体", "电子投标", "签字盖章", "投标文件组成", "付款方式",
    )
    resource_phrases = (
        "设备投入承诺", "人员证书", "岗位证书", "人员配置",
        "设备配置", "工具配置",
    )
    moves = []

    def route_children(parent, volume_type, path):
        kept = []
        for child in list(parent.get("children") or []):
            if not isinstance(child, dict):
                continue
            text = " ".join((
                str(child.get("title") or ""),
                str(child.get("description") or ""),
                str(child.get("semanticRole") or ""),
                str(child.get("contentForm") or ""),
            ))
            target = None
            quarantine = False
            phrase = ""
            for candidate in response_phrases:
                if candidate in text:
                    quarantine, phrase = True, candidate
                    break
            if not quarantine and volume_type != "resource_configuration":
                for candidate in resource_phrases:
                    if candidate in text:
                        target, phrase = resource_root, candidate
                        break
            if quarantine:
                moves.append({
                    "fromVolume": volume_type,
                    "toVolume": "non_technical_requirement_routing",
                    "title": str(child.get("title") or ""),
                    "phrase": phrase,
                    "path": path + [str(child.get("title") or "")],
                })
                continue
            if target is not None:
                if target is parent or (
                    volume_type == "resource_configuration"
                    and target is resource_root
                ):
                    route_children(
                        child,
                        volume_type,
                        path + [str(child.get("title") or "")],
                    )
                    kept.append(child)
                    continue
                target.setdefault("children", []).append(child)
                moves.append({
                    "fromVolume": volume_type,
                    "toVolume": str(target.get("volumeType") or ""),
                    "title": str(child.get("title") or ""),
                    "phrase": phrase,
                    "path": path + [str(child.get("title") or "")],
                })
                continue
            route_children(
                child,
                volume_type,
                path + [str(child.get("title") or "")],
            )
            kept.append(child)
        parent["children"] = kept

    for volume_type in (
        "overall_solution", "specific_service",
        "resource_configuration", "additional_service_response",
    ):
        root = roots.get(volume_type)
        if root:
            route_children(root, volume_type, [str(root.get("title") or "")])
    return {
        "version": "volume-boundary-enforcement-v1",
        "movedCount": len(moves),
        "moves": moves[:50],
    }


XIQUE_SCENE_TEMPLATES = {
    "diagnostic_narrative": [
        "项目场景与服务对象分析",
        "重点区域及高频业务难点",
        "季节性、突发性与协同难点",
    ],
    "organization": [
        "管理架构与指挥关系",
        "岗位配置与职责边界",
        "协同调度与替补保障",
    ],
    "schedule": [
        "总体阶段与任务衔接",
        "关键节点与动态跟踪",
        "偏差纠正与进度保障",
    ],
    "quality_control": [
        "质量标准与责任体系",
        "日常检查与专项复核",
        "问题整改与持续改进",
    ],
    "acceptance": [
        "验收对象与前置准备",
        "过程验收与问题复验",
        "成果移交与资料归档",
    ],
    "training_plan": [
        "分层分类培训安排",
        "现场实操与岗位辅导",
        "考核补训与效果跟踪",
    ],
    "service_process": [
        "日常服务受理与组织",
        "重点场景响应与协同",
        "回访复盘与服务改进",
    ],
    "risk_control": [
        "主要风险场景识别",
        "预警处置与应急联动",
        "恢复验证与复盘改进",
    ],
    "operational_plan": [
        "服务范围与作业对象",
        "日常作业与专项保障",
        "检查考核与整改闭环",
    ],
    "technical_process": [
        "技术条件与实施准备",
        "安装配置与联调测试",
        "试运行、验收与技术移交",
    ],
    "professional_narrative": [
        "业务范围与工作对象",
        "实施方法与重点场景",
        "检查验证与改进机制",
    ],
}


XIQUE_MODULE_TEMPLATES = {
    "diagnostic_narrative": ("现状特征与问题成因", "针对性策略与实施安排"),
    "organization": ("组织设置与岗位安排", "协作机制与履职检查"),
    "schedule": ("计划分解与节点控制", "跟踪纠偏与资源保障"),
    "quality_control": ("标准落实与过程检查", "复核整改与记录归档"),
    "acceptance": ("验收准备与实施程序", "问题复验与成果移交"),
    "training_plan": ("培训内容与组织实施", "考核评价与持续辅导"),
    "service_process": ("服务流程与关键动作", "异常升级与闭环验证"),
    "risk_control": ("预防监测与现场控制", "应急处置与恢复复盘"),
    "operational_plan": ("作业组织与实施步骤", "质量标准与检查记录"),
    "technical_process": ("技术实施与过程控制", "测试验证与交付保障"),
    "professional_narrative": ("工作内容与实施步骤", "控制标准与成果验证"),
}


XIQUE_DETAIL_BY_MODULE = {
    "现状特征与问题成因": (
        ("现状调查、信息采集与特征归纳", "形成对象、区域、时段和现有条件的调查结果"),
        ("问题成因研判与需求结论", "分析问题表现、形成原因、影响范围和改善需求"),
    ),
    "针对性策略与实施安排": (
        ("分场景改进策略与工作措施", "提出与现场问题逐项对应的专业解决措施"),
        ("实施次序、协同界面与成果要求", "明确先后次序、协作关系和阶段成果"),
    ),
    "组织设置与岗位安排": (
        ("项目组织设置与岗位配置", "说明管理层级、专业岗位及其配置逻辑"),
        ("岗位职责、授权边界与工作接口", "明确岗位责任、授权事项和上下游接口"),
    ),
    "协作机制与履职检查": (
        ("横向协作、纵向汇报与替补机制", "说明跨岗位协同、信息汇报和缺岗替补安排"),
        ("履职检查、绩效评价与责任追踪", "明确履职证据、检查方式和责任追踪方法"),
    ),
    "计划分解与节点控制": (
        ("任务分解、时序安排与节点计划", "把总体任务分解至阶段、周期和关键节点"),
        ("节点确认、前置条件与成果交付", "明确节点启动条件、完成标准和交付成果"),
    ),
    "跟踪纠偏与资源保障": (
        ("计划跟踪、偏差预警与动态纠正", "说明进度采集、偏差判断和纠正方法"),
        ("人员物资调配与关键节点保障", "明确资源调度和重点时段保障措施"),
    ),
    "标准落实与过程检查": (
        ("服务标准分解与岗位执行要求", "把质量标准转化为具体岗位动作和作业要求"),
        ("日常巡查、专项检查与结果判定", "明确检查层级、检查内容和判定依据"),
    ),
    "复核整改与记录归档": (
        ("问题登记、整改复核与闭环销项", "说明问题从发现到验证销项的全过程"),
        ("质量记录、台账汇总与资料归档", "明确记录表单、汇总周期和归档要求"),
    ),
    "验收准备与实施程序": (
        ("验收条件确认与资料准备", "明确验收对象、前置条件和资料清单"),
        ("现场验收程序与结果确认", "说明验收步骤、参与角色和结果签认"),
    ),
    "问题复验与成果移交": (
        ("验收问题整改与复验确认", "明确问题整改、复验和销项要求"),
        ("成果移交、资料归档与后续衔接", "说明成果清点、移交手续和后续安排"),
    ),
    "培训内容与组织实施": (
        ("培训对象、课程内容与分层安排", "按岗位和能力层级设计培训内容"),
        ("课堂讲解、现场实操与组织实施", "说明培训方式、组织过程和现场辅导"),
    ),
    "考核评价与持续辅导": (
        ("培训考核、效果评价与结果应用", "明确考核方式、合格标准和结果运用"),
        ("不合格补训与在岗持续辅导", "说明补训、跟踪和能力巩固安排"),
    ),
    "服务流程与关键动作": (
        ("服务受理、任务派发与现场执行", "展开从需求受理到现场处理的连续动作"),
        ("过程反馈、结果确认与服务回访", "明确进度反馈、完成确认和回访方法"),
    ),
    "异常升级与闭环验证": (
        ("异常识别、分级上报与协同处置", "说明异常分级、升级路径和协同要求"),
        ("处置验证、用户反馈与闭环改进", "明确结果验证、意见收集和改进措施"),
    ),
    "预防监测与现场控制": (
        ("风险辨识、预防检查与预警触发", "明确风险来源、检查方法和触发条件"),
        ("现场隔离、过程控制与影响抑制", "说明现场控制动作和影响范围管理"),
    ),
    "应急处置与恢复复盘": (
        ("应急响应、资源调度与联动处置", "展开响应启动、人员调度和联动流程"),
        ("恢复验证、事件复盘与措施完善", "明确恢复确认、原因复盘和持续完善"),
    ),
    "作业组织与实施步骤": (
        ("作业准备、人员分工与现场实施", "明确班前准备、岗位分工和连续作业步骤"),
        ("交接确认、异常处理与作业记录", "说明作业交接、异常处置和记录要求"),
    ),
    "质量标准与检查记录": (
        ("作业标准、检查指标与合格判定", "列明过程标准、检查指标和合格条件"),
        ("巡检复查、问题整改与台账留存", "形成检查、整改、复核和留痕闭环"),
    ),
    "技术实施与过程控制": (
        ("技术准备、条件确认与规范实施", "说明实施前提、工序要求和技术动作"),
        ("关键工序、接口协调与过程记录", "明确关键工序、专业接口和记录要求"),
    ),
    "测试验证与交付保障": (
        ("单项测试、联动验证与问题整改", "说明测试项目、验证方法和问题处理"),
        ("试运行确认、成果交付与技术移交", "明确试运行、交付和技术资料移交"),
    ),
    "工作内容与实施步骤": (
        ("任务分解、作业准备与连续实施", "把本模块展开为可执行的连续动作和岗位衔接"),
        ("关键环节、协作接口与过程记录", "明确关键环节、协作关系和过程留痕"),
    ),
    "控制标准与成果验证": (
        ("执行标准、检查要点与合格判定", "明确过程标准、检查重点和判定方法"),
        ("异常处置、成果复核与持续改进", "说明异常处理、成果复核和后续改进"),
    ),
}


def xique_detail_templates(module_title):
    return XIQUE_DETAIL_BY_MODULE.get(
        module_title,
        (
            ("具体工作内容与实施程序", "展开连续动作、岗位衔接和现场操作要求"),
            ("执行标准、检查方法与成果确认", "明确过程标准、检查重点和成果记录"),
        ),
    )


def xique_scene_titles(node, requirement_map):
    profile = section_profile(node.get("title"))
    form_factor = profile["formFactor"]
    titles = []
    for requirement_id in collect_node_requirement_ids(node):
        item = requirement_map.get(str(requirement_id))
        title = str((item or {}).get("title") or "").strip()
        title = re.sub(r"^(?:评分项|技术要求|服务要求|采购需求)[:：\s]*", "", title)
        if 4 <= len(title) <= 30 and title not in titles:
            titles.append(title)
        if len(titles) >= 2:
            break
    for title in XIQUE_SCENE_TEMPLATES.get(
        form_factor, XIQUE_SCENE_TEMPLATES["professional_narrative"]
    ):
        if norm(title) not in {norm(value) for value in titles}:
            titles.append(title)
        if len(titles) >= 3:
            break
    return titles[:3], form_factor


def expand_xique_outline(nodes, requirement_map, depth=1):
    expanded = []
    for node in nodes:
        current = json.loads(json.dumps(node, ensure_ascii=False))
        children = current.get("children") or []
        if children:
            current["children"] = expand_xique_outline(
                children, requirement_map, depth + 1
            )
            expanded.append(current)
            continue
        if depth >= 5:
            current["xiqueNode"] = True
            expanded.append(current)
            continue
        scenes, form_factor = xique_scene_titles(current, requirement_map)
        module_titles = XIQUE_MODULE_TEMPLATES.get(
            form_factor, XIQUE_MODULE_TEMPLATES["professional_narrative"]
        )
        requirement_ids = list(current.get("requirementIds") or [])
        scene_nodes = []
        for scene_index, scene_title in enumerate(scenes):
            module_nodes = []
            for module_index, module_title in enumerate(module_titles):
                detail_nodes = []
                for detail_index, (detail_title, detail_purpose) in enumerate(
                    xique_detail_templates(module_title)
                ):
                    detail_nodes.append({
                        "title": detail_title,
                        "description": (
                            f"围绕“{current.get('title')}—{scene_title}—"
                            f"{module_title}”{detail_purpose}"
                        ),
                        "requirementIds": (
                            requirement_ids
                            if (scene_index + module_index + detail_index) % 2 == 0
                            else []
                        ),
                        "children": [],
                        "xiqueNode": True,
                        "visualHint": (
                            "process_diagram" if detail_index == 0
                            else "control_table"
                        ),
                    })
                module_nodes.append({
                    "title": module_title,
                    "description": (
                        f"深化“{scene_title}”中的{module_title}，形成可装订的专业响应"
                    ),
                    "requirementIds": requirement_ids,
                    "children": detail_nodes,
                })
            scene_nodes.append({
                "title": scene_title,
                "description": (
                    f"从具体场景深化“{current.get('title')}”，避免仅作原则性描述"
                ),
                "requirementIds": requirement_ids,
                "children": module_nodes,
            })
        current["children"] = scene_nodes
        current["xiqueExpanded"] = True
        expanded.append(current)
    return expanded


def process_outline(job):
    job_id, project_id = job["jobId"], job["projectId"]
    requested_outline_mode = str(job.get("outlineMode") or "standard")
    outline_mode = (
        requested_outline_mode
        if requested_outline_mode in {"standard", "xique", "dynamic"}
        else "standard"
    )
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s", (job_id,))
        conn.execute("UPDATE projects SET status='outlining',progress=15,error_message=NULL,updated_at=now() WHERE id=%s", (project_id,))
        project = conn.execute("SELECT name FROM projects WHERE id=%s", (project_id,)).fetchone()
        doc_row = conn.execute("SELECT blocks FROM documents WHERE project_id=%s", (project_id,)).fetchone()
        old_outline_row = conn.execute(
            "SELECT content FROM outlines WHERE project_id=%s",
            (project_id,),
        ).fetchone()
        rows = conn.execute("""SELECT id,type,title,normalized_value,mandatory,ai_review_status FROM requirements WHERE project_id=%s AND review_status<>'rejected' ORDER BY mandatory DESC,created_at""", (project_id,)).fetchall()
        conn.commit()
    try:
        if not rows: raise ValueError("没有可用于生成大纲的招标要求")
        requirements = [{"id": str(row[0]), "type": row[1], "title": row[2], "requirement": row[3], "mandatory": bool(row[4]), "riskStatus": row[5]} for row in rows]
        blocks = doc_row[0] if doc_row and isinstance(doc_row[0], list) else (json.loads(doc_row[0]) if doc_row else [])
        selected = select_chunks(
            blocks, 18 if outline_mode == "dynamic" else 8
        )
        source_context = "\n".join(
            f"[{block.get('id')}] {block.get('text', '')}"
            for _, chunk, _ in selected for block in chunk
        )[:30000 if outline_mode == "dynamic" else 30000]
        planned = {}
        try:
            if outline_mode == "dynamic":
                planned = call_dynamic_outline_model(
                    requirements,
                    project[0] if project else "",
                    source_context,
                    project_id,
                    job_id,
                    (
                        (
                            old_outline_row[0].get(
                                "_outlinePlanningCheckpoint"
                            )
                            if isinstance(old_outline_row[0], dict)
                            else None
                        )
                        if old_outline_row else None
                    ),
                )
                generation_mode = "gpt_dynamic_semantic_v1"
                outline_model = GPT_TEXT_MODEL
            else:
                planned = call_outline_model(
                    requirements,
                    project[0] if project else "",
                    source_context,
                    project_id,
                    job_id,
                )
                generation_mode = "ai"
                outline_model = MODEL
            chapters = planned["chapters"]
            blueprint = sanitize_blueprint(planned, requirements, project[0] if project else "")
            project_analysis = sanitize_project_analysis(
                planned.get("projectAnalysis"),
                requirements,
                project[0] if project else "",
                blueprint,
            )
            scoring_tasks = sanitize_scoring_tasks(
                planned.get("scoringTasks"), requirements
            )
        except Exception as model_error:
            if outline_mode == "dynamic":
                raise RuntimeError(
                    f"GPT项目专属目录生成失败，未降级为模板目录: {model_error}"
                ) from model_error
            chapters = build_safe_outline(requirements, project[0] if project else "")
            blueprint = build_safe_blueprint(requirements, project[0] if project else "")
            project_analysis = build_safe_project_analysis(
                requirements, project[0] if project else "", blueprint
            )
            scoring_tasks = build_safe_scoring_tasks(requirements)
            generation_mode = "safe_fallback"
            outline_model = GPT_TEXT_MODEL if outline_mode == "dynamic" else MODEL
            print(json.dumps({"event": "outline_ai_fallback", "projectId": project_id, "error": str(model_error)[:200]}), flush=True)
        valid_ids = {item["id"] for item in requirements}
        valid_block_ids = {
            str(block.get("id") or "") for block in blocks
            if block.get("id")
        }
        clean = []
        for chapter in chapters[:20]:
            node = sanitize_outline_node(
                chapter, valid_ids, valid_block_ids
            )
            if node: clean.append(node)
        if not clean: raise ValueError("模型大纲没有有效章节")
        old_content = (
            old_outline_row[0]
            if old_outline_row and isinstance(old_outline_row[0], dict)
            else {}
        )
        if outline_mode != "dynamic":
            preserve_generated_content(
                clean, old_content.get("chapters") or []
            )
        volume_boundary_enforcement = (
            enforce_outline_volume_boundaries(clean)
            if outline_mode == "dynamic" else {}
        )
        requirement_map = {item["id"]: item for item in requirements}
        refine_requirement_links(clean, requirement_map)
        standard_chapters = None
        if outline_mode == "xique":
            if (
                old_content.get("generationSettings", {}).get("outlineMode")
                != "xique"
            ):
                standard_chapters = old_content.get("chapters") or []
            else:
                standard_chapters = old_content.get("standardChapters") or []
            clean = expand_xique_outline(clean, requirement_map)
            generation_mode = "xique_compiled"
        blueprint["projectProfile"]["deliveryArchetype"] = (
            project_analysis["deliveryArchetype"]
        )
        capability_plan = build_capability_plan(project_analysis, clean)
        planning_context = {
            **blueprint,
            "projectAnalysis": project_analysis,
            "scoringTasks": scoring_tasks,
            "capabilityPlan": capability_plan,
        }
        requirement_routing = apply_requirement_routing(
            clean, requirements, scoring_tasks
        )
        attach_section_briefs(clean, requirement_map, planning_context)
        document_budget = summarize_document_budget(clean)
        specificity_audit = outline_specificity_audit(clean)
        volume_boundary_audit = (
            outline_volume_boundary_audit(clean)
            if outline_mode == "dynamic" else {}
        )
        project_semantics = sanitize_project_semantics(
            planned.get("projectSemantics"), valid_block_ids
        )
        content = {
            "projectName": project[0] if project else "",
            "generationMode": generation_mode,
            "outlineModel": outline_model,
            "projectSemantics": project_semantics,
            "outlineSpecificityAudit": specificity_audit,
            "outlineVolumeBoundaryAudit": volume_boundary_audit,
            "outlineVolumeBoundaryEnforcement": volume_boundary_enforcement,
            "projectProfile": blueprint["projectProfile"],
            "implementationBlueprint": blueprint["implementationBlueprint"],
            "projectAnalysis": project_analysis,
            "scoringTasks": scoring_tasks,
            "capabilityPlan": capability_plan,
            "requirementRouting": requirement_routing,
            "documentBudget": document_budget,
            "generationSettings": {
                **(old_content.get("generationSettings") or {}),
                "outlineMode": outline_mode,
                **(
                    {"textModelMode": "gpt", "lengthMode": "xique"}
                    if outline_mode == "dynamic" else {}
                ),
                **({"lengthMode": "xique"} if outline_mode == "xique" else {}),
            },
            "standardChapters": standard_chapters or (
                old_content.get("standardChapters") or []
            ),
            "chapters": clean,
        }
        with psycopg.connect(DB) as conn:
            conn.execute("""INSERT INTO outlines(project_id,content,status,version,model,generated_at,updated_at) VALUES(%s,%s::jsonb,'ready',1,%s,now(),now()) ON CONFLICT(project_id) DO UPDATE SET content=excluded.content,status='ready',version=outlines.version+1,model=excluded.model,error_message=NULL,generated_at=now(),updated_at=now()""", (project_id, json.dumps(content, ensure_ascii=False), outline_model))
            conn.execute("UPDATE projects SET status='reviewing',progress=100,error_message=NULL,updated_at=now() WHERE id=%s", (project_id,))
            conn.execute("UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s", (job_id,))
            conn.commit()
        print(json.dumps({"event": "outline_generated", "projectId": project_id, "chapters": len(clean), "requirements": len(requirements), "outlineMode": outline_mode, "outlineModel": outline_model, "leafSections": len(outline_leaf_entries(clean)), "specificityAudit": specificity_audit, "volumeBoundaryAudit": volume_boundary_audit}, ensure_ascii=False), flush=True)
    except Exception as exc:
        with psycopg.connect(DB) as conn:
            conn.execute("""INSERT INTO outlines(project_id,status,error_message,updated_at) VALUES(%s,'failed',%s,now()) ON CONFLICT(project_id) DO UPDATE SET status='failed',error_message=excluded.error_message,updated_at=now()""", (project_id, str(exc)[:500]))
            conn.commit()
        fail_job(job_id, project_id, exc)


def outline_node_at(content, path):
    nodes = content.get("chapters") or []; node = None
    for index in path:
        if not isinstance(index, int) or index < 0 or index >= len(nodes): return None
        node = nodes[index]
        nodes = node.get("children") or []
    return node


def collect_node_requirement_ids(node):
    ids = list(node.get("requirementIds") or [])
    for child in node.get("children") or []:
        ids.extend(collect_node_requirement_ids(child))
    return list(dict.fromkeys(str(value) for value in ids))


def editor_relevant_requirements(node, requirement_map):
    """Reject legacy zero-score or cross-lane links before model generation."""
    entry = {"path": (), "node": node}
    section_lane = section_route_lane(node)
    selected = []
    rejected = []
    for requirement_id in collect_node_requirement_ids(node):
        item = requirement_map.get(str(requirement_id))
        if not item:
            continue
        item_lane = requirement_route_lane(item)
        score = route_match_score(item, entry)
        lane_matches = (
            item_lane == section_lane
            or (
                item_lane == "technical_parameter"
                and section_lane == "technical_solution"
            )
        )
        if lane_matches and score > 0:
            selected.append(item)
        else:
            rejected.append({
                "id": str(requirement_id),
                "title": item.get("title"),
                "requirementLane": item_lane,
                "sectionLane": section_lane,
                "score": score,
            })
    return selected, rejected


def safe_section_content(node, requirements):
    lines = [f"## {node.get('title', '响应方案')}", "", node.get("description") or "本节根据招标文件要求编制。", ""]
    if requirements:
        lines.extend(["### 招标要求响应", ""])
        for index, item in enumerate(requirements, 1):
            lines.append(f"{index}. **{item['title']}**")
            lines.append(f"   - 招标要求：{item['requirement']}")
            lines.append("   - 响应说明：我方将严格按照招标文件要求执行，具体证明材料、参数和实施细节将在定稿时结合投标方资料补充。")
    else:
        lines.extend(["### 编写要点", "", "- 按本章主题补充投标方实际方案、证明材料和项目数据。", "- 所有承诺和参数应与招标文件及投标方真实能力保持一致。"])
    lines.extend(["", "### 编制提示", "", "本内容为响应初稿，涉及企业资质、人员、业绩、报价和具体参数的内容需由投标人核实后定稿。"])
    return "\n".join(lines)


def section_source_context(node, blocks, blueprint=None, limit=9000):
    profile = (blueprint or {}).get("projectProfile") or {}
    seed = (
        str(node.get("title") or "")
        + str(node.get("description") or "")
        + " ".join(profile.get("scope") or [])
        + " ".join(profile.get("objectives") or [])
        + " ".join(profile.get("deliverables") or [])
    )
    terms = set(re.findall(r"[\u4e00-\u9fff]{2,8}", seed))
    generic = {"项目", "方案", "管理", "服务", "工作", "实施", "要求", "保障", "措施", "说明", "提供", "进行", "相关", "系统"}
    terms = {term for term in terms if term not in generic}
    ranked = []
    for index, block in enumerate(blocks):
        text = str(block.get("text") or "")
        score = sum(text.count(term) * (2 + len(term)) for term in terms)
        if score:
            ranked.append((score, index, block))
    if not ranked:
        ranked = [(0, index, block) for index, block in enumerate(blocks[:30])]
    selected = []
    seen = set()
    for _, index, _ in sorted(ranked, reverse=True)[:8]:
        for nearby in range(max(0, index - 1), min(len(blocks), index + 2)):
            block = blocks[nearby]
            block_id = str(block.get("id") or nearby)
            if block_id not in seen:
                selected.append(f"[{block_id}] {block.get('text', '')}")
                seen.add(block_id)
    return "\n".join(selected)[:limit]


REQUIREMENT_ATOM_CATEGORIES = (
    ("prohibition", ("不得", "严禁", "禁止", "无效", "否决", "拒绝")),
    ("frequency", ("每日", "每天", "每周", "每月", "每季度", "每年", "小时", "分钟", "频次", "定期")),
    ("record", ("记录", "台账", "报告", "归档", "留痕", "登记", "报表")),
    ("resource", ("人员", "岗位", "设备", "工具", "耗材", "车辆", "配置", "物资")),
    ("acceptance", ("验收", "考核", "检查", "复核", "抽查", "合格", "达标", "扣分")),
    ("scope", ("范围", "包括", "区域", "场地", "楼宇", "道路", "绿化", "服务内容")),
    ("standard", ("保持", "确保", "及时", "洁净", "整洁", "完好", "无明显", "符合")),
)


def requirement_atom_category(text):
    value = str(text or "")
    if re.search(r"扣\s*\d*(?:\.\d+)?\s*分", value):
        return "acceptance"
    for category, keywords in REQUIREMENT_ATOM_CATEGORIES:
        if any(keyword in value for keyword in keywords):
            return category
    return "action"


def split_requirement_clauses(text):
    value = re.sub(r"\s+", " ", str(text or "")).strip(" |")
    if not value:
        return []
    value = re.sub(
        r"(?<!\d)(?=(?:\d{1,2}[.、]|[（(][一二三四五六七八九十\d]+[）)]))",
        "\n",
        value,
    )
    parts = re.split(r"(?:[。；;]\s*|\n+|\s+\|\s+)", value)
    clauses = []
    for part in parts:
        clause = re.sub(r"^[|,\s]+|[|,\s]+$", "", part)
        if re.search(r"(?:不|扣|及|和|与|、|[:：])$", clause):
            continue
        if 8 <= len(clause) <= 420 and clause not in clauses:
            clauses.append(clause)
    return clauses


def title_relevance_terms(node):
    title = re.sub(r"[^\u4e00-\u9fff]", "", str((node or {}).get("title") or ""))
    generic = {
        "项目", "方案", "安排", "管理", "服务", "工作", "实施", "使用",
        "要求", "保障", "措施", "内容", "标准", "进行", "相关",
    }
    terms = set()
    for size in (4, 3, 2):
        for index in range(max(0, len(title) - size + 1)):
            term = title[index:index + size]
            if term not in generic and not any(word in term for word in generic):
                terms.add(term)
    return terms


def build_requirement_atoms(source_context, node=None, limit=36):
    """Turn only the section's matched tender excerpts into traceable atoms."""
    source = str(source_context or "")
    matches = list(re.finditer(
        r"(?ms)^\[([^\]]+)\]\s*(.*?)(?=^\[[^\]]+\]\s*|\Z)",
        source,
    ))
    ranked_atoms = []
    relevance_terms = title_relevance_terms(node)
    for match in matches:
        block_id, block_text = match.group(1).strip(), match.group(2).strip()
        for index, clause in enumerate(split_requirement_clauses(block_text)):
            numbers = re.findall(
                r"\d+(?:\.\d+)?(?:\s*[—~～\-至]\s*\d+(?:\.\d+)?)?"
                r"\s*(?:%|％|日历天|工作日|天|小时|分钟|秒|次|人|套|台|辆|"
                r"年|个月|月|MPa|kPa|Pa|m²|㎡|米|万元|元)?",
                clause,
                flags=re.IGNORECASE,
            )
            identity = hashlib.sha1(
                f"{block_id}|{index}|{clause}".encode("utf-8")
            ).hexdigest()[:16]
            relevance = sum(
                1 + min(3, len(term) - 1)
                for term in relevance_terms if term in clause
            )
            if relevance_terms and relevance == 0:
                continue
            atom = {
                "id": f"ra-{identity}",
                "sourceBlockId": block_id,
                "text": clause,
                "category": requirement_atom_category(clause),
                "mandatory": bool(re.search(
                    r"必须|应当|须|不得|严禁|禁止|确保|不低于|不少于",
                    clause,
                )),
                "numbers": [item.strip() for item in numbers if item.strip()],
                "relevance": relevance,
            }
            ranked_atoms.append((relevance, len(clause), atom))
    ranked_atoms.sort(key=lambda item: (item[0], item[1]), reverse=True)
    if relevance_terms and ranked_atoms:
        threshold = min(4, ranked_atoms[0][0])
        ranked_atoms = [
            item for item in ranked_atoms if item[0] >= threshold
        ]
    return [item[2] for item in ranked_atoms[:limit]]


def build_scene_action_cards(node, atoms, limit=12):
    """Build project-specific writing cues; no fixed industry schema is used."""
    cards = []
    grouped = {}
    for atom in atoms or []:
        grouped.setdefault(atom.get("category") or "action", []).append(atom)
    title = str((node or {}).get("title") or "")
    for category, items in grouped.items():
        excerpts = [item["text"] for item in items[:4]]
        block_ids = list(dict.fromkeys(
            item["sourceBlockId"] for item in items if item.get("sourceBlockId")
        ))
        action_terms = []
        for excerpt in excerpts:
            for fragment in re.split(r"[,，；;:：]", excerpt):
                dirty = re.search(
                    r"([^,，；;:：]{1,12})不\s*清洁(?:扣|$)", fragment
                )
                if dirty:
                    action_terms.append(
                        "清洁" + re.sub(r"^\d+[.、]?", "", dirty.group(1)).strip()
                    )
            action_terms.extend(re.findall(
                r"[\u4e00-\u9fff]{1,8}(?:检查|清洁|巡查|记录|复核|维护|"
                r"处置|整改|验收|消毒|保养|报告|移交|配置)",
                excerpt,
            ))
        cards.append({
            "scene": title,
            "focus": category,
            "sourceBlockIds": block_ids,
            "actions": list(dict.fromkeys(action_terms))[:8],
            "requirements": excerpts,
        })
        if len(cards) >= limit:
            break
    return cards


def extract_section_commitments(content, limit=80):
    """Extract reusable quantitative and operational commitments from output."""
    text = str(content or "")
    units = (
        r"%|％|日历天|工作日|天|小时|分钟|秒|次|人|套|台|辆|"
        r"年|个月|月|MPa|kPa|Pa|m²|㎡|米|万元|元"
    )
    value_pattern = re.compile(
        rf"(?:(?:\d+\s*[:：]\s*\d+)(?:\s*[—~～\-至]\s*\d+\s*[:：]\s*\d+)?|"
        rf"\d+(?:\.\d+)?(?:\s*[—~～\-至]\s*\d+(?:\.\d+)?)?\s*(?:{units}))",
        flags=re.IGNORECASE,
    )
    clauses = re.split(r"(?:\n+|[。；;]\s*)", text)
    commitments = []
    for clause in clauses:
        if clause.lstrip().startswith("#"):
            continue
        normalized = re.sub(r"^[#>*\-\s\d.、]+", "", clause).strip()
        values = [item.strip() for item in value_pattern.findall(normalized)]
        word_frequencies = re.findall(
            r"每日|每天|每周|每月|月末|上午|下午|下班前|到岗后|"
            r"使用高峰后|集中(?:取水|洗衣|使用)后|日常开放期间",
            normalized,
        )
        if (
            not values
            and word_frequencies
            and re.search(
                r"检查|巡查|保洁|清理|复查|汇总|提交|记录|完成|确认",
                normalized,
            )
        ):
            values = list(dict.fromkeys(word_frequencies))
        if not values or len(normalized) < 6:
            continue
        if re.search(
            r"每日|每天|每周|每月|月末|上午|下午|下班前|到岗后|"
            r"使用高峰后|集中(?:取水|洗衣|使用)后|日常开放期间|"
            r"每季度|频次|次/",
            normalized,
        ):
            kind = "frequency"
        elif re.search(r"工具|设备|配置|台|套|辆|人", normalized):
            kind = "resource"
        elif re.search(r"压力|浓度|比例|温度|阈值|MPa|kPa|%", normalized):
            kind = "parameter"
        elif re.search(r"\d{1,2}:\d{2}|小时|分钟|秒", normalized):
            kind = "time"
        else:
            kind = "quantitative"
        identity = hashlib.sha1(normalized.encode("utf-8")).hexdigest()[:16]
        commitments.append({
            "id": f"cm-{identity}",
            "kind": kind,
            "values": values,
            "text": normalized[:360],
        })
        if len(commitments) >= limit:
            break
    return commitments


def collect_outline_commitments(content, exclude_path=None, limit=240):
    items = []
    excluded = tuple(exclude_path or ())
    for entry in outline_leaf_entries((content or {}).get("chapters") or []):
        if tuple(entry["path"]) == excluded:
            continue
        node = entry["node"]
        for commitment in node.get("commitments") or []:
            if not isinstance(commitment, dict):
                continue
            items.append({
                **commitment,
                "sectionPath": list(entry["path"]),
                "sectionTitle": node.get("title"),
            })
            if len(items) >= limit:
                return items
    return items


def relevant_outline_commitments(content, node, exclude_path=None, limit=40):
    """Expose only commitments whose original scene overlaps this section."""
    terms = title_relevance_terms(node)
    if not terms:
        return []
    ranked = []
    for item in collect_outline_commitments(content, exclude_path):
        haystack = (
            str(item.get("sectionTitle") or "")
            + str(item.get("text") or "")
        )
        score = sum(
            1 + min(3, len(term) - 1)
            for term in terms if term in haystack
        )
        if score >= 4:
            ranked.append((score, item))
    ranked.sort(key=lambda pair: pair[0], reverse=True)
    return [item for _, item in ranked[:limit]]


def refresh_project_commitment_registry(project_id):
    with psycopg.connect(DB) as conn:
        row = conn.execute(
            "SELECT content FROM outlines WHERE project_id=%s FOR UPDATE",
            (project_id,),
        ).fetchone()
        if not row:
            return []
        content = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        items = collect_outline_commitments(content)
        content["commitmentRegistry"] = {
            "version": "section-commitment-registry-v1",
            "updatedAt": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "count": len(items),
            "items": items,
        }
        conn.execute(
            "UPDATE outlines SET content=%s::jsonb,updated_at=now() WHERE project_id=%s",
            (json.dumps(content, ensure_ascii=False), project_id),
        )
        conn.commit()
    return items


def clean_model_text_artifacts(value):
    removed = []

    def drop_latin_tail(match):
        removed.append(match.group(1))
        return ""

    cleaned = re.sub(
        r"(?<=[\u4e00-\u9fff])([a-z]{3,})(?=$|[\u4e00-\u9fff，。；：！？、）】\s])",
        drop_latin_tail,
        str(value or ""),
        flags=re.MULTILINE,
    )
    return cleaned, removed


FORMAL_BID_VOICE_MARKERS = {
    "从招标文件看": re.compile(r"从(?:本项目)?(?:采购|磋商|招标)文件(?:的)?(?:要求)?(?:可以)?看"),
    "材料解读口吻": re.compile(r"(?:根据|结合)(?:上述|输入|所给)(?:材料|内容)(?:可以)?(?:看出|判断)"),
    "旁观者结论": re.compile(r"(?:由此可见|可以看出|这说明|这意味着|需要指出的是|换言之)"),
    "评价采购人": re.compile(r"(?:体现|反映)出采购人对[^。；\n]{2,80}(?:重视|关注)"),
    "不是而是论述": re.compile(r"本项目(?:并非|不是)[^。；\n]{2,100}而是"),
    "不仅更是论述": re.compile(r"本项目不仅是[^。；\n]{2,100}更是"),
    "咨询建议口吻": re.compile(r"(?:建议采用|建议设置|建议配置|建议建立)"),
    "内部任务复述": re.compile(r"(?:围绕上述[^。；\n]{2,100}(?:展开|分解|建立|构建)|把本模块展开为|本模块(?:主要|重点)?(?:说明|阐述|展开))"),
    "目录路径复述": re.compile(r"围绕“[^”—\n]{2,180}(?:—[^”—\n]{2,180})+”(?:这一主线)?"),
    "写作过程说明": re.compile(r"(?:本节将|下文将|以下(?:将|从)[^。；\n]{0,40}(?:说明|阐述|展开|介绍))"),
    "空泛总起句": re.compile(r"(?:本项目服务质量的保障，核心在于|为确保[^。；\n]{2,80}，我方将围绕)"),
}


def formal_bid_voice_issues(value):
    text = str(value or "")
    return [
        label
        for label, pattern in FORMAL_BID_VOICE_MARKERS.items()
        if pattern.search(text)
    ]


def normalize_formal_bid_voice(value):
    """Safely remove model meta-language without changing technical facts."""
    text = str(value or "")
    replacements = [
        (
            re.compile(
                r"本项目[^。；\n]{0,80}围绕“[^”—\n]{2,180}"
                r"(?:—[^”—\n]{2,180})+”[^，。；\n]{0,40}展开[，,]将"
            ),
            "",
        ),
        (
            re.compile(
                r"我方将围绕上述业务范围(?:与|和)工作对象[，,]"
                r"将整体服务任务分解为可执行的连续动作[，,]"
                r"并明确各岗位的衔接关系[，,]确保服务启动有序、日常运行稳定。?"
            ),
            "服务任务按照作业对象、实施时序和岗位职责分解，前后岗位依照交接标准连续衔接。",
        ),
        (
            re.compile(
                r"为确保服务期内各项作业有序衔接、责任清晰、全程可追溯[，,]"
                r"我方将围绕日常作业执行、关键控制节点、多方协作接口及过程证据留存"
                r"建立统一的工作机制。?"
            ),
            "日常作业、关键控制节点、多方协作接口及过程证据统一纳入项目运行记录。",
        ),
        (
            re.compile(
                r"本项目服务质量的保障[，,]核心在于将招标文件规定的各项服务标准[，,]"
                r"转化为可执行、可检查、可判定的过程控制体系。?"
            ),
            "招标文件规定的各项服务标准分解为可执行、可检查、可判定的过程控制要求。",
        ),
        (
            re.compile(
                r"我方将围绕事前预防、过程检查、成果复核和问题整改四个环节[，,]"
                r"建立覆盖全部服务范围的质量管控机制[，,]"
                r"确保服务成果持续稳定地满足采购人要求。?"
            ),
            "质量管控按照事前预防、过程检查、成果复核和问题整改四个环节运行，覆盖全部服务范围。",
        ),
        (
            re.compile(
                r"我方将围绕上述业务范围(?:与|和)工作对象[，,]"
                r"建立从异常发现、快速处置到成果复核、持续改进的完整工作链条。?"
            ),
            "异常处置依次执行异常发现、快速处置、成果复核和持续改进四个环节。",
        ),
        (
            re.compile(
                r"为确保[^。；\n]{2,120}[，,]"
                r"我方将围绕“[^”\n]{2,240}”(?:这一主线)?[，,]\s*"
            ),
            "",
        ),
        (
            re.compile(
                r"我方将围绕“[^”\n]{2,240}”(?:这一主线)?[，,]\s*"
            ),
            "",
        ),
        (
            re.compile(r"从(?:本项目)?(?:采购|磋商|招标)文件(?:的)?(?:要求)?(?:可以)?看[，,：:]?\s*"),
            "",
        ),
        (
            re.compile(r"(?:由此可见|可以看出|这说明|这意味着|需要指出的是|换言之)[，,：:]?\s*"),
            "",
        ),
        (
            re.compile(
                r"本项目(?:并非|不是)[^，。；\n]{2,100}[，,]\s*"
                r"而是(?=(?:由|涵盖|包含|包括|涉及|要求))"
            ),
            "本项目",
        ),
        (
            re.compile(r"本项目不仅是[^，。；\n]{2,100}[，,]\s*更是"),
            "本项目是",
        ),
        (
            re.compile(
                r"我方将围绕上述[^，。；\n]{2,100}[，,]\s*(?:将)?"
            ),
            "",
        ),
        (
            re.compile(r"本项目服务质量的保障，核心在于将"),
            "",
        ),
        (
            re.compile(
                r"我方将围绕(?:事前|事中|事后|日常|专项|过程)[^，。；\n]{2,100}"
                r"(?:建立|构建)"
            ),
            "本项目建立",
        ),
        (
            re.compile(r"(?:本节将|下文将)[^。；\n]{2,100}[。；]?\s*"),
            "",
        ),
        (re.compile(r"我们将"), "我方将"),
        (re.compile(r"我们拟"), "我方拟"),
    ]
    applied = []
    for pattern, replacement in replacements:
        text, count = pattern.subn(replacement, text)
        if count:
            applied.append({"pattern": pattern.pattern, "count": count})
    return text, applied


GENERIC_BID_PHRASES = (
    "全面提升",
    "切实保障",
    "有序推进",
    "形成闭环",
    "建立健全",
    "进一步提高",
    "提供有力保障",
    "确保项目顺利实施",
    "具有重要意义",
    "奠定坚实基础",
)
ACTION_MARKERS = (
    "我方将", "本方案拟", "采用", "设置", "配置", "部署", "安装", "调试",
    "编制", "形成", "提交", "复核", "校验", "检查", "测试", "记录", "交付",
    "培训", "巡检", "整改", "验收",
)
VERIFICATION_MARKERS = (
    "检查", "复核", "校验", "测试", "验收", "确认", "记录", "报告", "清单",
    "台账", "图纸", "拓扑图", "点位图", "文档", "交付物", "签认", "试运行",
)
BIDDER_ACTOR_MARKERS = (
    "我方", "项目经理", "项目负责人", "保洁主管", "保洁人员",
    "宿管岗位", "维修岗位", "水电工", "作业人员",
)
META_EXPLANATION_MARKERS = (
    "招标要求明确", "对应考核要求", "评分中", "最高扣", "从招标文件要求看",
    "从招标文件看", "本节将", "下文将", "以下进行说明",
)
AWKWARD_BID_TERMS = (
    "通行面", "公共面", "外露面", "可视硬化区域", "巡回查看",
    "纳入公共区域保洁基础范围", "统筹边界确认", "组织现场落实",
    "按室外公共场地管理", "编入室外巡查路线",
    "接管浴室现场", "作业启动节点", "完成状态为", "现场岗位安排",
)
GENERIC_SECTION_HEADINGS = (
    "作业准备", "操作步骤", "安全控制", "复核确认",
    "工作目标", "实施原则", "保障措施", "持续改进",
)
ANALYSIS_SECTION_MARKERS = (
    "背景", "意义", "理解", "概述", "总体认识", "需求分析", "现状", "目标",
)
ANCHOR_PATTERN = re.compile(
    r"(?:[A-Za-z][A-Za-z0-9.+_-]{1,24}|"
    r"[\u4e00-\u9fffA-Za-z0-9]{2,20}"
    r"(?:系统|设备|软件|平台|方案|报告|清单|台账|图|表|室|墙|工程|"
    r"服务|作业|流程|岗位|场所|区域|标准|条款|资料|人员|考核|验收))"
)


def _quality_sentences(value):
    return [
        re.sub(r"[\s#*_`>|-]+", "", item)
        for item in re.split(r"[。！？；\n]+", str(value or ""))
        if len(re.sub(r"[\s#*_`>|-]+", "", item)) >= 18
    ]


def _trigram_similarity(left, right):
    def grams(value):
        return {value[index:index + 3] for index in range(max(0, len(value) - 2))}

    left_grams, right_grams = grams(left), grams(right)
    if not left_grams or not right_grams:
        return 0
    return len(left_grams & right_grams) / len(left_grams | right_grams)


def evaluate_section_quality(content, node=None, source_context="", requirements=None):
    text = str(content or "")
    title = str((node or {}).get("title") or "")
    characters = len(re.sub(r"\s+", "", text))
    scale = max(characters / 1000, 0.5)
    sentences = _quality_sentences(text)[:90]
    paragraphs = [
        item.strip()
        for item in re.split(r"\n\s*\n", text)
        if item.strip() and not item.lstrip().startswith("#")
    ]
    action_count = sum(text.count(marker) for marker in ACTION_MARKERS)
    verification_count = sum(text.count(marker) for marker in VERIFICATION_MARKERS)
    generic_hits = [
        {"phrase": phrase, "count": text.count(phrase)}
        for phrase in GENERIC_BID_PHRASES
        if phrase in text
    ]
    repeated_pairs = []
    for left_index, left in enumerate(sentences):
        for right_index in range(left_index + 1, len(sentences)):
            right = sentences[right_index]
            if abs(len(left) - len(right)) > max(len(left), len(right)) * 0.35:
                continue
            similarity = _trigram_similarity(left, right)
            if similarity >= 0.72:
                repeated_pairs.append({
                    "similarity": round(similarity, 2),
                    "first": left[:42],
                    "second": right[:42],
                })
                if len(repeated_pairs) >= 5:
                    break
        if len(repeated_pairs) >= 5:
            break

    source = str(source_context or "") + "\n" + "\n".join(
        str(item.get("title") or "") + " " + str(item.get("requirement") or "")
        for item in (requirements or [])
    )
    source += "\n" + "\n".join(
        str(item) for item in ((node or {}).get("generationFactAnchors") or [])
    )
    anchors = {
        match.group(0)
        for match in ANCHOR_PATTERN.finditer(source)
        if len(match.group(0)) >= 3
    }
    matched_anchors = sorted(
        (anchor for anchor in anchors if anchor in text),
        key=lambda value: (-len(value), value),
    )
    anchor_target = min(12, max(4, len(anchors) // 5)) if anchors else 0
    quantitative_count = len(re.findall(
        r"(?:\d+(?:\.\d+)?\s*(?:%|％|小时|分钟|日|天|年|台|套|项|个)|"
        r"[A-Za-z][A-Za-z0-9.+_-]{1,20})",
        text,
    ))
    table_rows = sum(
        1 for line in text.splitlines()
        if line.strip().startswith("|") and line.count("|") >= 3
    )
    analysis_section = any(marker in title for marker in ANALYSIS_SECTION_MARKERS)
    section_form = editor_section_form(node)
    boundary_section = section_form["id"] == "scope_definition"
    action_target = 3.0 if analysis_section or boundary_section else 7.0
    verification_target = 1.5 if boundary_section else 2.0 if analysis_section else 5.0
    action_density = action_count / scale
    verification_density = verification_count / scale
    bidder_actor_count = sum(text.count(marker) for marker in BIDDER_ACTOR_MARKERS)
    meta_explanation_hits = [
        {"phrase": phrase, "count": text.count(phrase)}
        for phrase in META_EXPLANATION_MARKERS
        if phrase in text
    ]
    awkward_term_hits = [
        {"phrase": phrase, "count": text.count(phrase)}
        for phrase in AWKWARD_BID_TERMS
        if phrase in text
    ]
    generic_heading_hits = [
        heading
        for heading in GENERIC_SECTION_HEADINGS
        if re.search(rf"^#{{2,5}}\s+(?:[一二三四五六七八九十\d.、 ]+)?{re.escape(heading)}\s*$", text, re.MULTILINE)
    ]
    grounding_score = (
        round(25 * min(1, len(matched_anchors) / anchor_target))
        if anchor_target else 15
    )
    action_score = round(25 * min(1, action_density / action_target))
    verification_score = round(
        20 * min(1, verification_density / verification_target)
    )
    specificity_score = round(15 * min(
        1,
        (len(matched_anchors) + min(quantitative_count, 6) + min(table_rows, 3) * 2)
        / 10,
    ))
    discipline_score = max(
        0,
        15
        - sum(item["count"] for item in generic_hits) * 2
        - len(repeated_pairs) * 3
        - sum(item["count"] for item in meta_explanation_hits) * 3
        - sum(item["count"] for item in awkward_term_hits) * 4
        - len(generic_heading_hits) * 2,
    )
    score = max(0, min(
        100,
        grounding_score
        + action_score
        + verification_score
        + specificity_score
        + discipline_score,
    ))
    issues = []
    if anchor_target and len(matched_anchors) < max(2, anchor_target // 2):
        issues.append({
            "code": "weak_grounding",
            "label": "项目依据偏少",
            "detail": "正文引用的项目专属对象和招标依据不足。",
        })
    if action_density < action_target * 0.55:
        issues.append({
            "code": "weak_action",
            "label": "实施动作偏少",
            "detail": "正文解释较多，但缺少我方拟采取的具体动作。",
        })
    if verification_density < verification_target * 0.5:
        issues.append({
            "code": "weak_verification",
            "label": "验证闭环偏少",
            "detail": "检查、测试、记录、交付或验收依据不足。",
        })
    if sum(item["count"] for item in generic_hits) >= 3:
        issues.append({
            "code": "generic_language",
            "label": "套话密度偏高",
            "detail": "、".join(item["phrase"] for item in generic_hits[:4]),
        })
    if repeated_pairs:
        issues.append({
            "code": "repetition",
            "label": "存在近似重复",
            "detail": f"检测到 {len(repeated_pairs)} 组近似句。",
        })
    if meta_explanation_hits:
        issues.append({
            "code": "meta_explanation",
            "label": "存在解说或评分口吻",
            "detail": "、".join(item["phrase"] for item in meta_explanation_hits[:4]),
        })
    if awkward_term_hits:
        issues.append({
            "code": "awkward_collocation",
            "label": "存在生造搭配",
            "detail": "、".join(item["phrase"] for item in awkward_term_hits[:4]),
        })
    if generic_heading_hits:
        issues.append({
            "code": "generic_headings",
            "label": "小标题过于通用",
            "detail": "、".join(generic_heading_hits[:4]),
        })
    if not analysis_section and bidder_actor_count < 2:
        issues.append({
            "code": "weak_bidder_voice",
            "label": "投标人主体不清",
            "detail": "实施安排缺少我方或具体执行岗位作为动作主体。",
        })
    if characters < 900:
        issues.append({
            "code": "thin_content",
            "label": "正文内容偏薄",
            "detail": "章节正文不足以支撑完整技术响应。",
        })
    grade = "优秀" if score >= 85 else "良好" if score >= 75 else "可用" if score >= 65 else "需优化"
    return {
        "score": score,
        "grade": grade,
        "sectionType": section_form["label"],
        "dimensions": {
            "projectGrounding": grounding_score,
            "actionability": action_score,
            "verifiability": verification_score,
            "specificity": specificity_score,
            "discipline": discipline_score,
        },
        "issues": issues[:5],
        "metrics": {
            "characters": characters,
            "paragraphs": len(paragraphs),
            "actionCount": action_count,
            "verificationCount": verification_count,
            "matchedAnchors": matched_anchors[:12],
            "genericPhraseCount": sum(item["count"] for item in generic_hits),
            "repeatedPairs": len(repeated_pairs),
            "bidderActorCount": bidder_actor_count,
            "metaExplanationCount": sum(
                item["count"] for item in meta_explanation_hits
            ),
            "awkwardTermCount": sum(
                item["count"] for item in awkward_term_hits
            ),
            "genericHeadingCount": len(generic_heading_hits),
        },
    }


def repair_unsupported_claims(content, unsupported, project_id, job_id):
    system = """你是投标文件事实校对编辑。输入草稿中列出的数字、时限或实质性措辞没有招标依据。
在完整保留章节结构、专业方法和表格的前提下重写全文：删除这些无依据表述，或改成“按招标文件要求”“按采购人确认结果”“及时”“定期”等审慎表达；不得新增任何数字、百分比、时限、人数、设备数量、签字盖章要求或量化承诺。只输出修订后的完整正文，不要解释。"""
    user = json.dumps(
        {
            "unsupportedClaims": unsupported,
            "draft": content,
        },
        ensure_ascii=False,
    )
    body = make_chat_body(
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        3600,
        0,
    )
    started = time.monotonic()
    run_id = str(uuid.uuid4())
    input_tokens = output_tokens = 0
    try:
        request = urllib.request.Request(
            AI_URL,
            data=body,
            headers={
                "Authorization": f"Bearer {AI_KEY}",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        with urllib.request.urlopen(request, timeout=240) as response:
            payload = json.loads(response.read().decode("utf-8"))
        usage = payload.get("usage") or {}
        input_tokens = int(usage.get("prompt_tokens") or 0)
        output_tokens = int(usage.get("completion_tokens") or 0)
        choice = payload["choices"][0]
        if choice.get("finish_reason") == "length":
            raise ValueError("事实声明修订达到长度上限")
        repaired = str(choice["message"]["content"] or "").strip()
        if len(repaired) < 700:
            raise ValueError("事实声明修订结果过短")
        save_run(
            run_id,
            project_id,
            job_id,
            1,
            "succeeded",
            input_tokens,
            output_tokens,
            int((time.monotonic() - started) * 1000),
            0,
            run_type="section_repair",
        )
        return repaired
    except Exception as exc:
        save_run(
            run_id,
            project_id,
            job_id,
            1,
            "failed",
            input_tokens,
            output_tokens,
            int((time.monotonic() - started) * 1000),
            0,
            exc,
            run_type="section_repair",
        )
        raise


def is_controlled_bidder_commitment(content, claim):
    unit_match = re.search(
        r"(?:%|％|个?工作日|天|小时|分钟|秒|次|人|套|台|家|年|个月|月)\s*$",
        str(claim or ""),
    )
    if not unit_match:
        return False
    unit = unit_match.group(0)
    if unit in {"人", "套", "台", "家", "年", "个月", "月", "天", "工作日", "个工作日"}:
        return False
    markers = (
        "我方承诺", "我方内部", "内部控制标准", "内部管理指标",
        "本方案设定", "本方案拟", "拟实行", "拟采用", "拟设置",
        "执行频次", "巡检频次", "检查频次", "培训安排",
        "质量目标", "服务目标", "响应分级",
    )
    tender_fact_markers = (
        "招标要求", "采购人规定", "采购文件规定", "合同约定",
        "服务期限", "工期为", "必须达到", "不得低于",
    )
    start = 0
    found = False
    while True:
        index = content.find(claim, start)
        if index < 0:
            break
        window = content[max(0, index - 100): min(len(content), index + len(claim) + 80)]
        if (
            any(marker in window for marker in markers)
            and not any(marker in window for marker in tender_fact_markers)
        ):
            found = True
            break
        start = index + len(claim)
    return found


def choose_section_text_model(node, requirements, blueprint, model_mode):
    if model_mode == "deepseek":
        return DEEPSEEK_TEXT_MODEL
    if model_mode == "gpt":
        return GPT_TEXT_MODEL

    title = str(node.get("title") or "")
    description = str(node.get("description") or "")
    section_text = title + " " + description
    deepseek_signals = (
        "资格", "商务", "报价", "价格", "响应表", "偏离表", "参数响应",
        "技术参数", "合同", "承诺函", "证明材料", "证书", "业绩", "人员简历",
        "法定代表人", "授权委托", "中小企业", "残疾人福利", "监狱企业",
    )
    gpt_signals = (
        "项目理解", "背景", "建设目标", "总体方案", "总体架构", "重点难点",
        "技术方案", "实施方案", "组织实施", "进度", "质量", "风险", "应急",
        "培训", "售后", "服务保障", "运维", "验收", "创新", "优化",
    )
    if any(signal in section_text for signal in deepseek_signals):
        return DEEPSEEK_TEXT_MODEL
    if any(signal in section_text for signal in gpt_signals):
        return GPT_TEXT_MODEL

    requirement_types = {str(item.get("type") or "") for item in requirements}
    if "scoring" in requirement_types and not requirement_types.intersection({"qualification", "commercial"}):
        return GPT_TEXT_MODEL
    brief = node.get("brief") if isinstance(node.get("brief"), dict) else {}
    form_factor = str(brief.get("formFactor") or "")
    if any(signal in form_factor for signal in ("analysis", "method", "plan", "narrative")):
        return GPT_TEXT_MODEL
    return DEEPSEEK_TEXT_MODEL


def call_section_model(node, requirements, project_name, source_context, project_id, job_id, blueprint, section_brief, generation_mode="deep"):
    if generation_mode == "deep":
        system = """你是政府采购投标文件技术方案主笔。根据项目名称、章节目标、关联招标要求和原文摘录，撰写本章节可直接编辑的专业正文。
写作要求：
1. 章节任务卡是后台写作指导，不是正文模板。必须按照任务卡中的formFactor和writingPattern选择自然段落、步骤、清单或表格；严禁逐项重复“目标、输入、实施步骤、责任角色、输出成果、质量关卡、风险处置、过程记录”。篇幅以sectionBrief.pageBudget中的目标页数和目标字符数为准，允许因内容完整性小幅浮动，但不要为了凑字数重复展开。若sectionBrief含componentContext，本次只撰写其中指定的写作组件，不得提前复述其他组件，也不得输出整章总结。
2. 可以运用通用的项目管理与专业实施方法，但不得虚构采购文件没有给出的期限、金额、数量、品牌、技术参数或法律结论。
3. 不得虚构投标人的资质、人员姓名、证书、业绩、设备数量和服务网点；必须依赖投标人资料的地方写“【待补充：……】”。
4. 对招标文件的明确要求应实质响应，不得改写其数值和含义。允许把不冲突且可履行的作业频次、巡检周期、响应分级、内部检查比例、培训安排、会议机制和质量目标设计为投标人主动承诺；这类内容必须明确写成“我方内部控制标准”“本方案拟实行”或“我方承诺”，不得冒充招标文件原文要求。不得自行编造合同期限、报价、法定时限、固定人员数量、设备数量或采购参数。
5. 项目背景、现状、覆盖范围、建设目标和预期成果只能来自输入摘录，不得根据项目名称自行推断；特别禁止自行添加“达到100%”“全域覆盖”“零事故”等目标。
6. 章节内容必须与项目实施蓝图保持一致，不能另造一套相互矛盾的阶段、角色或交付物；必须优先使用输入中的具体设备、系统、作业环节和验收对象，避免“全面保障、建立机制、形成闭环”等空话连续出现。
7. 同一章节内小节的表达方式应有变化：分析内容用自然段，连续动作使用编号步骤，重复字段才使用表格。不要把普通说明全部做成表格。
8. 本文面向政府采购评审专家，是投标人可直接装订的正式响应正文，不是咨询报告、需求分析报告、对话答复或对招标文件的解读。项目事实使用“本项目……”，投标人的实施动作和承诺使用“我方将……”“本方案拟……”，需采购人决策的事项使用“经采购人确认后……”。
9. 开篇直接陈述项目事实、建设目标或我方实施安排。禁止使用“从招标文件看/从招标文件要求看、根据输入材料可以看出、由此可见、这说明、这意味着、需要指出的是、换言之、体现出采购人重视、本项目并非/不是……而是……、本项目不仅是……更是……”等旁观者分析或宣传式口吻。
10. 不得使用“建议采用、建议设置、建议配置、建议建立”等咨询建议口吻；属于我方方案的改为“本方案拟采用/我方将设置”，属于采购人决策的改为“经采购人确认后实施”。
11. 后台的chapterTitle、writingGoal、componentContext和目录路径只用于约束写作，正文不得复述这些内部任务。禁止出现“围绕上述业务范围”“把本模块展开为”“将整体服务任务分解为可执行的连续动作”“本节将”“下文将”“以下进行说明”“核心在于”等写作过程或解说式表述。
12. 不要在每节开头重复项目全貌。标题是步骤、标准、检查、异常或交接时，直接进入该项作业程序、控制规则或岗位动作；不得先写一段“服务范围涵盖……为确保……我方将……”的总括性前言。
13. “我方将”只用于必要承诺，不得作为连续段落的固定开头。优先使用“实行、采用、设置、执行、记录、复核、移交”等正式制度句和操作句，使正文像既定实施方案，而不是向读者解释方案。
14. 使用Markdown二级或三级标题和正文段落；不要输出写作解释、风险提示或“招标要求/响应初稿”的机械模板。
15. requirementAtoms是本节招标原文的最小响应单元，sceneActionCards是由这些原文动态归纳的场景动作线索。逐项消化与本节有关的原子，但不得把字段名或原子清单原样输出；existingCommitments中的同类数值应保持一致，场景不同确需变化时写明适用条件。"""
        target_characters = int(
            (section_brief.get("pageBudget") or {}).get(
                "targetCharacters", 2800
            )
        )
        max_tokens = max(3200, min(7600, int(target_characters * 1.65)))
        timeout, temperature = 300, 0.2
    else:
        system = """你是政府采购投标文件主笔，以投标人身份撰写可直接编辑的正式响应正文。仅响应输入中明确给出的招标要求。
项目事实使用“本项目……”，履约动作使用“我方将……”“本方案拟……”，需采购人决策的事项使用“经采购人确认后……”。
禁止使用“从招标文件看、可以看出、这说明、本项目并非/不是……而是……、本项目不仅是……更是……”等解读或宣传式口吻，也不得使用“建议采用、建议配置”等咨询口吻。
不得杜撰企业资质、人员、业绩、品牌、参数或承诺；缺少投标方事实时使用“【待补充】”。使用Markdown，只输出正文，不要解释。"""
        max_tokens, timeout, temperature = 1100, 90, 0
    user = json.dumps({
        "projectName": project_name,
        "chapterTitle": node.get("title"),
        "writingGoal": node.get("description"),
        "projectProfile": blueprint.get("projectProfile"),
        "implementationBlueprint": blueprint.get("implementationBlueprint"),
        "sectionBrief": section_brief,
        "enterpriseFactPolicy": (
            "未经投标人资料证明，不得声称已有成熟供应链、长期合作厂家、"
            "现有设备、自有车辆、固定库存、认证人员或既有客服系统；"
            "应改写为本项目拟采取的动作、责任、检查点和验证方式。"
        ),
        "controlledCommitmentPolicy": {
            "allowed": [
                "作业频次与巡检节奏", "响应分级与内部时限",
                "内部抽检比例和服务质量目标", "培训与会议安排",
                "岗位动作、工作流程、记录表单和应急处置步骤",
            ],
            "wording": "明确标注为我方承诺、本方案拟实行或内部控制标准，不得表述为采购人要求",
            "forbidden": [
                "合同期限、报价和法定时限", "固定人员和设备数量",
                "企业资质、证书、业绩、客户名称", "自有设备、库存、网点",
                "品牌型号和招标文件未提供的强制技术参数",
            ],
        },
        "requirements": requirements,
        "requirementAtoms": node.get("requirementAtoms") or [],
        "sceneActionCards": node.get("sceneActionCards") or [],
        "tenderExcerpts": source_context,
    }, ensure_ascii=False)
    body = make_chat_body([{"role": "system", "content": system}, {"role": "user", "content": user}], max_tokens, temperature)
    started = time.monotonic(); run_id = str(uuid.uuid4()); input_tokens = output_tokens = 0; actual_cost = None
    try:
        request = urllib.request.Request(AI_URL, data=body, headers={"Authorization": f"Bearer {AI_KEY}", "Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(request, timeout=timeout) as response:
            client_request_id = response.headers.get("X-Client-Request-Id")
            payload = json.loads(response.read().decode("utf-8"))
        usage = payload.get("usage") or {}; input_tokens = int(usage.get("prompt_tokens") or 0); output_tokens = int(usage.get("completion_tokens") or 0)
        billing = fetch_billing(client_request_id)
        if billing: input_tokens, output_tokens, actual_cost = billing
        choice = payload["choices"][0]
        if choice.get("finish_reason") == "length":
            raise ValueError("模型输出达到长度上限，章节未完整结束")
        content = str(choice["message"]["content"] or "").strip()
        content, removed_artifacts = clean_model_text_artifacts(content)
        content, voice_normalizations = normalize_formal_bid_voice(content)
        if removed_artifacts:
            print(json.dumps({
                "event": "section_text_artifacts_removed",
                "projectId": project_id,
                "artifacts": removed_artifacts[:10],
            }, ensure_ascii=False), flush=True)
        if voice_normalizations:
            print(json.dumps({
                "event": "section_bid_voice_normalized",
                "projectId": project_id,
                "changes": voice_normalizations,
            }, ensure_ascii=False), flush=True)
        voice_issues = formal_bid_voice_issues(content)
        if voice_issues:
            print(json.dumps({
                "event": "section_bid_voice_warning",
                "projectId": project_id,
                "issues": voice_issues,
            }, ensure_ascii=False), flush=True)
        if len(content) < (700 if generation_mode == "deep" else 80): raise ValueError("模型返回的章节正文过短")
        source = (
            "\n".join(item["requirement"] for item in requirements)
            + "\n"
            + source_context
            + "\n"
            + json.dumps(blueprint, ensure_ascii=False)
            + "\n"
            + json.dumps(section_brief, ensure_ascii=False)
        )
        sensitive = [
            "放弃", "法律责任", "强制采购", "最低价", "授权代表",
            "加盖公章", "分项报价", "成熟的供应链", "成熟供应链",
            "长期合作厂家", "长期合作供应商", "现有设备", "自有车辆",
            "自有车队", "自有仓库", "固定库存", "厂家认证",
            "ISO9001", "ISO 9001", "集团客服系统",
        ]
        unsupported_text = [
            phrase for phrase in sensitive if phrase in content and phrase not in source
        ]
        numeric_claims = re.findall(
            r"\d+(?:\.\d+)?\s*(?:%|％|个?工作日|天|小时|分钟|秒|次|人|套|台|家|年|个月|月)",
            content,
        )
        unsupported_numbers = [
            claim for claim in numeric_claims
            if (
                norm(claim) not in norm(source)
                and not is_controlled_bidder_commitment(content, claim)
            )
        ]
        unsupported_claims = list(
            dict.fromkeys(unsupported_text + unsupported_numbers)
        )
        if unsupported_claims:
            content = repair_unsupported_claims(
                content,
                unsupported_claims[:16],
                project_id,
                job_id,
            )
            content, _ = clean_model_text_artifacts(content)
            content, _ = normalize_formal_bid_voice(content)
            unsupported_after_repair_text = [
                phrase
                for phrase in sensitive
                if phrase in content and phrase not in source
            ]
            if unsupported_after_repair_text:
                raise ValueError(
                    "修订后仍有无依据内容："
                    + "、".join(unsupported_after_repair_text)
                )
            repaired_claims = re.findall(
                r"\d+(?:\.\d+)?\s*(?:%|％|个?工作日|天|小时|分钟|秒|次|人|套|台|家|年|个月|月)",
                content,
            )
            unsupported_after_repair = [
                claim
                for claim in repaired_claims
                if (
                    norm(claim) not in norm(source)
                    and not is_controlled_bidder_commitment(content, claim)
                )
            ]
            if unsupported_after_repair:
                raise ValueError(
                    "修订后仍有无依据量化指标："
                    + "、".join(unsupported_after_repair[:5])
                )
        save_run(run_id, project_id, job_id, 1, "succeeded", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), 0, actual_cost=actual_cost, run_type="section_deep" if generation_mode == "deep" else "section")
        return content, "ai_deep" if generation_mode == "deep" else "ai"
    except Exception as exc:
        save_run(run_id, project_id, job_id, 1, "failed", input_tokens, output_tokens, int((time.monotonic() - started) * 1000), 0, exc, actual_cost, "section_deep" if generation_mode == "deep" else "section")
        if generation_mode == "deep":
            raise RuntimeError(f"AI深度生成暂时失败: {exc}") from exc
        return safe_section_content(node, requirements), "safe_fallback"


LENGTH_MODE_MULTIPLIERS = {
    "standard": 1.0,
    "detailed": 2.15,
    "extended": 3.15,
    "xique": 1.0,
}


COMPONENT_TEMPLATES = {
    "diagnostic_narrative": [
        "项目场景与关键难点成因",
        "分场景应对方法与实施安排",
        "控制验证、异常处置与改进",
        "跨专业协同与资源保障",
        "成果检查与闭环复盘",
    ],
    "organization": [
        "组织体系与管理界面",
        "岗位职责、授权与协作关系",
        "人员投入、替补与沟通升级",
        "跨单位协同和资源调度",
        "履职检查与组织效能改进",
    ],
    "schedule": [
        "阶段划分、任务依赖与里程碑",
        "进度计划执行与动态跟踪",
        "偏差分析、资源调整与纠偏",
        "关键路径和跨专业衔接",
        "节点验证与进度资料归档",
    ],
    "quality_control": [
        "质量目标、责任体系与控制依据",
        "过程检查、成果复核与问题整改",
        "记录留痕、版本控制与验证闭环",
        "关键工序和专业质量控制",
        "质量数据分析与持续改进",
    ],
    "acceptance": [
        "验收对象、依据与前置准备",
        "验收实施、问题整改与复验",
        "成果资料、移交培训与归档",
        "多方验收协同与争议处理",
        "交付后的跟踪验证",
    ],
    "training_plan": [
        "培训对象、场景与能力目标",
        "课程组织、实操演练与资料",
        "考核补训、记录与知识转移",
        "分层培训与现场辅导",
        "培训效果评价与持续支持",
    ],
    "service_process": [
        "服务受理、分类与响应组织",
        "处理升级、现场协同与验证关闭",
        "回访复盘、记录与持续改进",
        "重点场景保障与资源调度",
        "服务质量监测与成果报告",
    ],
    "risk_control": [
        "风险场景识别、分级与责任边界",
        "预防监测、现场控制与预警",
        "应急响应、恢复验证与复盘改进",
        "跨单位联动和资源保障",
        "风险记录、演练与持续优化",
    ],
    "operational_plan": [
        "作业对象、条件与总体安排",
        "分阶段实施方法和关键动作",
        "协同接口、过程控制与成果验证",
        "资源投入与异常替代",
        "记录移交与持续改进",
    ],
    "technical_process": [
        "技术条件、接口与实施准备",
        "安装配置、联调测试与问题整改",
        "试运行验证、验收与技术移交",
        "专业系统协同与变更控制",
        "性能复核和运行优化",
    ],
    "professional_narrative": [
        "项目事实、范围边界与响应目标",
        "实施方法、工作流程与专业措施",
        "协同控制、成果验证与异常处置",
        "资源配置和关键场景深化",
        "质量复核与持续改进",
    ],
}


def section_component_specs(node, requirements, section_brief, length_mode):
    form_factor = str(section_brief.get("formFactor") or "professional_narrative")
    matrix_factors = {
        "qualification_evidence", "commercial_response",
        "compliance_matrix", "technical_response_matrix",
    }
    if length_mode in {"standard", "xique"} or form_factor in matrix_factors:
        return []
    component_count = 3 if length_mode == "detailed" else 5
    titles = COMPONENT_TEMPLATES.get(
        form_factor, COMPONENT_TEMPLATES["professional_narrative"]
    )[:component_count]
    base_budget = section_brief.get("pageBudget") or {}
    total_characters = int(
        (base_budget.get("targetCharacters") or 2800)
        * LENGTH_MODE_MULTIPLIERS.get(length_mode, 1.0)
    )
    must_cover = list(section_brief.get("mustCover") or [])
    requirement_groups = [[] for _ in titles]
    for index, item in enumerate(requirements):
        requirement_groups[index % len(titles)].append(item)
    specs = []
    for index, title in enumerate(titles):
        component_must_cover = must_cover[index::len(titles)]
        target_characters = max(1800, int(total_characters / len(titles)))
        specs.append({
            "index": index,
            "title": title,
            "requirements": requirement_groups[index],
            "mustCover": component_must_cover,
            "targetCharacters": target_characters,
        })
    return specs


def component_section_brief(section_brief, spec, total):
    brief = json.loads(json.dumps(section_brief, ensure_ascii=False))
    brief["purpose"] = spec["title"]
    brief["mustCover"] = spec["mustCover"] or [
        f"围绕“{spec['title']}”形成可直接装订的正式投标正文"
    ]
    target_characters = spec["targetCharacters"]
    brief["pageBudget"] = {
        **(brief.get("pageBudget") or {}),
        "minPages": round(target_characters / 900, 1),
        "targetPages": round(target_characters / 720, 1),
        "maxPages": round(target_characters / 560, 1),
        "targetCharacters": target_characters,
    }
    brief["componentContext"] = {
        "componentIndex": spec["index"] + 1,
        "componentCount": total,
        "componentTitle": spec["title"],
        "scopeRule": "只撰写本组件；不得复述已完成或尚未完成的其他组件",
    }
    return brief


def normalize_component_content(content, component_title):
    lines = str(content or "").splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and re.match(r"^#{1,3}\s+", lines[0].strip()):
        lines.pop(0)
    body = "\n".join(lines).strip()
    return f"### {component_title}\n\n{body}"


def outline_node_json_path(path):
    result = ["chapters"]
    for depth, index in enumerate(path):
        if depth:
            result.append("children")
        result.append(str(index))
    return result


def persist_outline_node(project_id, path, node):
    with psycopg.connect(DB) as conn:
        conn.execute(
            """UPDATE outlines
               SET content=jsonb_set(content,%s::text[],%s::jsonb,true),
                   updated_at=now()
               WHERE project_id=%s""",
            (
                outline_node_json_path(path),
                json.dumps(node, ensure_ascii=False),
                project_id,
            ),
        )
        conn.commit()


def update_outline_node_fields(project_id, path, updates, removals=()):
    with psycopg.connect(DB) as conn:
        row = conn.execute(
            "SELECT content FROM outlines WHERE project_id=%s FOR UPDATE",
            (project_id,),
        ).fetchone()
        if not row:
            return None
        content = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        node = outline_node_at(content, path)
        if not node:
            return None
        for key, value in updates.items():
            node[key] = value
        for key in removals:
            node.pop(key, None)
        conn.execute(
            """UPDATE outlines
               SET content=jsonb_set(content,%s::text[],%s::jsonb,true),
                   updated_at=now()
               WHERE project_id=%s""",
            (
                outline_node_json_path(path),
                json.dumps(node, ensure_ascii=False),
                project_id,
            ),
        )
        conn.commit()
        return node


def component_fact_anchors(spec, source_context, limit=8):
    candidates = []
    for item in spec.get("requirements") or []:
        text = re.sub(
            r"\s+", " ",
            f"{item.get('title', '')}：{item.get('requirement', '')}",
        ).strip()
        if text:
            candidates.append((100, text[:220]))
    focus = (
        spec.get("title", "")
        + " "
        + " ".join(spec.get("mustCover") or [])
        + " "
        + " ".join(
            str(item.get("title") or "")
            for item in spec.get("requirements") or []
        )
    )
    focus_terms = set(re.findall(r"[\u4e00-\u9fff]{2,6}", focus))
    for line in str(source_context or "").splitlines():
        text = re.sub(r"^\[[^\]]+\]\s*", "", line).strip()
        text = re.sub(r"\s+", " ", text)
        if len(text) < 18:
            continue
        score = sum(2 + len(term) for term in focus_terms if term in text)
        if re.search(r"\d|不低于|不超过|应当|必须|严禁|验收|考核", text):
            score += 4
        if score:
            candidates.append((score, text[:220]))
    selected = []
    seen = set()
    for _, text in sorted(candidates, key=lambda value: -value[0]):
        key = norm(text)
        if not key or key in seen:
            continue
        seen.add(key)
        selected.append(text)
        if len(selected) >= limit:
            break
    return selected


def merge_component_contents(components):
    accepted = []
    comparable = []
    removed = []
    for component in components:
        blocks = re.split(r"\n\s*\n", str(component or "").strip())
        kept = []
        for block in blocks:
            clean = re.sub(r"\s+", " ", block).strip()
            is_heading = bool(re.match(r"^#{1,6}\s+", clean))
            is_table = clean.startswith("|")
            if not clean:
                continue
            duplicate = None
            if not is_heading and not is_table and len(clean) >= 45:
                for previous in comparable[-80:]:
                    clean_norm = norm(clean)
                    previous_norm = norm(previous)
                    clean_numbers = set(re.findall(
                        r"[A-Za-z]+[-_]?\d+(?:\.\d+)*|\d+(?:\.\d+)?",
                        clean,
                    ))
                    previous_numbers = set(re.findall(
                        r"[A-Za-z]+[-_]?\d+(?:\.\d+)*|\d+(?:\.\d+)?",
                        previous,
                    ))
                    if clean_numbers != previous_numbers:
                        continue
                    clean_anchors = set(ANCHOR_PATTERN.findall(clean))
                    previous_anchors = set(ANCHOR_PATTERN.findall(previous))
                    if (
                        clean_anchors and previous_anchors
                        and clean_anchors != previous_anchors
                    ):
                        continue
                    ratio = SequenceMatcher(
                        None, clean_norm, previous_norm, autojunk=False
                    ).ratio()
                    if clean_norm == previous_norm or ratio >= 0.94:
                        duplicate = {"similarity": round(ratio, 2), "text": clean[:100]}
                        break
            if duplicate:
                removed.append(duplicate)
                continue
            kept.append(block.strip())
            if not is_heading and not is_table and len(clean) >= 45:
                comparable.append(clean)
        if kept:
            accepted.append("\n\n".join(kept))
    return "\n\n".join(accepted), {
        "componentCount": len(components),
        "removedNearDuplicateBlocks": len(removed),
        "removedSamples": removed[:8],
    }


def call_section_with_length_mode(
    node, requirements, project_name, source_context, project_id, job_id,
    blueprint, section_brief, generation_mode, length_mode, path,
    target_model,
):
    specs = section_component_specs(
        node, requirements, section_brief, length_mode
    )
    if not specs:
        grounded_brief = json.loads(json.dumps(
            section_brief, ensure_ascii=False
        ))
        grounded_brief["factAnchors"] = component_fact_anchors({
            "title": node.get("title") or "",
            "mustCover": grounded_brief.get("mustCover") or [],
            "requirements": requirements,
        }, source_context)
        grounded_brief["groundingRule"] = (
            "优先使用factAnchors中的项目事实和原文约束；"
            "不得改写数值、期限、对象或考核含义"
        )
        node["generationFactAnchors"] = grounded_brief["factAnchors"]
        text, mode = call_section_model(
            node, requirements, project_name, source_context,
            project_id, job_id, blueprint, grounded_brief, generation_mode,
        )
        merged, merge_audit = merge_component_contents([text])
        node["componentMergeAudit"] = merge_audit
        return merged, mode, 1
    checkpoint = node.get("generationCheckpoint")
    if not isinstance(checkpoint, dict) or any((
        checkpoint.get("jobId") != job_id,
        checkpoint.get("lengthMode") != length_mode,
        checkpoint.get("model") != target_model,
        checkpoint.get("componentCount") != len(specs),
    )):
        checkpoint = {
            "version": "component-checkpoint-v1",
            "jobId": job_id,
            "lengthMode": length_mode,
            "model": target_model,
            "componentCount": len(specs),
            "components": [],
        }
    saved_components = {
        int(item.get("index")): item
        for item in checkpoint.get("components") or []
        if isinstance(item, dict)
        and item.get("status") == "ready"
        and isinstance(item.get("content"), str)
    }
    components = []
    all_fact_anchors = []
    for spec in specs:
        saved = saved_components.get(spec["index"])
        if saved:
            components.append(saved["content"])
            all_fact_anchors.extend(saved.get("factAnchors") or [])
            continue
        brief = component_section_brief(section_brief, spec, len(specs))
        brief["factAnchors"] = component_fact_anchors(
            spec, source_context
        )
        all_fact_anchors.extend(brief["factAnchors"])
        brief["groundingRule"] = (
            "优先使用factAnchors中的项目对象、考核事项、数量、时限和原文约束；"
            "只引用与本组件相关的事实，不得改写其数值和含义"
        )
        component_node = {
            **node,
            "title": spec["title"],
            "description": (
                f"本组件属于“{node.get('title', '')}”，"
                f"仅撰写{spec['title']}"
            ),
        }
        component_requirements = spec["requirements"]
        text, _ = call_section_model(
            component_node,
            component_requirements,
            project_name,
            source_context,
            project_id,
            job_id,
            blueprint,
            brief,
            generation_mode,
        )
        normalized = normalize_component_content(text, spec["title"])
        components.append(normalized)
        checkpoint["components"] = [
            item for item in checkpoint.get("components") or []
            if int(item.get("index", -1)) != spec["index"]
        ] + [{
            "index": spec["index"],
            "title": spec["title"],
            "status": "ready",
            "characters": len(normalized),
            "factAnchors": brief["factAnchors"],
            "content": normalized,
        }]
        checkpoint["components"].sort(key=lambda item: item["index"])
        node["generationCheckpoint"] = checkpoint
        persist_outline_node(project_id, path, node)
        print(json.dumps({
            "event": "section_component_checkpointed",
            "projectId": project_id,
            "path": path,
            "jobId": job_id,
            "component": spec["index"] + 1,
            "componentCount": len(specs),
            "characters": len(normalized),
        }, ensure_ascii=False), flush=True)
    node["generationFactAnchors"] = list(dict.fromkeys(
        all_fact_anchors
    ))[:24]
    merged, merge_audit = merge_component_contents(components)
    node["componentMergeAudit"] = merge_audit
    return merged, "ai_deep_multi_pass", len(specs)


def process_section(job):
    global MODEL, AI_URL, AI_KEY, AI_THINKING
    job_id, project_id, path = job["jobId"], job["projectId"], job.get("path") or []
    generation_mode = job.get("mode") or "deep"
    model_mode = job.get("modelMode") or "mixed"
    length_mode = job.get("lengthMode") or "standard"
    if length_mode not in LENGTH_MODE_MULTIPLIERS:
        length_mode = "standard"
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s", (job_id,))
        outline_row = conn.execute("SELECT content FROM outlines WHERE project_id=%s", (project_id,)).fetchone()
        requirement_rows = conn.execute("SELECT id,type,title,normalized_value,mandatory FROM requirements WHERE project_id=%s", (project_id,)).fetchall()
        project_row = conn.execute("SELECT name FROM projects WHERE id=%s", (project_id,)).fetchone()
        document_row = conn.execute("SELECT blocks FROM documents WHERE project_id=%s", (project_id,)).fetchone()
        conn.commit()
    try:
        if not outline_row: raise ValueError("项目大纲不存在")
        content = outline_row[0] if isinstance(outline_row[0], dict) else json.loads(outline_row[0])
        node = outline_node_at(content, path)
        if not node: raise ValueError("目标章节不存在")
        requirement_map = {
            str(row[0]): {
                "id": str(row[0]),
                "type": row[1],
                "title": row[2],
                "requirement": row[3],
                "mandatory": bool(row[4]),
            }
            for row in requirement_rows
        }
        requirements, rejected_links = editor_relevant_requirements(
            node, requirement_map
        )
        node["requirementLinkAudit"] = {
            "version": "semantic-zero-reject-v1",
            "acceptedIds": [item["id"] for item in requirements],
            "rejected": rejected_links,
        }
        blocks = document_row[0] if document_row and isinstance(document_row[0], list) else (json.loads(document_row[0]) if document_row else [])
        blueprint = {
            "projectProfile": content.get("projectProfile"),
            "implementationBlueprint": content.get("implementationBlueprint"),
            "projectAnalysis": content.get("projectAnalysis"),
            "scoringTasks": content.get("scoringTasks") or [],
            "capabilityPlan": content.get("capabilityPlan") or [],
        }
        if not isinstance(blueprint["projectProfile"], dict) or not isinstance(blueprint["implementationBlueprint"], dict):
            blueprint = build_safe_blueprint(
                list(requirement_map.values()),
                project_row[0] if project_row else "",
            )
            blueprint["projectAnalysis"] = build_safe_project_analysis(
                list(requirement_map.values()),
                project_row[0] if project_row else "",
                blueprint,
            )
            blueprint["scoringTasks"] = build_safe_scoring_tasks(
                list(requirement_map.values())
            )
            blueprint["capabilityPlan"] = build_capability_plan(
                blueprint["projectAnalysis"], content.get("chapters") or []
            )
            content["projectProfile"] = blueprint["projectProfile"]
            content["implementationBlueprint"] = blueprint["implementationBlueprint"]
            content["projectAnalysis"] = blueprint["projectAnalysis"]
            content["scoringTasks"] = blueprint["scoringTasks"]
            content["capabilityPlan"] = blueprint["capabilityPlan"]
        source_context = section_source_context(node, blocks, blueprint)
        requirement_atoms = build_requirement_atoms(source_context, node)
        scene_action_cards = build_scene_action_cards(node, requirement_atoms)
        existing_commitments = relevant_outline_commitments(
            content, node, path
        )
        node["requirementAtoms"] = requirement_atoms
        node["sceneActionCards"] = scene_action_cards
        section_brief = node.get("brief")
        if not isinstance(section_brief, dict):
            section_brief = build_section_brief(node, requirements, blueprint)
            node["brief"] = section_brief
        section_brief["requirementAtoms"] = requirement_atoms
        section_brief["sceneActionCards"] = scene_action_cards
        section_brief["existingCommitments"] = existing_commitments
        section_brief["structuredInputRule"] = (
            "先逐项消化requirementAtoms，再按sceneActionCards组织项目场景和动作；"
            "existingCommitments中同类数值应保持一致，确需变化时必须有场景理由。"
        )
        target_model = choose_section_text_model(node, requirements, blueprint, model_mode)
        original_model, original_url, original_key, original_thinking = MODEL, AI_URL, AI_KEY, AI_THINKING
        try:
            MODEL = target_model
            if target_model.startswith("gpt-"):
                if not IMAGE_KEY:
                    raise ValueError("sub2api GPT 密钥未配置")
                AI_URL = IMAGE_BASE_URL.rstrip("/") + "/v1/chat/completions"
                AI_KEY = IMAGE_KEY
            else:
                AI_URL = original_url
                AI_KEY = original_key
            AI_THINKING = "disabled"
            text, mode, generation_passes = call_section_with_length_mode(
                node,
                requirements,
                project_row[0] if project_row else "",
                source_context,
                project_id,
                job_id,
                blueprint,
                section_brief,
                generation_mode,
                length_mode,
                path,
                target_model,
            )
        finally:
            MODEL, AI_URL, AI_KEY, AI_THINKING = original_model, original_url, original_key, original_thinking
        node["content"] = text
        node["contentStatus"] = "ready"
        node["contentMode"] = mode
        node["generationModel"] = target_model
        node["generationStrategy"] = model_mode
        node["lengthMode"] = length_mode
        node["generationPasses"] = generation_passes
        node["commitments"] = extract_section_commitments(text)
        node.pop("generationCheckpoint", None)
        node["qualityAudit"] = evaluate_section_quality(
            text,
            node,
            source_context,
            requirements,
        )
        artifact_job_id = None
        with psycopg.connect(DB) as conn:
            conn.execute(
                """UPDATE outlines
                   SET content=jsonb_set(content,%s::text[],%s::jsonb,true),
                       updated_at=now()
                   WHERE project_id=%s""",
                (
                    outline_node_json_path(path),
                    json.dumps(node, ensure_ascii=False),
                    project_id,
                ),
            )
            conn.execute("UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s", (job_id,))
            latest_row = conn.execute(
                "SELECT content FROM outlines WHERE project_id=%s",
                (project_id,),
            ).fetchone()
            latest_content = (
                latest_row[0] if latest_row and isinstance(latest_row[0], dict)
                else json.loads(latest_row[0]) if latest_row else {}
            )
            leaf_nodes = artifact_leaf_sections(
                latest_content.get("chapters") or []
            )
            body_ready = bool(leaf_nodes) and all(
                str(item.get("contentStatus") or "") == "ready"
                and bool(str(item.get("content") or "").strip())
                for _, item in leaf_nodes
            )
            if body_ready:
                running_artifact = conn.execute(
                    """SELECT 1 FROM jobs
                       WHERE project_id=%s AND type='artifact'
                         AND status IN ('queued','running')
                       LIMIT 1""",
                    (project_id,),
                ).fetchone()
                if not running_artifact:
                    artifact_job_id = str(uuid.uuid4())
                    conn.execute(
                        "INSERT INTO jobs(id,project_id,type,status) VALUES(%s,%s,'artifact','queued')",
                        (artifact_job_id, project_id),
                    )
            conn.commit()
        refresh_project_commitment_registry(project_id)
        if artifact_job_id:
            R.lpush(
                "ai_bid:jobs",
                json.dumps(
                    {"jobId": artifact_job_id, "projectId": project_id, "type": "artifact"},
                    ensure_ascii=False,
                ),
            )
            print(
                json.dumps(
                    {
                        "event": "artifact_auto_queued",
                        "projectId": project_id,
                        "jobId": artifact_job_id,
                        "reason": "all_sections_ready",
                    },
                    ensure_ascii=False,
                ),
                flush=True,
            )
        print(json.dumps({"event": "section_generated", "projectId": project_id, "path": path, "mode": mode, "model": target_model, "strategy": model_mode, "requirements": len(requirements), "characters": len(text)}), flush=True)
    except Exception as exc:
        retry_count = int(job.get("retryCount") or 0)
        error_text = str(exc)
        transient = any(marker in error_text.lower() for marker in (
            "http error 500", "http error 502", "http error 503", "timed out",
            "timeout", "temporarily", "temporary", "connection reset",
            "remote end closed", "service unavailable", "bad gateway",
        ))
        if transient and retry_count < len(SECTION_RETRY_DELAYS):
            delay = SECTION_RETRY_DELAYS[retry_count]
            try:
                update_outline_node_fields(
                    project_id,
                    path,
                    {
                        "contentStatus": "retrying",
                        "retryCount": retry_count + 1,
                        "retryAfterSeconds": delay,
                    },
                )
                with psycopg.connect(DB) as conn:
                    conn.execute(
                        "UPDATE jobs SET status='retrying',error_message=%s WHERE id=%s",
                        (error_text[:500], job_id),
                    )
                    conn.commit()
                retry_job = dict(job)
                retry_job["retryCount"] = retry_count + 1
                R.zadd(DELAYED_QUEUE, {json.dumps(retry_job, ensure_ascii=False, sort_keys=True): time.time() + delay})
                print(json.dumps({"event": "section_retry_scheduled", "projectId": project_id, "path": path, "retry": retry_count + 1, "delaySeconds": delay, "error": error_text[:160]}, ensure_ascii=False), flush=True)
                return
            except Exception:
                traceback.print_exc()
        try:
            if outline_row:
                update_outline_node_fields(
                    project_id, path, {"contentStatus": "failed"}
                )
        finally:
            with psycopg.connect(DB) as conn:
                conn.execute("UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s", (str(exc)[:500], job_id)); conn.commit()
            traceback.print_exc()


def chapter_editor_role(index, count, title):
    if index == 0:
        return "建立项目事实、建设目标和投标响应立场，为本章后续内容提供统一起点"
    if index == count - 1:
        return "在前文范围和目标基础上收束关键问题，形成难点与措施一一对应的专业判断"
    if any(word in str(title or "") for word in ("范围", "需求", "边界")):
        return "明确服务对象、工作范围、实施边界和验收关注点，承接项目背景并引出重点难点"
    return "承接上一节结论并推进本章论证，不重复背景，不提前总结后续章节"


def formalize_editor_title(value):
    title = re.sub(r"\s+", " ", str(value or "")).strip()
    exact = {
        "管理重点难点对应楼宇卫生46分和校园环境卫生36.5分":
            "楼宇卫生与校园环境卫生重点难点统筹",
        "工作计划对应时间、内容、目标三项要素":
            "工作计划的时间、内容与目标统筹",
    }
    if title in exact:
        return exact[title]
    title = re.sub(
        r"^(.{2,18})覆盖.{4,100}的运行边界$",
        r"\1范围及责任分界",
        title,
    )
    title = re.sub(
        r"(\d+(?:\.\d+)?)\s*分(?:扣分)?风险控制",
        "质量风险控制",
        title,
    )
    title = re.sub(r"(\d+(?:\.\d+)?)\s*分扣分项", "质量问题", title)
    title = re.sub(r"(\d+(?:\.\d+)?)\s*分场景", "重点场景", title)
    title = re.sub(r"(\d+(?:\.\d+)?)\s*分", "", title)
    title = title.replace("扣分风险控制", "质量风险控制")
    title = title.replace("扣分项", "质量问题")
    title = title.replace("评分要求到本章内容的响应结构", "服务实施架构")
    title = title.replace("对应四类采购需求的", "四类物业服务")
    title = title.replace("管理重点难点专项策划", "重点难点分析与应对")
    title = title.replace("场景响应策划", "响应与闭环管理")
    title = re.sub(r"\s+", " ", title).strip(" ：:、")
    return title or str(value or "").strip()


def formalize_editor_chapter_titles(chapter):
    backup = []

    def visit(node, path):
        old_title = str(node.get("title") or "").strip()
        new_title = formalize_editor_title(old_title)
        if new_title != old_title:
            backup.append({
                "path": list(path),
                "title": old_title,
                "formalTitle": new_title,
            })
            node["title"] = new_title
        for index, child in enumerate(node.get("children") or []):
            visit(child, path + (index,))

    visit(chapter, ())
    return backup


def is_editor_controlled_commitment(node, content, claim):
    if is_controlled_bidder_commitment(content, claim):
        return True
    title = str((node or {}).get("title") or "")
    section_form = editor_section_form(node)["id"]
    claim_text = str(claim or "")
    if section_form in {
        "operating_procedure",
        "standard_checklist",
        "rectification_control",
    }:
        if re.search(
            r"(?:分钟|小时|秒|次|%|％|MPa|kPa|Pa|℃|毫升|ml|mL|克|g)\s*$",
            claim_text,
            re.I,
        ) or re.search(r"\d{1,2}[:：]\d{2}\s*$", claim_text):
            return True
        if re.search(r"(?:台|套|个)\s*$", claim_text):
            start = 0
            while True:
                index = content.find(claim_text, start)
                if index < 0:
                    break
                window = content[
                    max(0, index - 80): min(
                        len(content), index + len(claim_text) + 60
                    )
                ]
                if any(
                    marker in window
                    for marker in ("拟配备", "拟配置", "作业时使用", "配置")
                ):
                    return True
                start = index + len(claim_text)
    if not any(
        marker in title
        for marker in ("整改时限", "响应时限", "分级响应", "处置时限")
    ):
        return False
    if "我方内部控制标准" not in content:
        return False
    return bool(
        re.search(r"(?:分钟|小时|秒|次|%|％)\s*$", str(claim or ""))
    )


def editor_section_form(node, chapter_context=None):
    """Route each section to a bid-document form derived from its own semantics."""
    title = str((node or {}).get("title") or "")
    purpose = str((node or {}).get("description") or "")
    role = str((chapter_context or {}).get("sectionRole") or "")
    primary_text = f"{title} {purpose}"
    routes = (
        (
            "rectification_control",
            ("整改", "投诉", "闭环", "响应时限", "处置时限", "督办", "销项"),
            "编号式闭环机制",
            "按接收登记、分级响应、整改执行、复核销项、逾期升级组织；只有响应等级、责任主体、时限和输出物确需比较时使用一张简表。",
        ),
        (
            "standard_checklist",
            ("标准", "检查", "巡查", "验收", "判定", "考核", "质量要求"),
            "检查标准或判定清单",
            "先写适用对象，再按检查部位或质量特征列出可观察、可复核的判定要求；避免解释评分分值，确需并列比较时可使用检查表。",
        ),
        (
            "operating_procedure",
            ("流程", "作业", "清洁", "冲洗", "巡检", "实施", "操作", "维护"),
            "编号式作业规程",
            "从本节具体对象中提取4至7个常用作业事项作为小标题，例如浴室场景可写“清洁准备与安全提示、地面清杂与排水、污渍清洗与防滑处理、隔板墙面及管件清洁、地漏异味与设施报修、作业复查与记录”。不得使用“作业准备、操作步骤、安全控制、复核确认”，也不得生造“接管现场、作业启动节点、完成状态为”等流程设计术语。正文直接进入第一个事项，不写总起段。每个小标题下根据实际内容使用2至5条“1.1、1.2”式短规定；至少出现3种不同条目数，相同条目数不得覆盖多数标题，不要为了整齐把每节都填满5条。",
        ),
        (
            "scope_definition",
            ("范围", "边界", "对象", "区域划分", "职责分界", "交接"),
            "范围与职责界面说明",
            "按服务范围、区域划分、职责分界、交接机制组织。范围部分直接使用“包括、主要包括、以……为分界”等自然表述，不写“某岗位负责将……纳入基础范围”“统筹边界确认”“划定责任点位”等行政化句式；执行主体只在职责交接和异常处理处出现。不混入完整频次、应急和月报制度。",
        ),
        (
            "analysis_narrative",
            ("分析", "难点", "重点", "需求", "背景", "目标", "认识"),
            "项目分析与应对",
            "用项目事实引出专业判断，再按具体场景分别提出我方应对；不套用固定步骤，不用空泛价值判断。",
        ),
    )
    for semantic_text in (primary_text, role):
        for form_id, markers, label, instruction in routes:
            if any(marker in semantic_text for marker in markers):
                return {
                    "id": form_id,
                    "label": label,
                    "instruction": instruction,
                }
    return {
        "id": "professional_narrative",
        "label": "专业方案正文",
        "instruction": "根据本节实际任务选择自然段、编号步骤或少量表格，不强行套用统一结构；每一段只推进一个专业判断或实施动作。",
    }


def call_editor_model(
    node, current_text, requirements, project_name, source_context,
    chapter_context, project_id, job_id,
):
    section_form = editor_section_form(node, chapter_context)
    operating_procedure = section_form["id"] == "operating_procedure"
    detail_policy = {
        "mode": "scene_adaptive_bidder_commitment",
        "allowed": [
            "与本节场景直接相关的合理作业时段或巡查频次",
            "完成本节动作必需的通用工具或内部表单",
            "能够直接观察、检查或复核的完成标准",
        ],
        "requiredCategories": (
            [
                "至少1项明确作业时段或频次",
                "至少1项通用设备、工具或拟配置数量",
                "至少1项可执行参数、检查周期或质量阈值",
            ]
            if operating_procedure
            else (
                [
                    "至少1项与巡查对象直接相关的频次或触发条件",
                    "至少1项可观察的检查结果或异常处置条件",
                ]
                if section_form["id"] == "standard_checklist"
                else []
            )
        ),
        "instruction": (
            "根据sectionForm和sceneActionCards补充真正需要的专业细节。"
            "不得为了满足数量而硬塞清洁剂浓度、设备台数、压力、时间或阈值；"
            "只有作业动作确实需要时才写。"
        ),
    }
    system = """你是政府采购投标文件总编辑。你的任务不是回答问题，也不是重新生成一份咨询报告，而是把现有章节编辑成可直接装订的正式投标正文。
编辑规则：
1. 保留现有正文中有招标依据的项目对象、数量、时限、考核事项、工作范围、表格和实施动作，不得修改其数值或含义。
2. 删除“从招标文件看、可以看出、这说明、本项目并非……而是……、本节将、以下进行说明”等对话、解读和AI写作口吻。
3. 开篇直接陈述项目事实、我方实施安排或本节专业判断，不写摘要式套话。
4. 按chapterContext中的章节角色处理与相邻小节的关系：不重复上一节，不提前代写下一节，结尾自然过渡。
5. 严格围绕chapterContext.sectionRole写作。本节只承担一个明确任务；发现原稿混入其他章节内容时应删除或收束，不得为了保留篇幅而把范围、频次、应急、活动保障和报告管理混写在同一节。
6. 使用3至6个有业务含义的小标题；每个小标题推进一个不同问题。分析用自然段，连续动作使用编号步骤，只有真正可比较的重复字段才使用表格。
7. 每段先说明对象、判断或实施动作，再补充必要细节。一次连续列举原则上不超过6项，超过时按场景归类，不得形成招标对象名词串。
8. 使用正式、自然的投标文件语言，避免“通行面、公共面、外露面、可视硬化区域、巡回查看”等生造或含义模糊的搭配。优先使用“道路及附属区域、建筑外围、绿化景观区域、运动场地、巡查”等行业常用表达。
9. 不得逐项套用“目标、输入、步骤、责任、输出、质量、风险、记录”八段式模板。
10. 不得虚构企业资质证书、人员姓名和证书编号、既往业绩、客户案例、既有服务网点等需要证明材料支撑的企业事实。允许以“我方拟配备、作业时使用、我方作业标准”为口径新增合理的通用设备、工具、数量、作业时段、频次和内部控制参数。
11. 这是自然化编辑，不是强制压缩。删除重复和跑题内容，但保留使论述顺畅所需的连接、解释、实施逻辑和专业细节。
12. “本方案拟”原则上每节不超过1次；“我方内部控制标准”只在明确区分招标事实与自主承诺时使用，不得作为段落口头禅。
13. 输出Markdown正式正文，不要解释编辑过程，不要输出“修改稿、优化稿、编制说明、注意事项”等标签。
14. 将采购需求内化为我方履约安排。正文不得出现“招标要求明确、对应考核要求、评分中、扣分、最高扣分、从招标文件要求看”等面向读者解说评分规则的句子；确需保留的采购要求直接写入工作标准。
15. 开篇采用“具体项目场景或问题 + 我方行动”的投标人语气，优先使用“针对……，我方将……”“我方实行……”“项目经理/主管/作业人员负责……”；不得以“纳入……范围”“清洁质量对应考核要求”“学校考核中……”起笔。
16. 作业流程、整改处置等实施类段落应写明执行主体和动作，主体可为我方、项目经理、保洁主管、保洁人员、宿管岗位或维修岗位。范围界定、质量标准等说明类段落不强制反复出现岗位名称；在开篇或职责分界处交代责任即可，其余内容使用直接、客观的方案表述。
17. 每个编号项或自然段原则上为1至3句，只解决一个问题。超过6项的对象列举应按道路、楼宇、绿化、运动场等场景归类，不得把采购清单整段复制为长名词串。
18. 严格执行输入中的sectionForm。它是根据本节标题、用途和章节角色动态判定的写作体裁，并非跨行业固定模板；不得把其他体裁的栏目机械拼入本节。
19. 避免“负责将……纳入……基础范围、按……管理、统筹边界确认、组织现场落实、划定责任点位、编入巡查路线、不以……替代……”等为显得正式而制造的行政化搭配。能用“范围包括……”“日常保洁包括……”“交界处由双方衔接处理”说清楚时，不得改写成多层动宾结构。
20. 采用成熟企业作业制度的成稿口吻。除分析类章节外，不写大段背景、意义和设计理由；直接写作业项目、执行方法、控制要求和完成标准。
21. 小标题必须来自本项目的具体对象或动作。禁止连续使用“作业准备、操作步骤、安全控制、复核确认、实施原则、保障措施、持续改进”等可套用于任何行业的标题。
22. 采用“投标人自主承诺”模式补充专业细节。可根据行业常识给出合理的作业时间、每日或每周频次、通用设备和工具、设备数量、操作压力或浓度、检查周期、内部表单和岗位衔接；这些内容作为我方实施方案直接表达，不需要每条重复标注“内部控制标准”。不得虚构品牌、企业已有资产、人员证书、既往业绩或客户案例。
23. 一个编号项通常为1至2句，优先采用“动作名称＋具体要求”的短条款。除范围清单确有必要外，避免一个自然段超过180个汉字。
24. 作业规程类正文采用正式制度条款节奏：三级标题使用“1. 清洁准备与安全提示”形式，标题下使用“1.1、1.2、1.3”短条款。每条只写一个动作或标准，原则上不超过90个汉字。不同标题根据内容自然使用2至5条规定，至少形成3种不同条目数，相同条目数不得占全部标题的一半以上。
25. 作业规程类正文不要在标题后另写“针对……特点，我方实行……”的总起段，直接进入第一项具体规定；避免“接管现场、作业启动节点、完成状态为、现场岗位安排”等AI流程设计词，改用“开始清洁、作业前、清洁后、复查时”等日常专业表达。
26. requirementAtoms是本节招标原文的最小响应单元，sceneActionCards是由这些原文动态归纳的场景动作线索。应将其转化为项目化标题、作业动作和完成标准，不得输出字段名、原子编号或逐条问答；chapterContext.existingCommitments中的同类数值承诺应保持一致。"""
    user = json.dumps({
        "projectName": project_name,
        "sectionTitle": node.get("title"),
        "sectionPurpose": node.get("description"),
        "chapterContext": chapter_context,
        "sectionForm": section_form,
        "styleMode": "operational_manual",
        "professionalDetailPolicy": {
            **detail_policy,
            "examplesForMethodOnly": (
                [
                    "高峰结束后设置固定清洁时段",
                    "每日安排两次重点冲洗",
                    "拟配置高压水枪、吸水机或防滑垫",
                    "根据地材将冲洗压力控制在合理区间",
                ]
                if operating_procedure
                else [
                    "按使用高峰、日常开放和临时异常确定巡查触发条件",
                    "发现地面积水、设施异常或异味时记录位置并转入处置",
                    "以无明显积水、无散落杂物、异常已登记作为可复核结果",
                ]
            ),
        },
        "structureDiversity": {
            "headingCount": "根据内容使用4至7个",
            "clausesPerHeading": "每个标题2至5条",
            "hardRule": "至少出现3种不同条目数；任一相同条目数不得覆盖全部标题的一半以上；不得为了对称而补满条目。",
        },
        "requirements": requirements,
        "requirementAtoms": node.get("requirementAtoms") or [],
        "sceneActionCards": node.get("sceneActionCards") or [],
        "tenderExcerpts": source_context,
        "currentDraft": current_text,
        "lengthConstraint": {
            "currentCharacters": len(current_text),
            "minimumCharacters": int(
                len(current_text) * (0.82 if operating_procedure else 0.78)
            ),
            "targetCharacters": int(
                len(current_text) * (1.10 if operating_procedure else 0.88)
            ),
            "maximumCharacters": int(
                len(current_text) * (1.35 if operating_procedure else 1.05)
            ),
            "instruction": (
                "作业规程需要在保留有效内容的基础上新增明确时段、频次、设备工具、操作参数和检查标准，允许适度扩写，不能仅压缩或同义改写旧稿。"
                if operating_procedure else
                "这是自然化编辑任务，不按固定比例删减。删除重复和跑题内容，保留专业展开、必要连接、实施逻辑、检查点和有效表格，篇幅尽量接近 targetCharacters。"
            ),
        },
    }, ensure_ascii=False)
    max_tokens = max(4200, min(14000, int(len(current_text) * 1.25)))
    body = make_chat_body(
        [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        max_tokens,
        0.28 if operating_procedure else 0.15,
    )
    started = time.monotonic()
    run_id = str(uuid.uuid4())
    input_tokens = output_tokens = 0
    actual_cost = None
    try:
        request = urllib.request.Request(
            AI_URL,
            data=body,
            headers={
                "Authorization": f"Bearer {AI_KEY}",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        with urllib.request.urlopen(request, timeout=360) as response:
            client_request_id = response.headers.get("X-Client-Request-Id")
            payload = json.loads(response.read().decode("utf-8"))
        usage = payload.get("usage") or {}
        input_tokens = int(usage.get("prompt_tokens") or 0)
        output_tokens = int(usage.get("completion_tokens") or 0)
        billing = fetch_billing(client_request_id)
        if billing:
            input_tokens, output_tokens, actual_cost = billing
        choice = payload["choices"][0]
        if choice.get("finish_reason") == "length":
            raise ValueError("总编输出达到长度上限，章节未完整结束")
        edited = str(choice["message"]["content"] or "").strip()
        edited, _ = clean_model_text_artifacts(edited)
        edited, _ = normalize_formal_bid_voice(edited)
        minimum_edited_length = (
            max(1000, int(len(current_text) * 0.48))
            if section_form["id"] in {
                "operating_procedure",
                "standard_checklist",
                "rectification_control",
            }
            else max(800, int(len(current_text) * 0.70))
        )
        if len(edited) < minimum_edited_length:
            raise ValueError("总编结果删减过度，未保留原章节的主要内容")
        source = (
            current_text + "\n" + source_context + "\n"
            + "\n".join(item.get("requirement", "") for item in requirements)
        )
        if operating_procedure:
            clause_counts = {}
            for major, _minor in re.findall(
                r"(?m)^(\d+)\.(\d+)\s+",
                edited,
            ):
                clause_counts[major] = clause_counts.get(major, 0) + 1
            clause_values = list(clause_counts.values())
            dominant_count = max(
                (clause_values.count(value) for value in set(clause_values)),
                default=0,
            )
            if (
                len(clause_counts) >= 4
                and (
                    len(set(clause_values)) < 3
                    or dominant_count > len(clause_values) / 2
                )
            ):
                raise ValueError(
                    "作业规程各事项条目数过于整齐，仍有机械模板痕迹"
                )
            detail_categories = {
                "time_or_frequency": bool(re.search(
                    r"(?:\d{1,2}[:：]\d{2}|每日|每天|每周|每班|每次|"
                    r"\d+\s*(?:次|分钟|小时))",
                    edited,
                )),
                "equipment": bool(re.search(
                    r"(?:高压水枪|吸水机|刮水器|防滑垫|洗地机|"
                    r"喷壶|刷具|工具车|警示牌)",
                    edited,
                )),
                "parameter_or_threshold": bool(re.search(
                    r"(?:MPa|kPa|Pa|℃|浓度|压力|不超过|不少于|"
                    r"\d+\s*(?:分钟|小时|%|％))",
                    edited,
                    re.I,
                )),
            }
            if sum(detail_categories.values()) < 3:
                raise ValueError(
                    "作业规程缺少时段频次、设备工具或量化参数等专业细节"
                )
        unsupported_numbers = [
            value for value in re.findall(
                r"(?:\d{1,2}[:：]\d{2}|"
                r"\d+(?:\.\d+)?(?:\s*[-—~至]\s*\d+(?:\.\d+)?)?\s*"
                r"(?:%|％|个?工作日|天|小时|分钟|秒|次|人|套|台|家|年|个月|月|分|MPa|kPa|Pa|℃|毫升|ml|mL|克|g))",
                edited,
                flags=re.I,
            )
            if (
                norm(value) not in norm(source)
                and not is_editor_controlled_commitment(
                    node, edited, value
                )
            )
        ]
        if unsupported_numbers:
            raise ValueError(
                "总编结果出现无依据量化内容："
                + "、".join(list(dict.fromkeys(unsupported_numbers))[:6])
            )
        save_run(
            run_id, project_id, job_id, 1, "succeeded",
            input_tokens, output_tokens,
            int((time.monotonic() - started) * 1000),
            0, actual_cost=actual_cost, run_type="chapter_editor",
        )
        return edited
    except Exception as exc:
        save_run(
            run_id, project_id, job_id, 1, "failed",
            input_tokens, output_tokens,
            int((time.monotonic() - started) * 1000),
            0, exc, actual_cost, "chapter_editor",
        )
        raise


def process_chapter_editor(job):
    global MODEL, AI_URL, AI_KEY, AI_THINKING
    job_id, project_id = job["jobId"], job["projectId"]
    chapter_index = int(job.get("chapterIndex") or 0)
    with psycopg.connect(DB) as conn:
        conn.execute(
            "UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s",
            (job_id,),
        )
        outline_row = conn.execute(
            "SELECT content FROM outlines WHERE project_id=%s",
            (project_id,),
        ).fetchone()
        project_row = conn.execute(
            "SELECT name FROM projects WHERE id=%s", (project_id,)
        ).fetchone()
        requirement_rows = conn.execute(
            """SELECT id,type,title,normalized_value,mandatory
               FROM requirements WHERE project_id=%s""",
            (project_id,),
        ).fetchall()
        document_row = conn.execute(
            "SELECT blocks FROM documents WHERE project_id=%s",
            (project_id,),
        ).fetchone()
        conn.commit()
    original_model, original_url = MODEL, AI_URL
    original_key, original_thinking = AI_KEY, AI_THINKING
    current_path = [chapter_index]
    try:
        if not outline_row:
            raise ValueError("项目大纲不存在")
        content = (
            outline_row[0] if isinstance(outline_row[0], dict)
            else json.loads(outline_row[0])
        )
        chapters = content.get("chapters") or []
        if chapter_index < 0 or chapter_index >= len(chapters):
            raise ValueError("总编目标章节不存在")
        chapter = chapters[chapter_index]
        title_backup = formalize_editor_chapter_titles(chapter)
        if title_backup:
            chapter["editorialTitleBackup"] = {
                "version": "formal-title-v1",
                "savedAt": time.strftime(
                    "%Y-%m-%dT%H:%M:%SZ", time.gmtime()
                ),
                "items": title_backup,
            }
            with psycopg.connect(DB) as conn:
                conn.execute(
                    "UPDATE outlines SET content=%s::jsonb,updated_at=now() WHERE project_id=%s",
                    (json.dumps(content, ensure_ascii=False), project_id),
                )
                conn.commit()
        leaves = outline_leaf_entries(
            chapter.get("children") or [], (chapter_index,)
        )
        if not leaves:
            leaves = [{"path": (chapter_index,), "node": chapter}]
        requirement_map = {
            str(row[0]): {
                "id": str(row[0]), "type": row[1], "title": row[2],
                "requirement": row[3], "mandatory": bool(row[4]),
            }
            for row in requirement_rows
        }
        blocks = (
            document_row[0]
            if document_row and isinstance(document_row[0], list)
            else json.loads(document_row[0]) if document_row else []
        )
        blueprint = {
            "projectProfile": content.get("projectProfile") or {},
            "implementationBlueprint": content.get("implementationBlueprint") or {},
            "projectAnalysis": content.get("projectAnalysis") or {},
            "scoringTasks": content.get("scoringTasks") or [],
            "capabilityPlan": content.get("capabilityPlan") or [],
        }
        update_outline_node_fields(
            project_id, [chapter_index],
            {"editorStatus": "editing", "editorProgress": 0},
        )
        MODEL = GPT_TEXT_MODEL
        if not IMAGE_KEY:
            raise ValueError("sub2api GPT 密钥未配置")
        AI_URL = IMAGE_BASE_URL.rstrip("/") + "/v1/chat/completions"
        AI_KEY = IMAGE_KEY
        AI_THINKING = "disabled"
        titles = [entry["node"].get("title") for entry in leaves]
        editorial_errors = []
        for index, entry in enumerate(leaves):
            current_path = list(entry["path"])
            node = entry["node"]
            editorial_backup = node.get("previousEditorial") or {}
            current_text = str(
                editorial_backup.get("content") or node.get("content") or ""
            ).strip()
            if not current_text:
                continue
            requirements, rejected_links = editor_relevant_requirements(
                node, requirement_map
            )
            node["requirementLinkAudit"] = {
                "version": "semantic-zero-reject-v1",
                "acceptedIds": [item["id"] for item in requirements],
                "rejected": rejected_links,
            }
            source_context = section_source_context(node, blocks, blueprint)
            requirement_atoms = build_requirement_atoms(source_context, node)
            scene_action_cards = build_scene_action_cards(
                node, requirement_atoms
            )
            node["requirementAtoms"] = requirement_atoms
            node["sceneActionCards"] = scene_action_cards
            chapter_context = {
                "chapterTitle": chapter.get("title"),
                "chapterPurpose": chapter.get("description"),
                "sectionIndex": index + 1,
                "sectionCount": len(leaves),
                "sectionRole": chapter_editor_role(
                    index, len(leaves), node.get("title")
                ),
                "previousSectionTitle": titles[index - 1] if index else None,
                "nextSectionTitle": (
                    titles[index + 1] if index + 1 < len(titles) else None
                ),
                "existingCommitments": relevant_outline_commitments(
                    content, node, current_path
                ),
                "structuredInputRule": (
                    "逐项消化requirementAtoms并按sceneActionCards展开；"
                    "同类数值承诺与existingCommitments保持一致。"
                ),
            }
            try:
                edited = call_editor_model(
                    node, current_text, requirements,
                    project_row[0] if project_row else "",
                    source_context, chapter_context, project_id, job_id,
                )
            except Exception as exc:
                editorial_errors.append({
                    "path": current_path,
                    "title": node.get("title"),
                    "error": str(exc)[:180],
                })
                node["editorialStatus"] = "failed"
                node["editorialError"] = str(exc)[:300]
                persist_outline_node(project_id, current_path, node)
                update_outline_node_fields(
                    project_id, [chapter_index],
                    {
                        "editorStatus": "editing",
                        "editorProgress": round(
                            (index + 1) / len(leaves) * 100
                        ),
                    },
                )
                continue
            node["previousEditorial"] = {
                "content": current_text,
                "qualityAudit": node.get("qualityAudit"),
                "savedAt": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            }
            node["content"] = edited
            node["editorialStatus"] = "ready"
            node["editorialModel"] = GPT_TEXT_MODEL
            node["editorialUpdatedAt"] = time.strftime(
                "%Y-%m-%dT%H:%M:%SZ", time.gmtime()
            )
            node["commitments"] = extract_section_commitments(edited)
            node["qualityAudit"] = evaluate_section_quality(
                edited, node, source_context, requirements
            )
            persist_outline_node(project_id, current_path, node)
            update_outline_node_fields(
                project_id, [chapter_index],
                {
                    "editorStatus": "editing",
                    "editorProgress": round((index + 1) / len(leaves) * 100),
                },
            )
            refresh_project_commitment_registry(project_id)
        update_outline_node_fields(
            project_id, [chapter_index],
            {
                "editorStatus": "ready",
                "editorProgress": 100,
                "editorError": (
                    f"{len(editorial_errors)}节保留原稿，待定向复核"
                    if editorial_errors else None
                ),
                "editorialErrorDetails": editorial_errors[:30],
            },
        )
        with psycopg.connect(DB) as conn:
            conn.execute(
                "UPDATE jobs SET status='succeeded',finished_at=now(),error_message=%s WHERE id=%s",
                (
                    (
                        f"{len(editorial_errors)}节保留原稿，待定向复核"
                        if editorial_errors else None
                    ),
                    job_id,
                ),
            )
            conn.commit()
        print(json.dumps({
            "event": "chapter_edited",
            "projectId": project_id,
            "chapterIndex": chapter_index,
            "sections": len(leaves),
            "preservedOriginals": len(editorial_errors),
            "model": GPT_TEXT_MODEL,
        }, ensure_ascii=False), flush=True)
    except Exception as exc:
        update_outline_node_fields(
            project_id, [chapter_index],
            {"editorStatus": "failed", "editorError": str(exc)[:300]},
        )
        with psycopg.connect(DB) as conn:
            conn.execute(
                "UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s",
                (str(exc)[:500], job_id),
            )
            conn.commit()
        traceback.print_exc()
    finally:
        MODEL, AI_URL = original_model, original_url
        AI_KEY, AI_THINKING = original_key, original_thinking


def process_section_compare(job):
    global MODEL, AI_URL, AI_KEY, AI_THINKING
    job_id, project_id, path = job["jobId"], job["projectId"], job.get("path") or []
    target_model = job.get("model") or "gpt-5.5"
    original_model, original_url, original_key, original_thinking = MODEL, AI_URL, AI_KEY, AI_THINKING
    outline_row = None
    try:
        with psycopg.connect(DB) as conn:
            conn.execute(
                "UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s",
                (job_id,),
            )
            outline_row = conn.execute(
                "SELECT content FROM outlines WHERE project_id=%s",
                (project_id,),
            ).fetchone()
            requirement_rows = conn.execute(
                "SELECT id,type,title,normalized_value FROM requirements WHERE project_id=%s",
                (project_id,),
            ).fetchall()
            project_row = conn.execute(
                "SELECT name FROM projects WHERE id=%s",
                (project_id,),
            ).fetchone()
            document_row = conn.execute(
                "SELECT blocks FROM documents WHERE project_id=%s",
                (project_id,),
            ).fetchone()
            conn.commit()
        if not outline_row:
            raise ValueError("项目大纲不存在")
        content = outline_row[0] if isinstance(outline_row[0], dict) else json.loads(outline_row[0])
        node = outline_node_at(content, path)
        if not node:
            raise ValueError("目标章节不存在")
        requirement_map = {
            str(row[0]): {
                "id": str(row[0]),
                "type": row[1],
                "title": row[2],
                "requirement": row[3],
            }
            for row in requirement_rows
        }
        requirements = [
            requirement_map[value]
            for value in collect_node_requirement_ids(node)
            if value in requirement_map
        ]
        blocks = (
            document_row[0]
            if document_row and isinstance(document_row[0], list)
            else (json.loads(document_row[0]) if document_row else [])
        )
        blueprint = {
            "projectProfile": content.get("projectProfile"),
            "implementationBlueprint": content.get("implementationBlueprint"),
            "projectAnalysis": content.get("projectAnalysis"),
            "scoringTasks": content.get("scoringTasks") or [],
            "capabilityPlan": content.get("capabilityPlan") or [],
        }
        source_context = section_source_context(node, blocks, blueprint)
        section_brief = node.get("brief")
        if not isinstance(section_brief, dict):
            section_brief = build_section_brief(node, requirements, blueprint)
        MODEL = target_model
        if target_model.startswith("gpt-"):
            if not IMAGE_KEY:
                raise ValueError("sub2api GPT 密钥未配置")
            AI_URL = IMAGE_BASE_URL.rstrip("/") + "/v1/chat/completions"
            AI_KEY = IMAGE_KEY
        else:
            AI_URL = original_url
            AI_KEY = original_key
        AI_THINKING = "disabled"
        text, mode = call_section_model(
            node,
            requirements,
            project_row[0] if project_row else "",
            source_context,
            project_id,
            job_id,
            blueprint,
            section_brief,
            "deep",
        )
        with psycopg.connect(DB) as conn:
            run = conn.execute(
                """SELECT input_tokens,output_tokens,cost_usd,duration_ms,model
                   FROM ai_runs WHERE job_id=%s AND status='succeeded'
                   ORDER BY created_at DESC LIMIT 1""",
                (job_id,),
            ).fetchone()
            latest = conn.execute(
                "SELECT content FROM outlines WHERE project_id=%s",
                (project_id,),
            ).fetchone()
            latest_content = latest[0] if latest and isinstance(latest[0], dict) else json.loads(latest[0])
            latest_node = outline_node_at(latest_content, path)
            variants = latest_node.setdefault("comparisonVariants", {})
            variants[target_model] = {
                "status": "ready",
                "model": run[4] if run else target_model,
                "content": text,
                "contentMode": mode,
                "inputTokens": int(run[0]) if run else 0,
                "outputTokens": int(run[1]) if run else 0,
                "costUsd": float(run[2]) if run else 0,
                "durationMs": int(run[3]) if run else 0,
                "generatedAt": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "qualityAudit": evaluate_section_quality(
                    text,
                    latest_node,
                    source_context,
                    requirements,
                ),
            }
            conn.execute(
                "UPDATE outlines SET content=%s::jsonb,updated_at=now() WHERE project_id=%s",
                (json.dumps(latest_content, ensure_ascii=False), project_id),
            )
            conn.execute(
                "UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s",
                (job_id,),
            )
            conn.commit()
        print(json.dumps({
            "event": "section_compared",
            "projectId": project_id,
            "path": path,
            "model": target_model,
            "characters": len(text),
            "costUsd": float(run[2]) if run else None,
        }, ensure_ascii=False), flush=True)
    except Exception as exc:
        try:
            with psycopg.connect(DB) as conn:
                latest = conn.execute(
                    "SELECT content FROM outlines WHERE project_id=%s",
                    (project_id,),
                ).fetchone()
                if latest:
                    latest_content = latest[0] if isinstance(latest[0], dict) else json.loads(latest[0])
                    latest_node = outline_node_at(latest_content, path)
                    if latest_node:
                        variants = latest_node.setdefault("comparisonVariants", {})
                        variants[target_model] = {
                            "status": "failed",
                            "model": target_model,
                            "errorMessage": str(exc)[:500],
                        }
                        conn.execute(
                            "UPDATE outlines SET content=%s::jsonb,updated_at=now() WHERE project_id=%s",
                            (json.dumps(latest_content, ensure_ascii=False), project_id),
                        )
                conn.execute(
                    "UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s",
                    (str(exc)[:500], job_id),
                )
                conn.commit()
        finally:
            traceback.print_exc()
    finally:
        MODEL, AI_URL, AI_KEY, AI_THINKING = original_model, original_url, original_key, original_thinking


def artifact_text(value, limit=34):
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text if len(text) <= limit else text[: limit - 1] + "…"


def svg_text_lines(text, x, y, width=12, size=22, color="#17324d", weight=600):
    value = artifact_text(text, width * 3)
    lines = [value[i:i + width] for i in range(0, len(value), width)] or [""]
    start = y - ((len(lines) - 1) * size * 0.58)
    spans = "".join(
        f'<tspan x="{x}" y="{start + index * size * 1.18:.0f}">{html.escape(line)}</tspan>'
        for index, line in enumerate(lines[:3])
    )
    return (
        f'<text text-anchor="middle" font-family="Microsoft YaHei,Noto Sans CJK SC,sans-serif" '
        f'font-size="{size}" font-weight="{weight}" fill="{color}">{spans}</text>'
    )


def svg_document(title, subtitle, body):
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="1200" height="680" viewBox="0 0 1200 680">
<defs>
  <marker id="arrow" markerWidth="10" markerHeight="10" refX="8" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#5b7890"/></marker>
</defs>
<rect x="1" y="1" width="1198" height="678" rx="26" fill="#f7fafc" stroke="#d7e3eb" stroke-width="2"/>
<rect x="0" y="0" width="12" height="680" fill="#0d9488"/>
{svg_text_lines(title,600,58,28,30,"#12344d",700)}
{svg_text_lines(subtitle,600,96,48,16,"#60798c",400)}
{body}
<text x="1148" y="648" text-anchor="end" font-family="Microsoft YaHei,Noto Sans CJK SC,sans-serif" font-size="14" fill="#7890a2">由项目事实与实施方法自动编制</text>
</svg>"""


def build_flow_svg(title, subtitle, labels):
    labels = [artifact_text(item, 24) for item in labels if str(item or "").strip()][:6]
    if not labels:
        labels = ["需求确认", "方案设计", "组织实施", "检查验收"]
    count = len(labels)
    box_width = min(166, max(132, int((1040 - 34 * (count - 1)) / count)))
    gap = 34
    total = count * box_width + (count - 1) * gap
    start_x = (1200 - total) / 2
    parts = []
    for index, label in enumerate(labels):
        x = start_x + index * (box_width + gap)
        if index:
            parts.append(
                f'<line x1="{x-gap+5:.0f}" y1="337" x2="{x-7:.0f}" y2="337" stroke="#5b7890" '
                'stroke-width="3" marker-end="url(#arrow)"/>'
            )
        parts.append(
            f'<rect x="{x:.0f}" y="258" width="{box_width}" height="158" rx="18" '
            'fill="#ffffff" stroke="#b8d4df" stroke-width="2"/>'
        )
        parts.append(
            f'<circle cx="{x+box_width/2:.0f}" cy="282" r="20" fill="#176b87"/>'
            f'<text x="{x+box_width/2:.0f}" y="290" text-anchor="middle" font-family="Arial" '
            f'font-size="20" font-weight="700" fill="#fff">{index+1}</text>'
        )
        parts.append(svg_text_lines(label, x + box_width / 2, 354, 8, 20))
    parts.append(
        '<path d="M1010 455 C1010 540,190 540,190 455" fill="none" stroke="#0d9488" '
        'stroke-width="3" stroke-dasharray="9 7" marker-end="url(#arrow)"/>'
    )
    parts.append(svg_text_lines("反馈、纠偏与持续优化", 600, 566, 20, 18, "#0d766e", 600))
    return svg_document(title, subtitle, "".join(parts))


def build_architecture_svg(project_name, workstreams):
    labels = [artifact_text(item, 28) for item in workstreams if str(item or "").strip()][:6]
    if not labels:
        labels = ["需求与范围", "技术方案", "组织实施", "质量验收"]
    parts = [
        '<rect x="390" y="132" width="420" height="104" rx="24" fill="#176b87"/>',
        svg_text_lines(project_name, 600, 186, 20, 24, "#ffffff", 700),
    ]
    cols = 3
    box_width, box_height = 280, 112
    positions = []
    for index, label in enumerate(labels):
        row, col = divmod(index, cols)
        x = 100 + col * 360
        y = 330 + row * 168
        positions.append((x, y))
        parts.append(
            f'<path d="M600 236 C600 282,{x+box_width/2:.0f} 282,{x+box_width/2:.0f} {y-10}" '
            'fill="none" stroke="#66869a" stroke-width="2.5" marker-end="url(#arrow)"/>'
        )
        parts.append(
            f'<rect x="{x}" y="{y}" width="{box_width}" height="{box_height}" rx="18" '
            'fill="#fff" stroke="#9fc8d5" stroke-width="2"/>'
        )
        parts.append(svg_text_lines(label, x + box_width / 2, y + 58, 12, 21))
    return svg_document("项目总体交付架构图", "以项目目标为中心，将专业工作流统一纳入交付主线", "".join(parts))


def artifact_leaf_sections(nodes, prefix=()):
    output = []
    for index, node in enumerate(nodes or []):
        path = prefix + (index,)
        children = node.get("children") or []
        if children:
            output.extend(artifact_leaf_sections(children, path))
        else:
            output.append((path, node))
    return output


def artifact_target(outline, capability_id, title_phrases):
    ranked = []
    for order, (path, node) in enumerate(artifact_leaf_sections(outline.get("chapters") or [])):
        title = str(node.get("title") or "")
        brief = node.get("brief") or {}
        module_ids = {
            str(item.get("id") or "")
            for item in (brief.get("capabilityModules") or [])
            if isinstance(item, dict)
        }
        score = 180 if capability_id and capability_id in module_ids else 0
        for phrase_order, phrase in enumerate(title_phrases or []):
            if phrase and phrase in title:
                score += 240 - phrase_order * 20
        ranked.append((score, -order, path, title))
    if not ranked:
        return None
    _, _, path, title = max(ranked)
    return {"targetPath": list(path), "targetTitle": title, "placementMode": "planned"}


def planned_artifact(spec, outline, capability_id, title_phrases):
    target = artifact_target(outline, capability_id, title_phrases)
    if target:
        spec.update(target)
    spec["plannerVersion"] = "capability-artifact-planner-v2"
    return spec


def artifact_specs(outline):
    blueprint = outline.get("implementationBlueprint") or {}
    phases = blueprint.get("phases") or []
    phase_names = [item.get("name") for item in phases if isinstance(item, dict) and item.get("name")]
    analysis = outline.get("projectAnalysis") or {}
    workstreams = analysis.get("workstreams") or []
    stream_names = [
        item.get("name") if isinstance(item, dict) else str(item)
        for item in workstreams
    ]
    profile = outline.get("projectProfile") or {}
    if not stream_names:
        stream_names = profile.get("scope") or profile.get("deliverables") or []
    project_name = outline.get("projectName") or profile.get("projectType") or "项目总体目标"
    capability_ids = {
        str(item.get("id") or "")
        for item in (outline.get("capabilityPlan") or [])
        if isinstance(item, dict)
    }
    specs = [
        {
            "kind": "overall_architecture",
            "title": "项目总体交付架构图",
            "source": "projectAnalysis.workstreams",
            "svg": build_architecture_svg(project_name, stream_names),
        },
        {
            "kind": "implementation_route",
            "title": "项目实施路线图",
            "source": "implementationBlueprint.phases",
            "svg": build_flow_svg(
                "项目实施路线图",
                "阶段任务、质量门与交付成果沿同一主线推进",
                phase_names,
            ),
        },
        {
            "kind": "quality_closed_loop",
            "title": "全过程质量控制闭环图",
            "source": "cross-industry-quality-method",
            "svg": build_flow_svg(
                "全过程质量控制闭环图",
                "适用于货物、服务、软件、集成及施工类项目",
                ["需求基线", "方案审查", "过程检查", "成果复核", "整改验证", "验收归档"],
            ),
        },
    ]
    core_targets = (
        ("overall_solution", ("总体思路", "总体方案", "建设目标")),
        ("schedule_control", ("进度计划", "里程碑", "实施路径")),
        ("quality_control", ("质量目标", "质量管理", "质量控制")),
    )
    specs = [
        planned_artifact(spec, outline, capability_id, phrases)
        for spec, (capability_id, phrases) in zip(specs, core_targets)
    ]
    templates = [
        (
            110, "site_implementation", "site_implementation_flow", "现场实施与工序衔接图",
            "从进场条件确认到工序验收组织现场实施",
            ["条件确认", "资源进场", "技术交底", "工序实施", "过程检查", "成品保护"],
            ("现场实施", "施工组织", "工序", "安装实施"),
        ),
        (
            108, "system_integration", "system_integration_flow", "系统集成与联调验证图",
            "按接口、单元、分系统和端到端逐级完成联调验证",
            ["接口确认", "单元测试", "分系统联调", "端到端验证", "问题整改", "试运行确认"],
            ("系统集成", "联调", "调试", "测试"),
        ),
        (
            108, "software_delivery", "software_delivery_lifecycle", "软件交付与版本发布流程图",
            "以需求基线和版本控制贯穿设计、开发、测试及上线",
            ["需求基线", "设计评审", "开发配置", "测试整改", "部署上线", "版本归档"],
            ("软件", "开发", "部署", "测试"),
        ),
        (
            106, "service_operation", "service_operation_cycle", "日常服务作业闭环图",
            "按计划、执行、检查、处置和复核组织日常服务",
            ["岗前准备", "任务分派", "现场作业", "巡查检查", "问题处置", "记录复核"],
            ("工作方法", "工作流程", "服务实施", "日常服务"),
        ),
        (
            104, "procurement_supply", "procurement_supply_flow", "采购与供货控制流程图",
            "从选型确认到到货交接控制供应履约",
            ["选型确认", "订单下达", "生产跟踪", "出厂检查", "运输交接", "到货验收"],
            ("采购", "供货", "设备材料", "供应"),
        ),
        (
            102, "packaging_logistics", "logistics_delivery_flow", "包装运输与到货交接流程图",
            "按货物特性控制包装、运输、交接和异常处置",
            ["包装防护", "装车复核", "在途跟踪", "现场交接", "开箱清点", "货损处置"],
            ("包装运输", "物流", "到货", "交付"),
        ),
        (
            100, "risk_emergency", "emergency_response", "突发事件应急响应流程图",
            "按事件识别、分级响应、协同处置和恢复复盘推进",
            ["事件发现", "分级研判", "现场处置", "协同上报", "恢复验证", "复盘改进"],
            ("风险识别", "应急预案", "风险", "应急"),
        ),
        (
            98, "organization_resource", "organization_responsibility", "项目组织与岗位协同图",
            "以统一指挥、专业执行、监督复核和协同保障形成责任链",
            ["项目统筹", "任务分派", "专业执行", "过程协同", "质量监督", "结果反馈"],
            ("项目组织", "组织架构", "岗位职责", "人员安排"),
        ),
        (
            96, "requirements_analysis", "service_scope_map", "项目范围与任务分解图",
            "将采购要求转化为可执行、可检查、可留痕的项目任务",
            ["对象识别", "边界确认", "要求分解", "任务配置", "证据留存", "结果核验"],
            ("服务范围", "需求理解", "范围界定", "项目背景"),
        ),
        (
            94, "safety_confidentiality", "safety_control_system", "安全保密分层控制图",
            "从人员、场所、设备、网络、数据和制度实施分层防护",
            ["人员控制", "场所控制", "设备控制", "网络控制", "数据控制", "制度检查"],
            ("安全管理", "保密", "数据管理", "安全措施"),
        ),
        (
            92, "acceptance_handover", "acceptance_handover_flow", "验收与成果移交流程图",
            "以验收依据和成果证据推动整改复验及正式移交",
            ["验收准备", "资料核对", "现场验收", "问题整改", "复验确认", "成果移交"],
            ("验收", "资料移交", "成果移交", "交付"),
        ),
        (
            90, "aftersales_maintenance", "service_response_flow", "售后服务响应闭环图",
            "按受理、派单、处理、验证和复盘管理服务工单",
            ["服务受理", "问题分级", "任务派单", "处理反馈", "验证关闭", "复盘改进"],
            ("响应与售后", "售后服务", "运行维护", "运维"),
        ),
        (
            88, "training_transfer", "training_transfer_flow", "培训与知识转移路径图",
            "按培训对象和使用场景组织课程、实操、考核及补训",
            ["对象分析", "课程设计", "理论培训", "实操演练", "考核评价", "补训归档"],
            ("培训与知识转移", "培训", "知识转移"),
        ),
        (
            84, "document_delivery", "document_evidence_flow", "文档与过程证据链图",
            "统一编号、版本、审批、归档和交付形成可追溯证据",
            ["记录形成", "编号登记", "版本审核", "审批确认", "分类归档", "成果交付"],
            ("文档", "资料", "过程证据", "资料移交"),
        ),
    ]
    leaf_count = len(artifact_leaf_sections(outline.get("chapters") or []))
    artifact_budget = max(5, min(12, round(leaf_count / 3))) if leaf_count else 5
    for _, capability_id, kind, title, subtitle, labels, phrases in templates:
        if len(specs) >= artifact_budget:
            break
        if capability_id not in capability_ids:
            continue
        specs.append(
            planned_artifact(
                {
                    "kind": kind,
                    "title": title,
                    "source": f"capabilityPlan.{capability_id}",
                    "svg": build_flow_svg(title, subtitle, labels),
                },
                outline,
                capability_id,
                phrases,
            )
        )
    return specs


def outline_visual_context(outline):
    values = []
    analysis = outline.get("projectAnalysis") or {}
    for item in (analysis.get("domainSignals") or []) + (analysis.get("archetypeComponents") or []):
        if str(item or "").strip():
            values.append(str(item))
    for _, node in artifact_leaf_sections(outline.get("chapters") or []):
        brief = node.get("brief") or {}
        for item in brief.get("requirementTitles") or []:
            text = artifact_text(item, 44)
            if text and text not in values:
                values.append(text)
            if len(values) >= 14:
                break
        if len(values) >= 14:
            break
    return "；".join(values[:14])


def realistic_scene_specs(outline, visual_mode):
    if visual_mode not in {"mixed", "physical_priority"}:
        return []
    archetype = str((outline.get("projectAnalysis") or {}).get("deliveryArchetype") or "mixed")
    scene_library = {
        "operation_scene": (
            "项目现场作业写实示意图", "service_operation",
            ("工作方法", "服务实施", "服务范围", "日常服务"),
            "写实摄影风格，展示项目人员在真实现场按规范开展专业作业，体现人员、工具、环境和协同过程；不得出现品牌、单位标志和无法确认的数量。",
        ),
        "equipment_scene": (
            "项目设备物资写实示意图", "organization_resource",
            ("资源配置", "物资装备", "设备材料", "服务范围"),
            "写实产品摄影风格，展示与本项目交付形态相符的通用设备、工具和物资整齐配置于真实工作场地；不得出现品牌型号、夸张数量和文字标签。",
        ),
        "installation_scene": (
            "设备安装实施写实示意图", "site_implementation",
            ("安装实施", "现场实施", "施工组织", "实施路径"),
            "写实工程摄影风格，展示技术人员在规范防护条件下进行设备安装、接线或调试，环境整洁、动作专业；不得出现品牌型号、单位标志和虚构参数。",
        ),
        "software_scene": (
            "软件系统应用写实示意图", "software_delivery",
            ("软件", "系统应用", "部署", "测试"),
            "写实办公场景摄影风格，展示工作人员在电脑终端进行系统操作、测试和协同；屏幕内容仅为不可辨识的界面色块，不生成品牌、文字和虚构功能。",
        ),
        "site_scene": (
            "工程现场实施写实示意图", "site_implementation",
            ("现场实施", "施工组织", "工序", "安全管理"),
            "写实工程现场摄影风格，展示规范围护、个人防护、材料堆放和有序施工状态；不得出现具体单位标志、虚构工程量和未确认工艺参数。",
        ),
        "completed_scene": (
            "项目完成效果写实示意图", "acceptance_handover",
            ("建设目标", "验收", "完成效果", "成果移交"),
            "写实摄影风格，展示项目完成后的整洁、规范、可使用状态，体现交付成果和使用场景；不得出现品牌、单位标志、文字标牌和无法确认的设施型号。",
        ),
        "project_scene": (
            "项目综合场景写实示意图", "overall_solution",
            ("总体思路", "总体方案", "建设目标"),
            "写实摄影风格，展示与项目交付形态相符的综合现场场景，兼顾人员、环境、设备和成果；不得出现品牌、单位标志、文字和虚构型号。",
        ),
    }
    archetype_scenes = {
        "goods": ["equipment_scene", "completed_scene", "project_scene"],
        "equipment_integration": ["equipment_scene", "installation_scene", "completed_scene"],
        "software": ["software_scene", "completed_scene", "project_scene"],
        "professional_service": ["operation_scene", "completed_scene", "project_scene"],
        "operation_service": ["operation_scene", "equipment_scene", "completed_scene"],
        "construction": ["site_scene", "completed_scene", "equipment_scene"],
        "mixed": ["project_scene", "equipment_scene", "completed_scene"],
    }
    limit = 2 if visual_mode == "mixed" else 3
    output = []
    fact_context = outline_visual_context(outline)
    for image_type in archetype_scenes.get(archetype, archetype_scenes["mixed"])[:limit]:
        title, capability_id, phrases, prompt = scene_library[image_type]
        if fact_context:
            prompt += f" 项目事实线索：{fact_context}。只选择这些线索能够支持的对象和场景，不得自行替换为其他行业。"
        output.append(
            planned_artifact(
                {
                    "kind": f"gpt_{image_type}",
                    "imageType": image_type,
                    "title": title,
                    "source": "sub2api-gpt-image",
                    "userPrompt": prompt,
                    "estimatedCostUsd": 0.201,
                },
                outline,
                capability_id,
                phrases,
            )
        )
    return output


def render_svg_png(svg, target):
    document = fitz.open(stream=svg.encode("utf-8"), filetype="svg")
    try:
        page = document[0]
        pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
        pixmap.save(str(target))
    finally:
        document.close()


def process_artifact(job):
    job_id, project_id = job["jobId"], job["projectId"]
    try:
        with psycopg.connect(DB) as conn:
            conn.execute(
                "UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s",
                (job_id,),
            )
            row = conn.execute(
                "SELECT content FROM outlines WHERE project_id=%s AND status='ready'",
                (project_id,),
            ).fetchone()
            conn.commit()
        if not row:
            raise ValueError("项目大纲尚未就绪")
        outline = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
        generated = []
        specs = artifact_specs(outline)
        planned_kinds = [spec["kind"] for spec in specs]
        settings = outline.get("generationSettings") or {}
        visual_mode = str(settings.get("visualMode") or "diagrams")
        image_cost_confirmed = settings.get("visualImageCostConfirmed") is True
        scene_specs = realistic_scene_specs(outline, visual_mode) if image_cost_confirmed else []
        planned_scene_kinds = [spec["kind"] for spec in scene_specs]
        queued_images = []
        with psycopg.connect(DB) as conn:
            conn.execute(
                """DELETE FROM document_artifacts
                   WHERE project_id=%s
                     AND metadata->>'generator'='deterministic-svg-v1'
                     AND NOT (kind = ANY(%s::text[]))""",
                (project_id, planned_kinds),
            )
            conn.execute(
                """DELETE FROM document_artifacts
                   WHERE project_id=%s
                     AND metadata->>'imageNature'='ai-realistic-illustration'
                     AND NOT (kind = ANY(%s::text[]))""",
                (project_id, planned_scene_kinds),
            )
            for spec in specs:
                stem = f"{project_id}-{spec['kind']}"
                svg_target = ARTIFACT_DIR / f"{stem}.svg"
                png_target = ARTIFACT_DIR / f"{stem}.png"
                svg_target.write_text(spec["svg"], encoding="utf-8")
                render_svg_png(spec["svg"], png_target)
                artifact_id = str(uuid.uuid4())
                row_id = conn.execute(
                    """INSERT INTO document_artifacts(
                           id,project_id,kind,title,status,svg_path,png_path,metadata,error_message
                       ) VALUES(%s,%s,%s,%s,'ready',%s,%s,%s::jsonb,NULL)
                       ON CONFLICT(project_id,kind) DO UPDATE SET
                           title=excluded.title,status='ready',svg_path=excluded.svg_path,
                           png_path=excluded.png_path,
                           metadata=coalesce(document_artifacts.metadata,'{}'::jsonb) || excluded.metadata,
                           error_message=NULL,updated_at=now()
                       RETURNING id""",
                    (
                        artifact_id,
                        project_id,
                        spec["kind"],
                        spec["title"],
                        str(Path("artifacts") / svg_target.name),
                        str(Path("artifacts") / png_target.name),
                        json.dumps(
                            {
                                "source": spec["source"],
                                "generator": "deterministic-svg-v1",
                                "plannerVersion": spec.get("plannerVersion"),
                                "targetPath": spec.get("targetPath"),
                                "targetTitle": spec.get("targetTitle"),
                                "placementMode": spec.get("placementMode"),
                            },
                            ensure_ascii=False,
                        ),
                    ),
                ).fetchone()[0]
                generated.append({"id": str(row_id), "kind": spec["kind"], "bytes": png_target.stat().st_size})
            for scene in scene_specs:
                existing = conn.execute(
                    "SELECT id,status FROM document_artifacts WHERE project_id=%s AND kind=%s",
                    (project_id, scene["kind"]),
                ).fetchone()
                if existing and existing[1] == "ready" and not job.get("regenerateImages"):
                    continue
                artifact_id = str(existing[0]) if existing else str(uuid.uuid4())
                metadata = {
                    "source": scene["source"],
                    "generator": "gpt-image-v1",
                    "requestedModel": IMAGE_MODEL,
                    "estimatedCostUsd": scene["estimatedCostUsd"],
                    "plannerVersion": scene.get("plannerVersion"),
                    "targetPath": scene.get("targetPath"),
                    "targetTitle": scene.get("targetTitle"),
                    "placementMode": scene.get("placementMode"),
                    "visualMode": visual_mode,
                    "imageNature": "ai-realistic-illustration",
                }
                conn.execute(
                    """INSERT INTO document_artifacts(
                           id,project_id,kind,title,status,metadata,error_message
                       ) VALUES(%s,%s,%s,%s,'generating',%s::jsonb,NULL)
                       ON CONFLICT(project_id,kind) DO UPDATE SET
                           title=excluded.title,status='generating',
                           metadata=coalesce(document_artifacts.metadata,'{}'::jsonb) || excluded.metadata,
                           error_message=NULL,updated_at=now()""",
                    (
                        artifact_id,
                        project_id,
                        scene["kind"],
                        scene["title"],
                        json.dumps(metadata, ensure_ascii=False),
                    ),
                )
                image_job_id = str(uuid.uuid4())
                conn.execute(
                    "INSERT INTO jobs(id,project_id,type,status) VALUES(%s,%s,'image_artifact','queued')",
                    (image_job_id, project_id),
                )
                queued_images.append(
                    {
                        "jobId": image_job_id,
                        "projectId": project_id,
                        "type": "image_artifact",
                        "artifactId": artifact_id,
                        "imageType": scene["imageType"],
                        "userPrompt": scene["userPrompt"],
                    }
                )
            conn.execute("UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s", (job_id,))
            conn.commit()
        for image_job in reversed(queued_images):
            R.lpush("ai_bid:jobs", json.dumps(image_job, ensure_ascii=False))
        print(json.dumps({"event": "artifacts_generated", "projectId": project_id, "items": generated}, ensure_ascii=False), flush=True)
    except Exception as exc:
        with psycopg.connect(DB) as conn:
            conn.execute(
                "UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s",
                (str(exc)[:500], job_id),
            )
            conn.commit()
        traceback.print_exc()


def effect_image_prompt(outline, image_type, user_prompt=""):
    profile = outline.get("projectProfile") or {}
    project_name = artifact_text(outline.get("projectName") or profile.get("projectType") or "本项目", 80)
    scope = [artifact_text(item, 60) for item in (profile.get("scope") or [])[:5]]
    objectives = [artifact_text(item, 80) for item in (profile.get("objectives") or [])[:3]]
    context = "；".join(scope + objectives) or outline_visual_context(outline)
    prompts = {
        "culture_wall": (
            f"为“{project_name}”制作一张专业、真实、可落地的文化墙空间效果图。"
            "画面为正视略带透视的室内墙面，现代、庄重、简洁，层次清楚，"
            "预留标题区、主题图区和成果展示区；使用红色、暖金色与中性色，"
            "不要生成可辨识的文字、徽标、人物肖像和具体单位标志，文字位置以干净色块占位，便于后期排版。"
        ),
        "project_scene": (
            f"为“{project_name}”制作一张项目建成后的专业场景概念效果图。"
            f"项目范围线索：{context}。画面真实、整洁、可信，体现完成后的使用状态，"
            "不出现品牌商标、可辨识文字、施工尺寸和无法从项目资料确认的设备型号；"
            "作为投标方案概念示意，不作为施工图。"
        ),
        "solution_concept": (
            f"为“{project_name}”制作一张高端政府采购投标方案概念视觉图。"
            f"项目范围线索：{context}。采用现代蓝绿色、清晰空间层次和专业技术氛围，"
            "突出协同、智能、可靠与交付闭环；无文字、无品牌、无夸张科幻元素。"
        ),
        "overall_architecture": (
            f"为“{project_name}”制作一张横向16:9的高端政府采购投标文件信息图，主题为项目总体交付架构。"
            f"项目范围线索：{context}。使用象牙白背景、深蓝与青绿色、少量暖金色点缀，扁平矢量风格，"
            "结构清晰、留白充分、正式克制，呈现从总体目标向多个专业工作流协同交付的架构关系。"
            "只允许出现以下中文：项目总体交付架构图、总体目标、智慧安防、净水设备、校园文化墙、"
            "智慧教学、AI智能考勤、心理测试室、协同交付。文字必须准确，不得添加其他文字、英文、数字、徽标或水印。"
        ),
        "implementation_route": (
            f"为“{project_name}”制作一张横向16:9的高端政府采购投标文件流程信息图，主题为项目实施路线。"
            "使用象牙白背景、深蓝与青绿色、少量暖金色点缀，现代扁平矢量风格，流程从左到右，"
            "节点之间有清楚的箭头、里程碑和交付感，正式、专业、适合直接插入投标文件。"
            "只允许出现以下中文：项目实施路线图、项目启动、深化设计、采购备货、安装实施、"
            "联调测试、培训试运行、验收交付。文字必须准确，不得添加其他文字、英文、数字、徽标或水印。"
        ),
        "quality_closed_loop": (
            f"为“{project_name}”制作一张横向16:9的高端政府采购投标文件闭环信息图，主题为全过程质量控制。"
            "使用象牙白背景、深蓝与青绿色、少量暖金色点缀，现代扁平矢量风格，以环形或回流箭头表达持续改进，"
            "层次清晰、留白充分、正式克制，适合直接插入投标文件。"
            "只允许出现以下中文：全过程质量控制闭环图、需求基线、方案审查、过程检查、成果复核、"
            "整改验证、验收归档、反馈纠偏与持续优化。文字必须准确，不得添加其他文字、英文、数字、徽标或水印。"
        ),
        "operation_scene": (
            f"为“{project_name}”制作一张横向构图的项目现场作业写实示意图。"
            f"项目范围线索：{context}。采用真实商业摄影质感，自然光线，人物动作专业，"
            "同时呈现作业人员、通用工具和整洁现场环境。不得出现品牌、单位标志、可辨识文字、"
            "具体型号和夸张数量；图片仅作为投标方案实施场景示意，不作为既有业绩或实物证明。"
        ),
        "equipment_scene": (
            f"为“{project_name}”制作一张横向构图的设备、工具和物资写实示意图。"
            f"项目范围线索：{context}。采用真实产品摄影与现场陈列风格，物品排列专业、整洁、可信，"
            "背景为与项目交付形态相符的工作场地。不得出现品牌、商标、型号、可辨识文字和未经确认的数量；"
            "图片仅为配置思路示意，不作为实际供货品牌、库存或数量承诺。"
            "仅选择项目事实线索明确支持的工具与物资；若线索属于运营服务、安保、保洁或物业维修，"
            "应优先展示清洁推车、清扫工具、警示设施、巡检用品、个人防护和常用维修工具，"
            "不得生成工厂仓库、重型发电机、大型线缆盘或无关工业设备。"
        ),
        "installation_scene": (
            f"为“{project_name}”制作一张横向构图的设备安装与调试现场写实示意图。"
            "展示技术人员在规范防护条件下进行安装、接线或调试，现场整洁有序，动作符合专业常识。"
            "不得出现品牌、单位标志、可辨识文字、设备型号和虚构参数；仅作为实施场景示意。"
        ),
        "software_scene": (
            f"为“{project_name}”制作一张横向构图的软件系统应用写实示意图。"
            "展示工作人员在电脑终端进行系统操作、测试和协同，办公环境真实克制。"
            "屏幕只呈现不可辨识的抽象界面色块，不出现品牌、文字、具体功能名称和虚构数据；仅作为应用场景示意。"
        ),
        "site_scene": (
            f"为“{project_name}”制作一张横向构图的工程现场实施写实示意图。"
            "展示规范围护、个人防护、材料分类堆放和有序施工状态，采用真实工程摄影质感。"
            "不得出现单位标志、品牌、可辨识文字、虚构工程量和未确认工艺参数；仅作为施工组织场景示意。"
        ),
        "completed_scene": (
            f"为“{project_name}”制作一张横向构图的项目完成效果写实示意图。"
            f"项目范围线索：{context}。展示项目完成后的整洁、规范、可使用状态，体现交付成果与使用场景。"
            "不得出现品牌、单位标志、可辨识文字、未经确认的设施型号和数量；仅作为预期效果示意。"
        ),
        "custom": (
            f"为“{project_name}”制作专业投标文件配图。"
            "图片应真实、克制、可用于正式文档，无品牌、无水印、无可辨识文字。"
        ),
    }
    prompt = prompts.get(image_type, prompts["custom"])
    if str(user_prompt or "").strip():
        prompt += " 用户补充要求：" + artifact_text(user_prompt, 800)
    return prompt


def process_image_artifact(job):
    job_id, project_id = job["jobId"], job["projectId"]
    artifact_id = job.get("artifactId")
    image_type = job.get("imageType") or "project_scene"
    run_id = str(uuid.uuid4())
    started = time.monotonic()
    input_tokens = output_tokens = 0
    actual_cost = None
    try:
        if not IMAGE_KEY:
            raise ValueError("sub2api 图片密钥未配置")
        with psycopg.connect(DB) as conn:
            conn.execute(
                "UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s",
                (job_id,),
            )
            if artifact_id:
                conn.execute(
                    "UPDATE document_artifacts SET status='generating',error_message=NULL,updated_at=now() WHERE id=%s AND project_id=%s",
                    (artifact_id, project_id),
                )
            row = conn.execute(
                """SELECT p.name,o.content FROM projects p
                   JOIN outlines o ON o.project_id=p.id
                   WHERE p.id=%s AND o.status='ready'""",
                (project_id,),
            ).fetchone()
            conn.commit()
        if not row:
            raise ValueError("项目大纲尚未就绪")
        outline = row[1] if isinstance(row[1], dict) else json.loads(row[1])
        prompt = effect_image_prompt(outline, image_type, job.get("userPrompt") or "")
        image_size = "1536x1024" if image_type in {
            "overall_architecture", "implementation_route", "quality_closed_loop"
        } else "1024x1024"
        body = json.dumps({
            "model": IMAGE_MODEL,
            "prompt": prompt,
            "size": image_size,
            "quality": "low",
            "output_format": "png",
            "n": 1,
        }, ensure_ascii=False).encode("utf-8")
        request = urllib.request.Request(
            IMAGE_URL,
            data=body,
            headers={"Authorization": f"Bearer {IMAGE_KEY}", "Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(request, timeout=210) as response:
            client_request_id = response.headers.get("X-Client-Request-Id")
            payload = json.loads(response.read().decode("utf-8"))
        item = (payload.get("data") or [{}])[0]
        encoded = item.get("b64_json") or item.get("b64")
        if not encoded:
            raise ValueError("图片模型未返回图像数据")
        image_bytes = base64.b64decode(encoded, validate=True)
        if len(image_bytes) < 10000:
            raise ValueError("图片模型返回的数据不完整")
        ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
        stem = f"{project_id}-gpt-{image_type}"
        png_target = ARTIFACT_DIR / f"{stem}.png"
        png_target.write_bytes(image_bytes)
        usage = payload.get("usage") or {}
        input_tokens = int(usage.get("input_tokens") or usage.get("prompt_tokens") or 0)
        output_tokens = int(usage.get("output_tokens") or usage.get("completion_tokens") or 0)
        billing = fetch_billing(client_request_id)
        if billing:
            input_tokens, output_tokens, actual_cost = billing
        metadata = {
            "source": "sub2api",
            "generator": "gpt-image-v1",
            "requestedModel": IMAGE_MODEL,
            "returnedModel": payload.get("model") or IMAGE_MODEL,
            "requestedQuality": "low",
            "returnedQuality": payload.get("quality"),
            "returnedSize": payload.get("size"),
            "actualCostUsd": actual_cost,
            "prompt": prompt,
        }
        title_map = {
            "culture_wall": "文化墙概念效果图",
            "project_scene": "项目建成场景概念图",
            "solution_concept": "项目方案概念视觉图",
            "operation_scene": "项目现场作业写实示意图",
            "equipment_scene": "项目设备物资写实示意图",
            "installation_scene": "设备安装实施写实示意图",
            "software_scene": "软件系统应用写实示意图",
            "site_scene": "工程现场实施写实示意图",
            "completed_scene": "项目完成效果写实示意图",
            "custom": "自定义项目效果图",
            "overall_architecture": "项目总体交付架构图",
            "implementation_route": "项目实施路线图",
            "quality_closed_loop": "全过程质量控制闭环图",
        }
        kind = image_type if image_type in {
            "overall_architecture", "implementation_route", "quality_closed_loop"
        } else f"gpt_{image_type}"
        with psycopg.connect(DB) as conn:
            if artifact_id:
                updated = conn.execute(
                    """UPDATE document_artifacts SET title=%s,status='ready',svg_path=NULL,
                           png_path=%s,
                           metadata=coalesce(metadata,'{}'::jsonb) || %s::jsonb,
                           error_message=NULL,updated_at=now()
                       WHERE id=%s AND project_id=%s RETURNING id""",
                    (
                        title_map.get(image_type, title_map["custom"]),
                        str(Path("artifacts") / png_target.name),
                        json.dumps(metadata, ensure_ascii=False),
                        artifact_id,
                        project_id,
                    ),
                ).fetchone()
            else:
                updated = None
            if not updated:
                artifact_id = str(uuid.uuid4())
                conn.execute(
                    """INSERT INTO document_artifacts(
                           id,project_id,kind,title,status,png_path,metadata
                       ) VALUES(%s,%s,%s,%s,'ready',%s,%s::jsonb)
                       ON CONFLICT(project_id,kind) DO UPDATE SET
                           title=excluded.title,status='ready',png_path=excluded.png_path,
                           metadata=coalesce(document_artifacts.metadata,'{}'::jsonb) || excluded.metadata,
                           error_message=NULL,updated_at=now()""",
                    (
                        artifact_id,
                        project_id,
                        kind,
                        title_map.get(image_type, title_map["custom"]),
                        str(Path("artifacts") / png_target.name),
                        json.dumps(metadata, ensure_ascii=False),
                    ),
                )
            conn.execute("UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s", (job_id,))
            conn.commit()
        save_run(
            run_id, project_id, job_id, 1, "succeeded", input_tokens, output_tokens,
            int((time.monotonic() - started) * 1000), 0, actual_cost=actual_cost,
            run_type="image_generation", model_override=payload.get("model") or IMAGE_MODEL,
        )
        print(json.dumps({
            "event": "image_artifact_generated",
            "projectId": project_id,
            "imageType": image_type,
            "bytes": len(image_bytes),
            "actualCostUsd": actual_cost,
        }, ensure_ascii=False), flush=True)
    except Exception as exc:
        with psycopg.connect(DB) as conn:
            if artifact_id:
                conn.execute(
                    "UPDATE document_artifacts SET status='failed',error_message=%s,updated_at=now() WHERE id=%s AND project_id=%s",
                    (str(exc)[:500], artifact_id, project_id),
                )
            conn.execute(
                "UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s",
                (str(exc)[:500], job_id),
            )
            conn.commit()
        save_run(
            run_id, project_id, job_id, 1, "failed", input_tokens, output_tokens,
            int((time.monotonic() - started) * 1000), 0, error=exc,
            actual_cost=actual_cost, run_type="image_generation", model_override=IMAGE_MODEL,
        )
        traceback.print_exc()


def set_docx_font(style, name, size, bold=None, color=None):
    style.font.name = name; style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if bold is not None: style.font.bold = bold
    if color: style.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    paragraph.add_run(" 页")


def add_rich_paragraph(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    add_rich_runs(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    if not style:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)
    return paragraph


def add_rich_runs(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part: continue
        if part.startswith("**") and part.endswith("**"): paragraph.add_run(part[2:-2]).bold = True
        else: paragraph.add_run(part)
    return paragraph


def markdown_table_cells(line):
    value = str(line or "").strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def markdown_table_separator(line):
    cells = markdown_table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", cell or "") for cell in cells)


def word_table_widths(rows, total_width=9360):
    column_count = max((len(row) for row in rows), default=1)
    weights = []
    for index in range(column_count):
        lengths = [
            min(40, max(2, len(re.sub(r"\*\*", "", row[index]))))
            for row in rows
            if index < len(row)
        ]
        weights.append(max(lengths or [4]))
    minimum = 720 if column_count >= 6 else 960
    base = minimum * column_count
    if base >= total_width:
        return [total_width // column_count] * column_count
    remaining = total_width - base
    weight_total = sum(weights) or column_count
    widths = [minimum + int(remaining * weight / weight_total) for weight in weights]
    widths[-1] += total_width - sum(widths)
    return widths


def set_word_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        row_properties.append(no_split)
        for index, cell in enumerate(row.cells):
            cell_width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(cell_width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell_properties = cell._tc.get_or_add_tcPr()
            tc_width = cell_properties.get_or_add_tcW()
            tc_width.set(qn("w:w"), str(cell_width))
            tc_width.set(qn("w:type"), "dxa")
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for edge, value in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
                element = margins.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    margins.append(element)
                element.set(qn("w:w"), str(value))
                element.set(qn("w:type"), "dxa")


def add_markdown_table(doc, rows, has_header=True):
    column_count = max((len(row) for row in rows), default=0)
    if not column_count:
        return None
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(normalized):
        for column_index, value in enumerate(values):
            cell = table.rows[row_index].cells[column_index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if row_index == 0 and has_header:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_rich_runs(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9.5)
                run.font.name = "SimSun"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if row_index == 0 and has_header:
                    run.bold = True
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                    run.font.color.rgb = RGBColor(31, 58, 95)
            if row_index == 0 and has_header:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E8EEF5")
                cell._tc.get_or_add_tcPr().append(shading)
    if has_header:
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        header_properties.append(repeat_header)
    set_word_table_geometry(table, word_table_widths(normalized))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_markdown_content(doc, text, skip_title=None):
    title_skipped = False
    lines = str(text or "").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and markdown_table_separator(lines[index + 1].strip()):
            rows = [markdown_table_cells(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(markdown_table_cells(lines[index].strip()))
                index += 1
            add_markdown_table(doc, rows, has_header=True)
            continue
        heading_text = re.sub(r"^#{1,3}\s+", "", line).strip()
        if not title_skipped and skip_title and line.startswith("#") and norm(heading_text) == norm(skip_title):
            title_skipped = True
            index += 1
            continue
        if line.startswith("#### "): add_rich_paragraph(doc, line[5:], "Heading 3")
        elif line.startswith("### "): add_rich_paragraph(doc, line[4:], "Heading 3")
        elif line.startswith("## "): add_rich_paragraph(doc, line[3:], "Heading 2")
        elif line.startswith("# "): add_rich_paragraph(doc, line[2:], "Heading 2")
        elif re.match(r"^\d+\.\s+", line): add_rich_paragraph(doc, re.sub(r"^\d+\.\s+", "", line), "List Number")
        elif line.startswith(("- ", "* ")): add_rich_paragraph(doc, line[2:], "List Bullet")
        else: add_rich_paragraph(doc, line)
        index += 1


def flatten_outline(nodes, prefix=()):
    output = []
    for index, node in enumerate(nodes, 1):
        number = prefix + (index,)
        output.append((number, node))
        output.extend(flatten_outline(node.get("children") or [], number))
    return output


def outline_title_path(chapters, number):
    titles = []
    nodes = chapters or []
    for index in number:
        if index < 1 or index > len(nodes):
            break
        node = nodes[index - 1]
        titles.append(str(node.get("title") or ""))
        nodes = node.get("children") or []
    return titles


def artifact_section_score(artifact, number, node, chapters):
    if node.get("children"):
        return -1
    kind = str(artifact.get("kind") or "")
    artifact_title = str(artifact.get("title") or "")
    node_title = str(node.get("title") or "")
    context = " ".join(outline_title_path(chapters, number))
    score = 0
    fixed_rules = {
        "overall_architecture": (
            ("总体建设目标", 180),
            ("总体方案", 150),
            ("总体理解", 120),
            ("项目概述", 90),
        ),
        "implementation_route": (
            ("整体实施进度计划", 200),
            ("实施进度", 170),
            ("进度计划", 150),
            ("组织与实施", 100),
        ),
        "quality_closed_loop": (
            ("质量控制与风险管理", 200),
            ("质量控制", 170),
            ("质量管理", 150),
            ("验收保障", 120),
        ),
        "service_scope_map": (
            ("服务范围与需求理解", 220),
            ("服务范围", 190),
            ("需求理解", 170),
            ("项目背景", 80),
        ),
        "service_operation_cycle": (
            ("工作方法与协同机制", 220),
            ("实施路径与工作流程", 200),
            ("服务实施", 170),
            ("工作流程", 150),
        ),
        "organization_responsibility": (
            ("项目组织架构", 240),
            ("岗位职责与人员安排", 190),
            ("组织架构", 180),
            ("人员安排", 130),
        ),
        "inspection_rectification": (
            ("过程检查与成果审核", 240),
            ("质量控制", 180),
            ("过程检查", 170),
            ("整改", 140),
        ),
        "emergency_response": (
            ("风险识别与应急预案", 250),
            ("应急预案", 210),
            ("风险识别", 180),
            ("安全管理", 100),
        ),
        "gpt_culture_wall": (
            ("设计效果图", 180),
            ("文化墙", 140),
        ),
        "gpt_project_scene": (
            ("总体建设目标", 180),
            ("建设效果", 160),
            ("项目概述", 100),
        ),
        "gpt_solution_concept": (
            ("总体方案", 180),
            ("总体建设目标", 160),
            ("项目概述", 100),
        ),
    }
    for phrase, weight in fixed_rules.get(kind, ()):
        if phrase in node_title:
            score += weight
        elif phrase in context:
            score += weight // 2
    shared_keywords = (
        "文化墙", "效果图", "总体", "架构", "实施", "进度", "质量", "风险",
        "服务", "范围", "作业", "组织", "岗位", "巡检", "整改", "应急",
        "安防", "净水", "教学", "考勤", "测试室", "布局", "拓扑", "原理",
    )
    for keyword in shared_keywords:
        if keyword in artifact_title and keyword in context:
            score += 35
    return score


def assign_artifacts_to_sections(outline_content, artifacts):
    chapters = outline_content.get("chapters") or []
    flat = flatten_outline(chapters)
    available_numbers = {number for number, _ in flat}
    assigned = {}
    unmatched = []
    for artifact in artifacts or []:
        metadata = artifact.get("metadata") or {}
        target_path = metadata.get("targetPath")
        if (
            isinstance(target_path, list)
            and target_path
            and all(isinstance(item, int) and item >= 0 for item in target_path)
        ):
            target_number = tuple(item + 1 for item in target_path)
            if target_number in available_numbers:
                assigned.setdefault(target_number, []).append(artifact)
                continue
        candidates = [
            (artifact_section_score(artifact, number, node, chapters), number)
            for number, node in flat
        ]
        score, number = max(candidates, default=(-1, ()), key=lambda item: item[0])
        if score > 0:
            assigned.setdefault(number, []).append(artifact)
        else:
            unmatched.append(artifact)
    return assigned, unmatched


def add_docx_artifact(doc, artifact, figure_number):
    image_path = Path(artifact["path"])
    if not image_path.exists():
        return False
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(str(image_path), width=Inches(6.15))
    caption = doc.add_paragraph(f"图 {figure_number}  {artifact['title']}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.space_after = Pt(10)
    for run in caption.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 112, 126)
    return True


def add_parameter_matrix_table(doc, parameter_items):
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("序号", "产品/设备", "标识", "招标技术要求", "投标响应", "偏离情况", "证明材料")
    header_row = table.rows[0]
    for index, label in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8EEF5")
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    header_properties = header_row._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:tblHeader"))
    for item in parameter_items:
        cells = table.add_row().cells
        values = (
            str(item.get("itemIndex") or ""),
            f"{item.get('productNo') or ''} {item.get('productName') or ''}".strip(),
            str(item.get("marker") or "一般"),
            str(item.get("requirement") or ""),
            str(item.get("responseValue") or "【待填写：拟投产品具体参数】"),
            str(item.get("deviationStatus") or "待核对"),
            str(item.get("evidenceReference") or item.get("proofRequirement") or "【待补证明材料】"),
        )
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    set_word_table_geometry(
        table, [560, 1160, 620, 2450, 1900, 960, 1710]
    )
    return table


def configure_bid_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)


def set_run_typeface(run, east_asia, latin=None):
    latin = latin or east_asia
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_bid_header_footer(section, project_name):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].clear()
    table = header.add_table(rows=1, cols=2, width=Cm(16.2))
    set_word_table_geometry(table, [7200, 2160])
    for cell in table.rows[0].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        borders = cell_properties.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            cell_properties.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            borders.append(border)
    left = table.cell(0, 0).paragraphs[0]
    left.text = str(project_name or "")[:42]
    right = table.cell(0, 1).paragraphs[0]
    right.text = "投标文件"
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for paragraph in (left, right):
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_typeface(run, "宋体", "SimSun")
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(96, 108, 124)
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)
    for run in footer.runs:
        set_run_typeface(run, "宋体", "SimSun")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(96, 108, 124)


def add_chapter_opener(doc, number, node):
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(58)
    label.paragraph_format.space_after = Pt(10)
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = label.add_run(f"CHAPTER {number:02d}")
    set_run_typeface(run, "微软雅黑", "Arial")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(82, 110, 138)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(15)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(f"第{number}章  {str(node.get('title') or '')}")
    set_run_typeface(run, "黑体", "Microsoft YaHei")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 58, 95)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8FA5BA")
    border.append(bottom)
    rule._p.get_or_add_pPr().append(border)
    description = str(node.get("description") or "").strip()
    if description:
        lead = doc.add_paragraph()
        lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lead.paragraph_format.left_indent = Cm(0.7)
        lead.paragraph_format.right_indent = Cm(0.7)
        lead.paragraph_format.line_spacing = 1.5
        lead.paragraph_format.space_after = Pt(18)
        run = lead.add_run(description)
        set_run_typeface(run, "宋体", "SimSun")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(75, 86, 100)


def build_docx(project_name, outline_content, target, parameter_items=None, artifacts=None):
    doc = Document()
    cover_section = doc.sections[0]
    configure_bid_section(cover_section)
    cover_section.header.is_linked_to_previous = False
    cover_section.footer.is_linked_to_previous = False
    cover_section.header.paragraphs[0].clear()
    cover_section.footer.paragraphs[0].clear()

    normal = doc.styles["Normal"]
    set_docx_font(normal, "SimSun", 12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.widow_control = True
    for name, size, color, before, after in (
        ("Heading 1", 18, "1F3A5F", 20, 12),
        ("Heading 2", 15, "1F3A5F", 16, 8),
        ("Heading 3", 13, "334E68", 12, 6),
    ):
        style = doc.styles[name]
        set_docx_font(style, "Microsoft YaHei", size, True, color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    cover_top = doc.add_paragraph()
    cover_top.paragraph_format.space_before = Pt(92)
    cover_top.paragraph_format.space_after = Pt(22)
    marker = doc.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker.paragraph_format.space_after = Pt(22)
    marker_run = marker.add_run("投  标  文  件")
    set_run_typeface(marker_run, "黑体", "Microsoft YaHei")
    marker_run.bold = True
    marker_run.font.size = Pt(15)
    marker_run.font.color.rgb = RGBColor(82, 110, 138)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(15)
    title_run = title.add_run(project_name)
    set_run_typeface(title_run, "黑体", "Microsoft YaHei")
    title_run.bold = True
    title_run.font.size = Pt(25)
    title_run.font.color.rgb = RGBColor(31, 58, 95)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(58)
    sub_run = subtitle.add_run("响应文件")
    set_run_typeface(sub_run, "黑体", "Microsoft YaHei")
    sub_run.bold = True
    sub_run.font.size = Pt(17)
    sub_run.font.color.rgb = RGBColor(75, 86, 100)
    for label in ("投标人（盖章）：【待补充】", "项目编号：【待补充】", "编制日期：【待补充】"):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(12)
        run = paragraph.add_run(label)
        set_run_typeface(run, "宋体", "SimSun")
        run.font.size = Pt(12)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_bid_section(main_section)
    add_bid_header_footer(main_section, project_name)
    doc.add_heading("目录", level=1)
    flat = flatten_outline(outline_content.get("chapters") or [])
    section_artifacts, unmatched_artifacts = assign_artifacts_to_sections(
        outline_content,
        artifacts or [],
    )
    for number, node in flat:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65 * max(len(number) - 1, 0))
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(
            ".".join(map(str, number)) + "  " + str(node.get("title") or "")
        )
        set_run_typeface(run, "黑体" if len(number) == 1 else "宋体", "Microsoft YaHei" if len(number) == 1 else "SimSun")
        run.bold = len(number) == 1
        run.font.size = Pt(11 if len(number) == 1 else 10.5)
    doc.add_page_break()

    figure_number = 1
    chapter_seen = 0
    for number, node in flat:
        if len(number) == 1:
            if chapter_seen:
                doc.add_page_break()
            chapter_seen += 1
            add_chapter_opener(doc, number[0], node)
            continue
        level = min(len(number) - 1, 3)
        doc.add_heading(
            ".".join(map(str, number)) + " " + str(node.get("title") or ""),
            level=level,
        )
        if (
            parameter_items
            and not node.get("children")
            and "产品技术参数响应表" in str(node.get("title") or "")
        ):
            add_parameter_matrix_table(doc, parameter_items)
        elif node.get("content") and not node.get("children"):
            add_markdown_content(
                doc, node["content"], str(node.get("title") or "")
            )
        elif not node.get("children"):
            add_rich_paragraph(doc, "【本章节正文待生成】")
        for artifact in section_artifacts.get(number, []):
            if add_docx_artifact(doc, artifact, figure_number):
                figure_number += 1
    if unmatched_artifacts:
        doc.add_page_break()
        add_chapter_opener(
            doc,
            chapter_seen + 1,
            {"title": "项目图示汇编", "description": "本章汇集尚未指定正文插入位置的项目图示，供定稿时结合章节内容复核使用。"},
        )
        for artifact in unmatched_artifacts:
            doc.add_heading(artifact["title"], level=1)
            if add_docx_artifact(doc, artifact, figure_number):
                figure_number += 1
    props = doc.core_properties
    props.title = project_name
    props.subject = "投标响应文件"
    props.author = "投标人"
    props.keywords = "投标文件,响应文件"
    doc.save(target)


def process_export(job):
    job_id, project_id = job["jobId"], job["projectId"]
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE jobs SET status='running',attempts=attempts+1,started_at=now() WHERE id=%s", (job_id,))
        conn.execute("UPDATE document_exports SET status='running',updated_at=now() WHERE project_id=%s", (project_id,))
        project = conn.execute("SELECT name FROM projects WHERE id=%s", (project_id,)).fetchone()
        outline = conn.execute("SELECT content FROM outlines WHERE project_id=%s AND status='ready'", (project_id,)).fetchone()
        parameter_rows = conn.execute(
            """SELECT item_index,product_no,product_name,parameter_no,marker,
                      requirement_text,source_page,proof_requirement,response_value,
                      deviation_status,evidence_reference
               FROM technical_parameter_items
               WHERE project_id=%s ORDER BY item_index""",
            (project_id,),
        ).fetchall()
        artifact_rows = conn.execute(
            """SELECT kind,title,png_path,metadata FROM document_artifacts
               WHERE project_id=%s AND status='ready' AND png_path IS NOT NULL
               ORDER BY created_at,kind""",
            (project_id,),
        ).fetchall()
        conn.commit()
    try:
        if not project or not outline: raise ValueError("项目或大纲不存在")
        EXPORT_DIR.mkdir(parents=True, exist_ok=True)
        safe_name = re.sub(r'[\\/:*?"<>|]+', "_", project[0]).strip()[:80] or "投标文件"
        file_name = safe_name + "-投标文件.docx"
        stored_name = f"{project_id}-{int(time.time())}.docx"
        target = EXPORT_DIR / stored_name
        content = outline[0] if isinstance(outline[0], dict) else json.loads(outline[0])
        parameter_items = [
            {
                "itemIndex": row[0],
                "productNo": row[1],
                "productName": row[2],
                "parameterNo": row[3],
                "marker": row[4],
                "requirement": row[5],
                "sourcePage": row[6],
                "proofRequirement": row[7],
                "responseValue": row[8],
                "deviationStatus": row[9],
                "evidenceReference": row[10],
            }
            for row in parameter_rows
        ]
        artifacts = [
            {
                "kind": row[0],
                "title": row[1],
                "path": UP.parent / row[2],
                "metadata": row[3] if isinstance(row[3], dict) else json.loads(row[3] or "{}"),
            }
            for row in artifact_rows
        ]
        build_docx(project[0], content, target, parameter_items, artifacts)
        with psycopg.connect(DB) as conn:
            conn.execute("UPDATE document_exports SET status='ready',file_name=%s,stored_path=%s,error_message=NULL,finished_at=now(),updated_at=now() WHERE project_id=%s", (file_name, str(Path("exports") / stored_name), project_id))
            conn.execute("UPDATE jobs SET status='succeeded',finished_at=now() WHERE id=%s", (job_id,)); conn.commit()
        print(json.dumps({"event": "docx_exported", "projectId": project_id, "file": stored_name, "bytes": target.stat().st_size}), flush=True)
    except Exception as exc:
        with psycopg.connect(DB) as conn:
            conn.execute("UPDATE document_exports SET status='failed',error_message=%s,updated_at=now() WHERE project_id=%s", (str(exc)[:500], project_id))
            conn.execute("UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s", (str(exc)[:500], job_id)); conn.commit()
        traceback.print_exc()


def fail_job(job_id, project_id, exc):
    with psycopg.connect(DB) as conn:
        conn.execute("UPDATE projects SET status='failed',error_message=%s,updated_at=now() WHERE id=%s", (str(exc)[:500], project_id)); conn.execute("UPDATE jobs SET status='failed',error_message=%s,finished_at=now() WHERE id=%s", (str(exc)[:500], job_id)); conn.commit()
    traceback.print_exc()


def process(job):
    if job.get("type") == "export": process_export(job)
    elif job.get("type") == "image_artifact": process_image_artifact(job)
    elif job.get("type") == "artifact": process_artifact(job)
    elif job.get("type") == "chapter_editor": process_chapter_editor(job)
    elif job.get("type") == "section_compare": process_section_compare(job)
    elif job.get("type") == "section": process_section(job)
    elif job.get("type") == "outline": process_outline(job)
    elif job.get("type") == "audit": process_audit(job)
    elif job.get("type") == "extract" or "storedName" not in job: process_extract(job)
    else: process_parse(job)


def promote_delayed_jobs():
    now = time.time()
    jobs = R.zrangebyscore(DELAYED_QUEUE, 0, now, start=0, num=10)
    for raw in jobs:
        if R.zrem(DELAYED_QUEUE, raw):
            R.lpush("ai_bid:jobs", raw)


def recover_interrupted_outline_jobs():
    recovered = []
    with psycopg.connect(DB) as conn:
        rows = conn.execute(
            """
            SELECT j.id,j.project_id,o.content
            FROM jobs j
            LEFT JOIN outlines o ON o.project_id=j.project_id
            WHERE j.type='outline' AND j.status='running'
            ORDER BY j.started_at
            """
        ).fetchall()
        for job_id, project_id, content in rows:
            content = content if isinstance(content, dict) else {}
            checkpoint = content.get("_outlinePlanningCheckpoint") or {}
            outline_mode = (
                "dynamic"
                if checkpoint.get("version") == "dynamic-outline-checkpoint-v1"
                else "dynamic"
            )
            conn.execute(
                """
                UPDATE jobs
                SET status='queued',error_message=NULL,finished_at=NULL
                WHERE id=%s
                """,
                (job_id,),
            )
            recovered.append({
                "jobId": str(job_id),
                "projectId": str(project_id),
                "type": "outline",
                "outlineMode": outline_mode,
            })
        conn.commit()
    for job in recovered:
        R.lpush("ai_bid:jobs", json.dumps(job, ensure_ascii=False))
    if recovered:
        print(json.dumps({
            "event": "interrupted_outline_jobs_recovered",
            "count": len(recovered),
            "jobs": [item["jobId"] for item in recovered],
        }, ensure_ascii=False), flush=True)


def main():
    print(json.dumps({"event": "worker_started", "model": MODEL, "aiConfigured": bool(AI_KEY)}), flush=True)
    recover_interrupted_outline_jobs()
    while True:
        try:
            promote_delayed_jobs()
            item = R.brpop("ai_bid:jobs", timeout=5)
            if item: process(json.loads(item[1]))
        except Exception: traceback.print_exc(); time.sleep(3)


if __name__ == "__main__": main()
